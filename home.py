import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import os
import zipfile
from datetime import date, timedelta

# --- MEMÓRIA (BANCO DE DADOS TEMPORÁRIO) ---
# Nota: Na nuvem gratuita, este arquivo reseta quando o servidor reinicia.
DB_FILE = "db_siarcon.json"

DEFAULTS = {
    "tecnico": [
        "Fornecimento de mão de obra para fabricação de dutos",
        "Fornecimento de mão de obra para preparação e montagem de dutos, incluindo suportações",
        "Fornecimento de mão de obra para isolamento dos dutos",
        "Fornecimento de mão de obra para instalação dos elementos de difusão de ar",
        "Fornecimento de mão de obra para isolamento e instalação das caixas terminais",
        "Fornecimento de mão de obra para fabricação de peças de ajuste em campo",
        "Fornecimento de mão de obra para preparação e execução dos testes de estanqueidade",
        "Fornecimento de mão de obra para marcação de abertura de lajes, paredes e forro"
    ],
    "qualidade": [
        "Higienização e tamponamento dos dutos com plástico filme",
        "Todas as curvas deverão possuir veios internos",
        "Aplicação de selante tipo PU nas emendas longitudinais e TDC",
        "Ao término dos trabalhos, equipe deve realizar limpeza e organização"
    ],
    "docs_sms": [
        "NR-10 (Elétrica)", "NR-12 (Máquinas)", "NR-33 (Espaço Confinado)", "NR-35 (Trabalho em Altura)"
    ]
}

def carregar_banco():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except: return DEFAULTS
    return DEFAULTS

def salvar_banco(dados):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(dados, f, indent=4, ensure_ascii=False)

if 'db' not in st.session_state:
    st.session_state['db'] = carregar_banco()

st.set_page_config(page_title="Gerador SIARCON Pro", page_icon="🏗️", layout="wide")

def aprender_item(categoria, novo_item):
    if novo_item and novo_item not in st.session_state['db'][categoria]:
        st.session_state['db'][categoria].append(novo_item)
        salvar_banco(st.session_state['db'])
        return True
    return False

# --- FUNÇÃO GERADORA DO WORD ---
def gerar_docx(dados):
    document = Document()
    try:
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    except: pass

    # TÍTULO
    title = document.add_heading('Escopo de fornecimento - Rede de dutos', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = document.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {dados['revisao']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()

    # 1. OBJETIVO
    document.add_heading('1. OBJETIVO E RESUMO', level=1)
    table = document.add_table(rows=5, cols=2)
    try: table.style = 'Table Grid'
    except: pass
    
    ref_proj = dados['projetos_ref']
    if dados['nomes_anexos']:
        ref_proj += " | Anexos: " + ", ".join(dados['nomes_anexos'])

    infos = [
        ("Cliente:", dados['cliente']),
        ("Local/Obra:", dados['obra']),
        ("Projetos Ref:", ref_proj),
        ("Fornecedor:", dados['fornecedor']),
        ("Responsável:", dados['responsavel'])
    ]
    for i, (k, v) in enumerate(infos):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = v
    
    document.add_paragraph()
    document.add_paragraph("Resumo do Escopo:", style='List Bullet')
    document.add_paragraph(dados['resumo_escopo'])

    # 2. TÉCNICO
    document.add_heading('2. DETALHAMENTO TÉCNICO', level=1)
    if not dados['itens_tecnicos'] and not dados['tecnico_livre']:
        document.add_paragraph("[Nenhum item técnico selecionado]", style='List Bullet')
    for item in dados['itens_tecnicos']:
        document.add_paragraph(item, style='List Bullet')
    if dados['tecnico_livre']:
        document.add_paragraph(dados['tecnico_livre'], style='List Bullet')

    # 3. QUALIDADE
    document.add_heading('3. ESPECIFICAÇÕES DE QUALIDADE', level=1)
    if not dados['itens_qualidade'] and not dados['qualidade_livre']:
        document.add_paragraph("[Nenhum item de qualidade selecionado]", style='List Bullet')
    for item in dados['itens_qualidade']:
        document.add_paragraph(item, style='List Bullet')
    if dados['qualidade_livre']:
        document.add_paragraph(dados['qualidade_livre'], style='List Bullet')

    # 4. MATRIZ
    document.add_heading('4. MATRIZ DE RESPONSABILIDADES', level=1)
    table_m = document.add_table(rows=1, cols=3)
    try: table_m.style = 'Table Grid'
    except: pass
    
    h = table_m.rows[0].cells
    h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = dados['fornecedor'].upper()
    
    for item, resp in dados['matriz'].items():
        row = table_m.add_row().cells
        row[0].text = item
        if resp == "SIARCON":
            row[1].text = "X"
            row[1].paragraphs[0].alignment = 1
        else:
            row[2].text = "X"
            row[2].paragraphs[0].alignment = 1
    document.add_paragraph()

    # 5. SEGURANÇA
    document.add_heading('5. OBRIGAÇÕES DE SMS', level=1)
    document.add_paragraph("Documentos Padrão Obrigatórios:", style='List Bullet')
    padrao = ["ASO", "Livro de Registro", "Ficha de EPI", "Ordem de Serviço", "Treinamento NR-06", "Documentos Pessoais"]
    for d in padrao: document.add_paragraph(f"- {d}")
    
    if dados['nrs_selecionadas']:
        document.add_paragraph("Documentos Específicos/Adicionais:", style='List Bullet')
        for doc in dados['nrs_selecionadas']:
            document.add_paragraph(f"- {doc}")

    # 6. CRONOGRAMA
    document.add_heading('6. CRONOGRAMA', level=1)
    document.add_paragraph(f"Início: {dados['data_inicio'].strftime('%d/%m/%Y')}")
    document.add_paragraph(f"Integração: {dados['dias_integracao']} dias antes.")
    if dados['data_fim']:
        document.add_paragraph(f"Término Previsto: {dados['data_fim'].strftime('%d/%m/%Y')}")

    # 7. OBS
    if dados['obs_gerais']:
        document.add_heading('7. OBSERVAÇÕES', level=1)
        document.add_paragraph(dados['obs_gerais'])

    # 8. COMERCIAL
    document.add_heading('8. COMERCIAL', level=1)
    document.add_paragraph(f"Valor Global: {dados['valor_total']}")
    document.add_paragraph(f"Condição Pagamento: {dados['condicao_pgto']}")
    if dados['info_comercial']: document.add_paragraph(dados['info_comercial'])
    
    document.add_paragraph("_"*60)
    document.add_paragraph(f"DE ACORDO: {dados['fornecedor']}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- FRONT END ---
st.title("🏗️ Gerador de Escopos - SIARCON")
st.caption("Gera documentos .docx e empacota anexos automaticamente.")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["1. Cadastro & Anexos", "2. Técnico & Qualidade", "3. Matriz", "4. SMS", "5. Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        cliente = st.text_input("Cliente", "Hitachi")
        obra = st.text_input("Local", "Guarulhos - SP")
        fornecedor = st.text_input("Fornecedor (Contratada)")
    with c2:
        responsavel = st.text_input("Responsável pelo preenchimento")
        revisao = st.text_input("Revisão", "R-00")
        projetos_ref = st.text_input("Projetos Ref. (Texto)")
    
    resumo_escopo = st.text_area("Resumo do Escopo", "Fornecimento de mão de obra...")
    
    st.markdown("---")
    st.markdown("### 📎 Anexar Projetos")
    arquivos_anexos = st.file_uploader(
        "Arraste arquivos aqui (PDF, DWG, XLSX, Imagens)", 
        accept_multiple_files=True
    )
    if arquivos_anexos:
        st.success(f"{len(arquivos_anexos)} arquivos anexados. Eles serão baixados junto com o contrato em um ZIP.")

with tab2:
    st.markdown("### Detalhamento Técnico")
    itens_tecnicos = st.multiselect(
        "Selecione os serviços (Inicia vazio):", 
        options=st.session_state['db']['tecnico'],
        default=[] 
    )
    
    c_add, c_free = st.columns(2)
    with c_add:
        novo_tec = st.text_input("➕ Novo item (Técnico)")
        if st.button("Salvar Item Técnico"):
            if aprender_item('tecnico', novo_tec):
                st.success("Salvo!"); st.rerun()
    with c_free:
        tecnico_livre = st.text_area("Texto Livre (Técnico)")

    st.markdown("---")
    st.markdown("### Qualidade")
    itens_qualidade = st.multiselect(
        "Itens de Qualidade (Inicia vazio):",
        options=st.session_state['db']['qualidade'],
        default=[]
    )
    qualidade_livre = st.text_input("Texto Livre (Qualidade)")

with tab3:
    st.markdown("### Matriz de Responsabilidades")
    if not fornecedor:
        st.warning("Preencha o Fornecedor na Aba 1.")
        escolhas_matriz = {}
    else:
        escolhas_matriz = {}
        itens_matriz = [
            ("Materiais de dutos (Chapas, cantos)", "SIARCON"),
            ("Elementos de Difusão", "SIARCON"),
            ("Consumíveis (Parafusos, brocas)", "CONTRATADA"),
            ("Vedação (PU, Fita)", "SIARCON"),
            ("Plataformas elevatórias e andaimes", "SIARCON"),
            ("Escadas manuais tipo A", "CONTRATADA"),
            ("Ferramental manual", "CONTRATADA"),
            ("Alimentação/Transporte Local", "CONTRATADA"),
            ("Recolhimentos trabalhistas", "CONTRATADA"),
            ("Transporte Viagem/Hospedagem", "CONTRATADA"),
            ("Uniformes e EPIs", "CONTRATADA")
        ]
        
        opcoes = ["SIARCON", fornecedor.upper()]
        for item, padrao in itens_matriz:
            c_m1, c_m2 = st.columns([3, 2])
            c_m1.write(f"**{item}**")
            idx = 1 if padrao == "CONTRATADA" else 0
            escolhas_matriz[item] = c_m2.radio(f"r_{item}", opcoes, index=idx, horizontal=True, label_visibility="collapsed")
            st.divider()

with tab4:
    st.markdown("### Segurança")
    nrs = st.multiselect(
        "Selecione documentos adicionais (Inicia vazio):",
        options=st.session_state['db']['docs_sms'],
        default=[]
    )
    
    novo_doc = st.text_input("➕ Novo Documento/NR")
    if st.button("Salvar Documento"):
        if aprender_item('docs_sms', novo_doc): st.success("Salvo!"); st.rerun()

    st.markdown("### Cronograma")
    d_ini = st.date_input("Início", date.today())
    d_int = st.number_input("Dias Integração", 5)
    usa_fim = st.checkbox("Definir fim?")
    d_fim = st.date_input("Fim", date.today()+timedelta(days=30)) if usa_fim else None

with tab5:
    valor = st.text_input("Valor (R$)", "R$ 155.000,00")
    pgto = st.text_area("Condição Pagamento")
    info = st.text_input("Info Extra")
    obs = st.text_area("Observações")

st.markdown("---")
if st.button("📦 GERAR PACOTE DE CONTRATAÇÃO", type="primary", use_container_width=True):
    if not fornecedor:
        st.error("Preencha o nome do Fornecedor!")
    else:
        nomes_anexos = [f.name for f in arquivos_anexos] if arquivos_anexos else []

        dados = {
            'cliente': cliente, 'obra': obra, 'fornecedor': fornecedor,
            'responsavel': responsavel, 'revisao': revisao, 'projetos_ref': projetos_ref,
            'resumo_escopo': resumo_escopo,
            'itens_tecnicos': itens_tecnicos, 'tecnico_livre': tecnico_livre,
            'itens_qualidade': itens_qualidade, 'qualidade_livre': qualidade_livre,
            'matriz': escolhas_matriz, 'nrs_selecionadas': nrs,
            'data_inicio': d_ini, 'dias_integracao': d_int, 'data_fim': d_fim,
            'obs_gerais': obs, 'valor_total': valor,
            'condicao_pgto': pgto, 'info_comercial': info,
            'nomes_anexos': nomes_anexos
        }
        
        docx_buffer = gerar_docx(dados)
        nome_docx = f"Escopo_{fornecedor.replace(' ', '_')}.docx"
        
        if arquivos_anexos:
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                zf.writestr(nome_docx, docx_buffer.getvalue())
                for arquivo in arquivos_anexos:
                    zf.writestr(arquivo.name, arquivo.getvalue())
            zip_buffer.seek(0)
            st.success("✅ Pacote ZIP gerado!")
            st.download_button("📥 BAIXAR PACOTE (.ZIP)", zip_buffer, f"Pacote_{fornecedor}.zip", "application/zip")
        else:
            st.success("✅ Documento Word gerado!")
            st.download_button("📥 BAIXAR ESCOPO (.DOCX)", docx_buffer, nome_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
