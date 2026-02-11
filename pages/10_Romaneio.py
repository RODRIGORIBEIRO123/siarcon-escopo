import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import os
import time
from datetime import datetime, date
import utils_db

# --- 🔒 SEGURANÇA ---
if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado. Faça login.")
    st.stop()

st.set_page_config(page_title="Romaneio de Materiais", page_icon="📦", layout="wide")

# Carrega opções do banco
if 'opcoes_db' not in st.session_state: 
    st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

# ============================================================================
# FUNÇÕES DE DOCUMENTO (DOCX)
# ============================================================================
def set_cell_background(cell, color_hex):
    """Pinta o fundo da célula"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def gerar_romaneio_docx(dados, df_materiais):
    doc = Document()
    
    # --- CONFIGURAÇÃO PAISAGEM (LANDSCAPE) ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Tamanho A4 Paisagem
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    
    # Margens (0.5 polegada = 1.27cm)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # --- CABEÇALHO (Logo + Título apenas) ---
    # Agora apenas 2 colunas para dar mais espaço ao título
    table_head = doc.add_table(rows=1, cols=2)
    table_head.style = 'Table Grid'
    
    # Célula 1: Logo
    c0 = table_head.cell(0, 0)
    c0.width = Inches(2.5)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_path = "logo_siarcon.png"
    if os.path.exists(logo_path):
        try:
            run = p0.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
        except:
            p0.add_run("SIARCON").bold = True
    else:
        p0.add_run("SIARCON").bold = True
    
    # Célula 2: Título (Agora ocupa o resto e sem data)
    c1 = table_head.cell(0, 1)
    p1 = c1.paragraphs[0]
    run_title = p1.add_run("ROMANEIO DE ENVIO DE MATERIAIS")
    run_title.bold = True
    run_title.font.size = Pt(16) # Fonte maior
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("") 

    # --- DADOS EMPRESA/OBRA/DESTINO/DATA (Fundo Cinza) ---
    # Tabela com 3 linhas e 2 colunas
    # Linha 1 e 2: Mescladas
    # Linha 3: Destino (Esq) | Data (Dir)
    t_sub = doc.add_table(rows=3, cols=2)
    t_sub.style = 'Table Grid'
    t_sub.autofit = False # Desliga autofit para respeitar as larguras
    
    # Define larguras da tabela de dados (Total ~10.7)
    for row in t_sub.rows:
        row.cells[0].width = Inches(8.7) # Coluna larga para textos
        row.cells[1].width = Inches(2.0) # Coluna para data
    
    # Linha 1: Empresa (Mesclada)
    c_emp = t_sub.cell(0, 0)
    c_emp.merge(t_sub.cell(0, 1))
    c_emp.text = "Empresa: SIARCON ENGENHARIA LTDA - EPP"
    set_cell_background(c_emp, "D9D9D9")
    c_emp.paragraphs[0].runs[0].bold = True
    
    # Linha 2: Obra (Mesclada)
    c_obra = t_sub.cell(1, 0)
    c_obra.merge(t_sub.cell(1, 1))
    c_obra.text = f"Obra: {dados['obra'].upper()}"
    set_cell_background(c_obra, "D9D9D9")
    c_obra.paragraphs[0].runs[0].bold = True

    # Linha 3: Destino (Esq) e Data (Dir)
    c_dest = t_sub.cell(2, 0)
    c_dest.text = f"Destino: {dados['destino'].upper()}"
    set_cell_background(c_dest, "D9D9D9")
    c_dest.paragraphs[0].runs[0].bold = True
    
    c_data = t_sub.cell(2, 1)
    c_data.text = f"DATA: {dados['data']}"
    set_cell_background(c_data, "D9D9D9")
    c_data.paragraphs[0].runs[0].bold = True
    c_data.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("") 

    # --- TABELA DE MATERIAIS ---
    t_mat = doc.add_table(rows=1, cols=4)
    t_mat.style = 'Table Grid'
    t_mat.autofit = False 
    
    # --- AJUSTE FINO DE LARGURAS (Soma = 10.6 inches) ---
    # Descrição (4.0) | Detalhe (4.4) | Unid (1.1) | Qtd (1.1)
    # Isso garante que caiba na margem de 10.7 disponível
    w_desc = Inches(4.0)
    w_det = Inches(4.4)
    w_unid = Inches(1.1)
    w_qtd = Inches(1.1)
    
    # Cabeçalho
    hdr = t_mat.rows[0].cells
    hdr[0].width = w_desc; hdr[0].text = "DESCRIÇÃO DO MATERIAL"
    hdr[1].width = w_det;  hdr[1].text = "DETALHE"
    hdr[2].width = w_unid; hdr[2].text = "UNID."
    hdr[3].width = w_qtd;  hdr[3].text = "QTD."
    
    for cell in hdr:
        set_cell_background(cell, "D9D9D9")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Dados
    for index, row in df_materiais.iterrows():
        desc = str(row.get('MATERIAL', '')).strip()
        if desc:
            cells = t_mat.add_row().cells
            
            # Aplica larguras em cada célula nova
            cells[0].width = w_desc
            cells[1].width = w_det
            cells[2].width = w_unid
            cells[3].width = w_qtd
            
            cells[0].text = desc
            cells[1].text = str(row.get('DETALHE', ''))
            
            cells[2].text = str(row.get('UNID', ''))
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            qtd = row.get('QTD', 0)
            try: 
                if float(qtd).is_integer(): qtd = int(qtd)
            except: pass
            cells[3].text = str(qtd)
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linhas em branco
    for _ in range(max(0, 12 - len(df_materiais))): 
        row = t_mat.add_row()
        row.cells[0].width = w_desc
        row.cells[1].width = w_det
        row.cells[2].width = w_unid
        row.cells[3].width = w_qtd

    doc.add_paragraph(""); doc.add_paragraph("")

    # --- RODAPÉ ---
    t_footer = doc.add_table(rows=2, cols=2)
    t_footer.autofit = True
    # Força largura total para alinhar assinaturas nas pontas
    t_footer.rows[0].cells[0].width = Inches(5.3)
    t_footer.rows[0].cells[1].width = Inches(5.3)
    
    c_vistos = t_footer.cell(0, 0)
    c_vistos.merge(t_footer.cell(0, 1))
    c_vistos.text = "Vistos:"
    c_vistos.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_sig = t_footer.rows[1]
    p_s = r_sig.cells[0].paragraphs[0]
    p_s.add_run("_______________________________\n").bold = True
    p_s.add_run(f"{dados['responsavel_envio']}\n")
    p_s.add_run("Siarcon Engenharia Ltda - EPP")
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_c = r_sig.cells[1].paragraphs[0]
    p_c.add_run("_______________________________\n").bold = True
    p_c.add_run(f"Responsável: {dados['responsavel_recebimento']}\n")
    p_c.add_run("Recebimento no Destino")
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# ============================================================================
# INTERFACE
# ============================================================================
st.title("📦 Romaneio de Materiais")

# 1. DADOS
c1, c2, c3 = st.columns([2, 2, 1])
obra = c1.text_input("Nome da Obra (Cliente)", value="EUROFARMA")
destino = c2.text_input("Destino (Local de Entrega)", placeholder="Ex: Almoxarifado Central / Prédio Anexo")
data_atual = date.today().strftime("%d/%m/%Y")
data_doc = c3.text_input("Data", value=data_atual)

st.write("### 📋 Lista de Materiais")

opcoes_materiais = sorted(st.session_state['opcoes_db'].get('materiais_romaneio', []))

# --- CORREÇÃO DO DATAFRAME ---
cols_esperadas = ["MATERIAL", "DETALHE", "UNID", "QTD"]
if 'df_romaneio' not in st.session_state:
    st.session_state['df_romaneio'] = pd.DataFrame([{"MATERIAL": "", "DETALHE": "", "UNID": "PC", "QTD": 0} for _ in range(5)])
else:
    df_atual = st.session_state['df_romaneio']
    for col in cols_esperadas:
        if col not in df_atual.columns: df_atual[col] = "" if col != "QTD" else 0
    st.session_state['df_romaneio'] = df_atual[cols_esperadas]

# Configuração da Tabela
config_colunas = {
    "MATERIAL": st.column_config.SelectboxColumn("Descrição (Banco)", options=opcoes_materiais, width="large", required=True),
    "DETALHE": st.column_config.TextColumn("Detalhe / Complemento", width="large"),
    "UNID": st.column_config.SelectboxColumn("Unid.", options=["PC", "KG", "M", "CX", "PAR", "BR", "RL", "CJ", "UN"], width="small", required=True),
    "QTD": st.column_config.NumberColumn("Qtd.", min_value=0, step=1, format="%d", width="small", required=True)
}

df_editado = st.data_editor(
    st.session_state['df_romaneio'],
    column_config=config_colunas,
    num_rows="dynamic",
    use_container_width=True,
    hide_index=True
)

# 4. CADASTRO EM LOTE
st.markdown("---")
with st.expander("➕ Cadastro de Materiais em Lote (Copiar e Colar)"):
    st.info("Cole a lista de materiais novos (um por linha) para adicionar ao banco de opções.")
    texto_lista = st.text_area("Lista:", height=100, placeholder="Parafuso Sextavado\nArruela Lisa\nPorca 3/8")
    if st.button("💾 Cadastrar Lista"):
        if texto_lista:
            novos = [x.strip() for x in texto_lista.split('\n') if x.strip()]
            if novos:
                for i in novos: utils_db.aprender_novo_item("materiais_romaneio", i)
                st.session_state['opcoes_db'] = utils_db.carregar_opcoes()
                st.success(f"{len(novos)} materiais cadastrados!"); time.sleep(1); st.rerun()

st.markdown("---")

# 5. RODAPÉ E AÇÕES
col_env, col_rec = st.columns(2)
resp_envio = col_env.text_input("Responsável Envio (Siarcon)", value="João Bigoni")
resp_receb = col_rec.text_input("Responsável Recebimento", value="Global")

col_b1, col_b2 = st.columns(2)

dados_romaneio = {
    'data': data_doc, 'obra': obra, 'destino': destino,
    'responsavel_envio': resp_envio, 'responsavel_recebimento': resp_receb,
    'materiais_json': df_editado.to_json(orient="records")
}

if col_b1.button("💾 Salvar Romaneio"):
    if utils_db.registrar_romaneio(dados_romaneio): st.success("Salvo!")
    else: st.error("Erro ao salvar.")

if col_b2.button("📄 Gerar DOCX (Paisagem)", type="primary"):
    df_final = df_editado[df_editado["MATERIAL"].astype(str).str.strip() != ""]
    if not df_final.empty:
        arquivo_docx = gerar_romaneio_docx(dados_romaneio, df_final)
        nome_arq = f"Romaneio - {obra} - {datetime.now().strftime('%d-%m-%Y')}.docx"
        st.download_button("📥 Baixar DOCX", arquivo_docx, nome_arq, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else: st.warning("Preencha a lista de materiais.")
