import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import ast
from datetime import date, datetime
import utils_db

# ==============================================================================
# CONTROLE DE ACESSO
# ==============================================================================
if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado."); st.stop()

DISCIPLINA_ATUAL = "GOM"

# ==============================================================================
# ESTRUTURA INTEGRAL - GUIA ORIENTATIVO DE MANUTENÇÃO SIARCON (15 SEÇÕES)
# ==============================================================================
ESTRUTURA_PMOC_SIARCON = {
    "Ventiladores": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de danos e limpar conjunto.", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e eliminar focos de corrosão.", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar fixação, vibrações e ruídos anormais.", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar aquecimento anormal dos mancais.", "frequencia": "S", "tipo": "P"},
                {"item": "Lubrificar os mancais, se aplicável", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar vazamentos nas junções flexíveis", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o estado dos amortecedores de vibração.", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a operação dos controles de vazão.", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o estado e a instalação dos dispositivos de proteção.", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar sistema de drenagem.", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar conjunto mecânico observar itens da seção 3.13", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ]
        }
    },
    "Aquecedor de ar (liquido ou gás)": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de agentes que possam prejudicar a troca térmica", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar as superfícies do lado de ar", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar os fluxos de ar/liquido, vapor ou gás", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar as temperaturas e pressões, na condição de plena vazão de ambos os fluidos e nos pontos de entrada e saída", "frequencia": "S", "tipo": "NP"},
                {"item": "Verificar isolamento térmico do componente (inspeção visual)", "frequencia": "T", "tipo": "P"}
            ]
        }
    },
    "Aquecedor de ar elétrico (resistências)": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de agentes que possam prejudicar a troca térmica", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar as resistências elétricas do lado de ar", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar funcionamento dos dispositivos de segurança", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar os valores de tensão, corrente e isolação elétrica", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a existência de aterramento do componente", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar isolamento térmico do componente (inspeção visual)", "frequencia": "T", "tipo": "P"}
            ]
        }
    },
    "Resfriadores de ar (liquido)": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de agentes que possam prejudicar a troca térmica.", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar as superfícies do lado de ar", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar os fluxos de ar/liquido", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e eliminar a existência de ar do lado de liquido", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar as temperaturas e pressões, na condição de plena vazão de ambos os fluidos e nos pontos de entrada e saída", "frequencia": "S", "tipo": "P"},
                {"item": "Limpar o sistema de drenagem", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a existência de sujeira, danos, corrosão e fixação do eliminador de gotas", "frequencia": "M", "tipo": "P"}
            ]
        }
    },
    "Condicionadores de ar “Expansão direta”": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de agentes que possam prejudicar a troca térmica.", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e eliminar sujeira, danos e corrosão no gabinete, na moldura da serpentina e na bandeja;", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar as serpentinas e bandejas", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a operação dos controles de vazão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar sistema de drenagem", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar isolamento térmico e acústico do componente (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a vedação dos painéis de fechamento do gabinete", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar os filtros (observar itens da seção de filtros)", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Medir e registrar os valores de tensão, corrente e isolação elétrica", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar ventiladores (observar seção de ventiladores)", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Verificar conjunto mecânico observar itens da seção mecânica", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ]
        }
    },
    "Condicionadores de ar tipo “Expansão indireta” (Chillers)": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de agentes que possam prejudicar a troca térmica.", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a eventual perda de refrigerante e a presença de umidade", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e corrigir a carga refrigerante", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar suportes, amortecedores de vibração e fixação", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar as serpentinas e bandejas", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar presença de vazamentos em juntas e válvulas de refrigerantes", "frequencia": "T", "tipo": "P"},
                {"item": "Reparar e corrigir os vazamentos eventuais", "frequencia": "T", "tipo": "NP"},
                {"item": "Verificar a operação dos controles de vazão", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar sistema de drenagem", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar isolamento térmico e acústico do componente (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar funcionamento dos dispositivos de segurança", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a vedação dos painéis de fechamento do gabinete", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar os filtros (observar itens da seção de filtros)", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Medir e registrar os valores de tensão, corrente e isolação elétrica", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar a operação da chave de fluxo de água gelada e bomba", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar e medir a queda de pressão no filtro de óleo", "frequencia": "T", "tipo": "P"},
                {"item": "Inspecionar a vedação da bomba de água", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar e reapertar as conexões elétricas", "frequencia": "A", "tipo": "P"},
                {"item": "Inspecione todos os contatores e relés substituindo os necessários", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar e precisão dos termistores e transdutores por meio de instrumentos calibrados", "frequencia": "A", "tipo": "P"},
                {"item": "Certificar-se de que exista a concentração adequada de anticongelante no circuito de água gelada", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar se o circuito de água possui tratamento adequado", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar filtros em acordo com a seção de filtros", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar a condição e fixação das pás do ventilador no eixo do motor", "frequencia": "A", "tipo": "P"},
                {"item": "Executar o teste de serviço do devido equipamento para confirmar a operação de todos os componentes", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar se existe uma aproximação excessiva da temperatura de saída de água gelada, caso haja realizar limpeza completa da carcaça pois pode indicar incrustação", "frequencia": "A", "tipo": "P"},
                {"item": "Realizar análise do óleo, realizar troca se necessário", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar ventiladores (observar seção de ventiladores)", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Verificar conjunto mecânico observar itens da seção mecânica", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ]
        }
    },
    "Evaporadores (fluído frigorífico ou liquido)": {
        "subitens": {
            "Geral": [
                {"item": "Verificar a existência de agentes que possam prejudicar a troca térmica.", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar as superfícies do lado de ar ou de liquido refrigerado", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar os fluxos de fluidos frigoríficos e refrigerados", "frequencia": "S", "tipo": "P"},
                {"item": "Eliminar a existência de ar do lado do líquido refrigerado", "frequencia": "A", "tipo": "NP"},
                {"item": "Medir e registrar as temperaturas e pressões, na condição de plena vazão de ambos os fluidos e nos pontos de entrada e saída", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar isolamento térmico do componente (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar funcionamento do sistema anticongelamento (fluido frigorífico refrigerado a ar)", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Em casos de soluções aquosas, verificar a concentração do anticongelante", "frequencia": "T", "tipo": "P"},
                {"item": "Corrigir a concentração de anticongelantena", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar sistema de drenagem", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a existência de vazamentos de fluídos refrigerantes, liquido ou ar", "frequencia": "T", "tipo": "P"},
                {"item": "Para evaporador fluído frigorífico, efetuar análise de água quanto a característica", "frequencia": "S", "tipo": "P"},
                {"item": "Efetuar correção da característica da água", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar os filtros (observar seção de filtros)", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Medir e registrar os valores de tensão, corrente e isolação elétrica", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar ventiladores (observar seção de ventiladores)", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Verificar conjunto mecânico observar itens da seção mecânica", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ]
        }
    },
    "Trocador de calor de contracorrente ou corrente-cruzada": {
        "subitens": {
            "Geral": [
                {"item": "Verificar o funcionamento do sistema de purga de ar (no caso de liquido/liquido)", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar as temperaturas e pressões, na condição de plena vazão de ambos os fluidos e nos pontos de entrada e saída", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar isolamento térmico do componente (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a operação dos sistemas de segurança", "frequencia": "M", "tipo": "P"}
            ]
        }
    },
    "Filtros de ar": {
        "subitens": {
            "Filtro rotativo automático": [
                {"item": "Verificar a existência de danos, limpar e vedar frestas na moldura", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Medir e registrar o diferencial de pressão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a operação da alimentação do elemento filtrante", "frequencia": "M", "tipo": "P"},
                {"item": "Completar o fluído de medição do manômetro diferencial", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar estado do material filtrante no alimentador", "frequencia": "M", "tipo": "P"},
                {"item": "Substituir o elemento filtrante", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar conjunto mecânico observar itens da seção mecânica", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ],
            "Filtro seco": [
                {"item": "Verificar a existência de danos, limpar e vedar frestas na moldura", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Medir e registrar o diferencial de pressão", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o ajuste da moldura do filtro na estrutura", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar o elemento filtrante (se recuperável)", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar estado do material filtrante no alimentador", "frequencia": "T", "tipo": "P"},
                {"item": "Substituir o elemento filtrante", "frequencia": "Quando necessário", "tipo": "NP"}
            ],
            "Filtro eletrostático": [
                {"item": "Verificar a existência de danos, sujeira e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar danos e focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar e limpar o módulo eletrostático", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar e vedar frestas da estrutura", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar a tensão elétrica nos módulos eletrostáticos", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a existência de danos no ionizador", "frequencia": "T", "tipo": "P"},
                {"item": "Substituir ionizador", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar e corrigir a ocorrência de descargas elétricas", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar estado e fixação dos isoladores", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar a tensão e corrente elétrica", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a tensão elétrica no modulo eletrostático", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o estado dos dispositivos de proteção elétrica", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar sistema de drenagem", "frequencia": "T", "tipo": "P"}
            ],
            "Filtros absorvente e adsorventes": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar e vedar frestas na estrutura", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar o ajuste do elemento filtrante", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar saturação do elemento filtrante", "frequencia": "T", "tipo": "P"},
                {"item": "Substituir elemento filtrante", "frequencia": "Quando necessário", "tipo": "NP"}
            ],
            "Filtros de alta eficiência para partículas (HEPA) ou superiores": [
                {"item": "NOTA: As atividades de manutenção programada deste tipo de filtro em especifico não podem ser generalizadas, devendo ser analisado caso a caso, em função das particularidades de condição de instalação e operação.", "frequencia": "-", "tipo": "-"}
            ],
            "Filtros embebidos em óleo": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar e vedar frestas na estrutura", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar o diferencial de pressão", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o ajuste da moldura do filtro na estrutura", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar filtro", "frequencia": "M", "tipo": "P"},
                {"item": "Aplicar óleo no elemento filtrante", "frequencia": "M", "tipo": "P"},
                {"item": "Substituir elemento filtrante", "frequencia": "Quando necessário", "tipo": "NP"}
            ]
        }
    },
    "Umidificadores de ar e eliminadores de gotas": {
        "subitens": {
            "Umidificadores com lavador de ar incorporado": [
                {"item": "Verificar a existência de sujeira, sedimentos, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar foco de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o funcionamento do sistema de alimentação e distribuição de água", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o nível de água", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar o funcionamento do extravasor e do sistema de drenagem de água", "frequencia": "T", "tipo": "P"},
                {"item": "Desobstruir o extravasor e o sistema de drenagem", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o funcionamento dos bicos pulverizadores de água", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a impermeabilização e a estanqueidade do sistema", "frequencia": "S", "tipo": "P"}
            ],
            "Umidificadores de ar com gerador de vapor elétrico incorporado": [
                {"item": "Verificar a existência de sujeiras, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o funcionamento do sistema de alimentação e nível da água", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o sistema extravasor e o sistema de drenagem", "frequencia": "T", "tipo": "P"},
                {"item": "Desobstruir o extravasor e o sistema de drenagem", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o funcionamento dos bicos pulverizadores e do sistema de distribuição de vapor", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o funcionamento das válvulas solenoides", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar vazamentos e danos nas linhas de vapor e condensado", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar a tensão e corrente elétrica na entrada", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar operação dos sistemas de segurança", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar o isolamento dos elementos elétricos", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a existência de aterramento dos elementos elétricos", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar filtro de água de acordo com seção de filtros", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ],
            "Umidificadores de ar com vapor de rede externa": [
                {"item": "Verificar o funcionamento das linhas de distribuição do vapor e condensado", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar vazamentos e danos nas linhas de vapor e condensado", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o filtro de vapor", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar o filtro de vapor", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o funcionamento da válvula de controle", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a existência de danos na isolação térmica das linhas de vapor (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Medir e registrar a pressão do vapor antes e depois da válvula de controle", "frequencia": "A", "tipo": "P"}
            ],
            "Geradores de vapor": [
                {"item": "Verificar a existência de sujeira, sedimentos, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar os focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o funcionamento de todas as válvulas", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar a pressão e temperatura do vapor", "frequencia": "A", "tipo": "P"},
                {"item": "Verificar o funcionamento do sistema de aquecimento e seus elementos", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o funcionamento dos sistemas de alimentação de água e controle do nível da água", "frequencia": "M", "tipo": "P"},
                {"item": "NOTA: Verificar o funcionamento dos dispositivos de medição, controle e segurança, de acordo com a NR-13 do Ministério do Trabalho e ASME Boiler and pressure Vessel Code – sections IV,V,VII.", "frequencia": "-", "tipo": "-"}
            ],
            "Eliminadores de gotas": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Corrigir fixação", "frequencia": "Quando necessário", "tipo": "NP"}
            ]
        }
    },
    "Componentes de distribuição e difusão de ar": {
        "subitens": {
            "Venezianas, grelhas e difusores": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Ajuste para reestabelecimento das condições de referência", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar funcionamento mecânico", "frequencia": "S", "tipo": "P"},
                {"item": "Lubrificar mancais de acionamento", "frequencia": "S", "tipo": "P"}
            ],
            "Damper corta-fogo (Registro de ar)": [
                {"item": "Verificar a existência de sujeira nos elementos de fechamento, trava e abertura", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos de fechamento, trava e reabertura", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar funcionamento mecânico", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o posicionamento do indicador de posição", "frequencia": "S", "tipo": "P"},
                {"item": "Lubrificar mancais de acionamento", "frequencia": "S", "tipo": "P"}
            ],
            "Damper (Registro de ar)": [
                {"item": "Verificar a existência de sujeira, danos e corrosão.", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar funcionamento mecânico", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar atuadores do registro", "frequencia": "S", "tipo": "P"},
                {"item": "Lubrificar mancais de acionamento", "frequencia": "S", "tipo": "P"}
            ],
            "Dutos e câmara plenum para ar": [
                {"item": "Verificar a existência de sujeira, danos e corrosão interna e externa, mediante portas de inspeção", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar o conjunto", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar sistema de drenagem", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a vedação das portas de inspeção", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a existência de danos na isolação térmica (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a vedação das conexões", "frequencia": "T", "tipo": "P"}
            ],
            "Unidades de indução": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar o funcionamento dos injetores de indução", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar o conjunto", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Ajustar os injetores de indução", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a existência de danos na isolação térmica (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a estanqueidade das conexões", "frequencia": "S", "tipo": "P"},
                {"item": "Limpar a câmara plenum", "frequencia": "M", "tipo": "P"},
                {"item": "Substituir filtros", "frequencia": "T", "tipo": "P"}
            ],
            "Dispositivos para expansão e mistura": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o funcionamento dos controles de vazão", "frequencia": "M", "tipo": "P"}
            ]
        }
    },
    "Sistemas e quadros elétricos": {
        "subitens": {
            "Sistemas elétricos e eletrônicos": [
                {"item": "Verificar instalações e suas condições locais", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar os elementos quanto ao funcionamento eletromecânico e fixação", "frequencia": "T", "tipo": "P"},
                {"item": "Reapertar terminais, barramentos e elementos de fixação", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar tensão e corrente elétrica dos equipamentos ligados ao quadro", "frequencia": "S", "tipo": "P"},
                {"item": "Regular os elementos de proteção, operação e controle conforme as condições de referência", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o funcionamento de alarmes visuais e sonoros", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o funcionamento nas condições de operação", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar barramento, fiação e sistema de aterramento", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar as tensões de entrada no quadro elétrico", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar aquecimento excessivo em conexões elétricas", "frequencia": "T", "tipo": "P"}
            ],
            "Sistemas de comando pneumático": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o sistema de geração e/ou alimentação de ar comprido", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e regular dispositivos de controle e segurança", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar o sistema de drenagem", "frequencia": "T", "tipo": "P"},
                {"item": "Drenar o reservatório de ar comprimido", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar os elementos filtrantes", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar ou substituir os elementos filtrantes", "frequencia": "S", "tipo": "P"}
            ]
        }
    },
    "Elementos de transmissão e acionamento mecânico": {
        "subitens": {
            "Motores elétricos": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o sentido de rotação", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar vibração e ruídos anormais (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Lubrificar os mancais e/ou rolamentos", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar instalação e fixação das proteções mecânicas", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar a tensão e a corrente elétrica", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar o isolamento elétrico", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar condição da fiação elétrica", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar aterramento", "frequencia": "S", "tipo": "P"}
            ],
            "Polias e correias": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a tensão de esticamento e alinhamento", "frequencia": "T", "tipo": "P"},
                {"item": "Substituir jogo de correias", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Ajustar conjunto", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a fixação e aderência das polias", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a instalação e fixação das proteções mecânicas", "frequencia": "T", "tipo": "P"}
            ],
            "Acoplamentos": [
                {"item": "Verificar a existência de sujeira, danos e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar alinhamento", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar ruídos e vibrações anormais", "frequencia": "M", "tipo": "P"},
                {"item": "Substituir o lubrificante", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a instalação e fixação das proteções mecânicas", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar os elementos de interligação", "frequencia": "S", "tipo": "P"},
                {"item": "Substituir os elementos de interligação", "frequencia": "Quando necessário", "tipo": "NP"}
            ],
            "Correias e engrenagens": [
                {"item": "Verificar a existência de sujeira, danos e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar alinhamento", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar ruídos e vibrações anormais", "frequencia": "M", "tipo": "P"},
                {"item": "Substituir o lubrificante", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a instalação e fixação das proteções mecânicas", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar condição dos eixos e engrenagens", "frequencia": "T", "tipo": "P"},
                {"item": "Realizar troca do óleo e limpeza interna", "frequencia": "S", "tipo": "P"}
            ],
            "Redutores": [
                {"item": "Verificar a existência de sujeira, danos e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os elementos", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a vibração e ruído anormal", "frequencia": "M", "tipo": "P"},
                {"item": "Substituir o óleo e limpar internamente", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a existência de vazamento de óleo lubrificante", "frequencia": "T", "tipo": "P"}
            ]
        }
    },
    "Sistemas hidráulicos": {
        "subitens": {
            "Bombas": [
                {"item": "Verificar a existência de sujeira, danos e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar o equipamento", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar os focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a vibração e ruído anormal", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a vedação do selo mecânico", "frequencia": "T", "tipo": "P"},
                {"item": "Ajustar a prensa-gaxeta", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar nível de óleo", "frequencia": "T", "tipo": "P"},
                {"item": "Completar o nível de óleo", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar o sistema de drenagem", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar itens mecânicos seção mecânica", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ],
            "Válvulas de controle e bloqueio": [
                {"item": "Verificar a existência de sujeira, danos e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar o componente", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar os focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a vibração e ruído anormal", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a existência de vazamentos (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Ajustar elementos de vedação", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar nível de óleo", "frequencia": "T", "tipo": "P"},
                {"item": "Lubrificar o mecanismo de acionamento", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar fiação e conexões dos atuadores e seu funcionamento correto", "frequencia": "A", "tipo": "P"}
            ],
            "Filtros": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar atuação das válvulas (se houver)", "frequencia": "S", "tipo": "P"},
                {"item": "Limpar o componente", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar os focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar danos no elemento filtrante", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar o diferencial de pressão na entrada e na saída do filtro", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar nível de saturação do filtro por meio do diferencial de pressão e condição", "frequencia": "T", "tipo": "P"},
                {"item": "Executar limpeza do filtro quando reutilizável", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Executar troca do elemento filtrante", "frequencia": "Quando necessário", "tipo": "NP"}
            ],
            "Tubulações, tanques e acessórios": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar o componente", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar os focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a existência de vazamentos", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar os tanques internamente", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar isolamento (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar juntas de expansão (inspeção visual)", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar o nível de liquido no tanque de expansão", "frequencia": "T", "tipo": "P"},
                {"item": "Ajustar nível de liquido no tanque de expansão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar os dispositivos de segurança e controle", "frequencia": "S", "tipo": "P"},
                {"item": "Realizar purga do ar", "frequencia": "M", "tipo": "P"},
                {"item": "Drenar para eliminação de sujeira", "frequencia": "T", "tipo": "P"}
            ],
            "Compressores": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar o componente", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar os focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar vibrações e ruídos anormais de fixação", "frequencia": "T", "tipo": "P"},
                {"item": "Medir e registrar a pressão de sucção junto ao compressor", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a temperatura de sucção de gás", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a pressão de descarga", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a temperatura de descarga", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a temperatura da linha de liquido após o condensador", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a temperatura da linha de liquido antes do dispositivo de expansão", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o nível de óleo no visor", "frequencia": "T", "tipo": "P"},
                {"item": "Completar nível de óleo", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Medir e registrar a pressão do óleo", "frequencia": "S", "tipo": "P"},
                {"item": "Ajustar a pressão do óleo nas unidades centrífugas", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Medir e registrar a temperatura do óleo antes e depois do resfriador", "frequencia": "S", "tipo": "P"},
                {"item": "Medir e registrar a temperatura do fluido refrigerante antes e depois do resfriador", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o funcionamento do separador de óleo", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar o funcionamento do sistema de aquecimento", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar operação durante a partida, do dispositivo de redução de capacidade", "frequencia": "T", "tipo": "P"},
                {"item": "Verificar a hermeticidade do selo de vedação do eixo", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o funcionamento das válvulas de serviço", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar existência de vazamentos", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar o funcionamento dos sistemas de segurança", "frequencia": "T", "tipo": "P"}
            ]
        }
    },
    "Circuitos de fluido frigorifico": {
        "subitens": {
            "Tubulações": [
                {"item": "Verificar a existência de danos, corrosão e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar a existência de danos no isolamento", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a existência de vazamentos com detector eletrônico ou outro processo externo", "frequencia": "A", "tipo": "P"},
                {"item": "Reapertar conexões", "frequencia": "S", "tipo": "P"}
            ],
            "Válvulas": [
                {"item": "Verificar a existência de danos, corrosão e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar externamente", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Ajustar parâmetros de operação", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar a existência de vazamentos com detector eletrônico ou outro processo externo", "frequencia": "A", "tipo": "P"}
            ],
            "Acessórios": [
                {"item": "Verificar a existência de danos, corrosão e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar externamente", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar operação", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar a existência de vazamentos com detector eletrônico ou outro processo externo", "frequencia": "A", "tipo": "P"}
            ],
            "Torre de resfriamento": [
                {"item": "Verificar a existência de danos, corrosão e fixação", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar externamente", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar e revisar componentes internos", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar ventiladores em acordo com seção de ventiladores", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Verificar alimentação e distribuição de água", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar o nível da água na bacia", "frequencia": "S", "tipo": "P"},
                {"item": "Ajustar o controlador do nível de água", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar o sistema de purga", "frequencia": "T", "tipo": "P"},
                {"item": "Efetuar análise da água", "frequencia": "S", "tipo": "P"},
                {"item": "Corrigir a característica da água", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Ajustar o volume de purga conforme recomendações técnicas pela análise da água", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Limpar o sistema de drenagem", "frequencia": "T", "tipo": "P"},
                {"item": "Limpar o filtro", "frequencia": "Periodicidade indicada na seção", "tipo": "P"},
                {"item": "Verificar o funcionamento do sistema de acionamento dos ventiladores", "frequencia": "S", "tipo": "P"},
                {"item": "Verificar o funcionamento do termostato", "frequencia": "T", "tipo": "P"},
                {"item": "Ajustar a regulagem dos acionadores dos ventiladores", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar bomba em acordo com seção de bombas", "frequencia": "Periodicidade indicada na seção", "tipo": "P"}
            ],
            "Instrumentação": [
                {"item": "Verificar a existência de sujeira, danos e corrosão", "frequencia": "M", "tipo": "P"},
                {"item": "Limpar externamente", "frequencia": "M", "tipo": "P"},
                {"item": "Eliminar focos de corrosão", "frequencia": "Quando necessário", "tipo": "NP"},
                {"item": "Verificar se o instrumento está fornecendo a informação sobre a grandeza que está medindo", "frequencia": "M", "tipo": "P"},
                {"item": "Verificar e registrar a validade do período de calibração do instrumento, por meio de etiqueta, selo ou certificado.", "frequencia": "A", "tipo": "P"},
                {"item": "Registrar e informar quais os instrumentos que necessitam de calibração ou substituição", "frequencia": "Quando necessário", "tipo": "NP"}
            ]
        }
    }
}

# ==============================================================================
# FUNÇÕES DE ESTILIZAÇÃO AVANÇADA (PYTHON-DOCX)
# ==============================================================================
def aplicar_fundo_celula(celula, cor_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{cor_hex}"/>')
    celula._tc.get_or_add_tcPr().append(shading_elm)

def aplicar_bordas_celula(celula, top="single", bottom="single", left="single", right="single", color="CCCCCC", sz="4"):
    tcPr = celula._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def configurar_cabecalho_siarcon(doc, cod_doc="GOM-2026-00-00"):
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        header = section.header
        header.is_linked_to_previous = False
        
        t_head = header.add_table(rows=1, cols=3, width=Inches(6.5))
        t_head.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        c0, c1, c2 = t_head.rows[0].cells
        c0.width = Inches(1.5); c1.width = Inches(3.5); c2.width = Inches(1.5)
        
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists("logo_siarcon.png"):
            run0 = p0.add_run()
            run0.add_picture("logo_siarcon.png", width=Inches(1.3))
        else:
            r0 = p0.add_run("SIARCON\nEngenharia")
            r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = RGBColor(0, 102, 204)
            
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("GUIA ORIENTATIVO DE\nMANUTENÇÃO")
        r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(13)
        r1.font.color.rgb = RGBColor(102, 102, 102)
        
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(cod_doc)
        r2.bold = True; r2.font.name = "Calibri"; r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(85, 85, 255)
        
        for c in (c0, c1, c2):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            aplicar_bordas_celula(c, color="888888", sz="6")

def configurar_rodape_siarcon(doc):
    for section in doc.sections:
        section.bottom_margin = Inches(0.8)
        footer = section.footer
        footer.is_linked_to_previous = False
        
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        r_emp = p_foot.add_run("SIARCON ENGENHARIA\n")
        r_emp.bold = True; r_emp.font.name = "Calibri"; r_emp.font.size = Pt(9)
        r_emp.font.color.rgb = RGBColor(120, 120, 120)
        
        r_end = p_foot.add_run(
            "Rua: Prof. Estevan Lange Adrien,450-Jd. Nossa Senhora do Amparo\n"
            "CEP: 13482-280- Limeira- SP. – Fone:(19) 3701-7300\n"
            "siarcon@siarcon.com.br – "
        )
        r_end.font.name = "Calibri"; r_end.font.size = Pt(8.5)
        r_end.font.color.rgb = RGBColor(120, 120, 120)
        
        r_site = p_foot.add_run("www.siarcon.com.br")
        r_site.underline = True; r_site.font.name = "Calibri"; r_site.font.size = Pt(8.5)
        r_site.font.color.rgb = RGBColor(0, 102, 204)

# ==============================================================================
# FUNÇÃO DE CONVERSÃO BLINDADA
# ==============================================================================
def converter_para_estrutura(valor, tipo_esperado=list):
    if isinstance(valor, tipo_esperado):
        return valor
    if isinstance(valor, str) and valor.strip():
        try:
            res = ast.literal_eval(valor)
            if isinstance(res, tipo_esperado):
                return res
        except:
            pass
    return [] if tipo_esperado == list else {}

# ==============================================================================
# GERADOR DE DOCX — GUIA ORIENTATIVO DE MANUTENÇÃO SIARCON
# ==============================================================================
def gerar_docx_pmoc(dados):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    cod_doc = dados.get('codigo_doc', f'GOM-{datetime.now().year}-00-00')
    configurar_cabecalho_siarcon(doc, cod_doc)
    configurar_rodape_siarcon(doc)

    # --------------------------------------------------------------------------
    # CAPA DO DOCUMENTO
    # --------------------------------------------------------------------------
    doc.add_paragraph()
    h_capa = doc.add_heading('GUIA ORIENTATIVO DE MANUTENÇÃO', 0)
    h_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p_capa = doc.add_paragraph(
        f"Cliente: {dados.get('nome_proprietario', '-')}\n"
        f"Endereço: {dados.get('end_ambiente', '-')}, {dados.get('num_ambiente', '')} - {dados.get('cidade_ambiente', '')}/{dados.get('uf_ambiente', '')}\n"
        f"Contato: {dados.get('tel_proprietario', '-')}\n"
        f"Ano de Referência: {datetime.now().year}"
    )
    p_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --------------------------------------------------------------------------
    # GERAÇÃO DINÂMICA DO SUMÁRIO
    # --------------------------------------------------------------------------
    selecao_itens = dados.get('selecao_subitens', {})
    secoes_sumario = [
        "RESUMO",
        "1. DADOS GERAIS",
        "  1.1. Identificação do Ambiente ou Conjunto de Ambientes",
        "  1.2. Identificação do Proprietário",
        "  1.3. Identificação do Responsável Técnico pela Obra",
        "2. MAPEAMENTO DO SISTEMA HVAC",
        "  2.1 Relação de ambientes climatizados",
        "  2.2 Relação de equipamentos presentes no sistema",
        "3. PLANO DE MANUTENÇÃO E CONTROLE"
    ]

    idx_cat = 1
    for cat_nome, subitens_selecionados in selecao_itens.items():
        if subitens_selecionados:
            num_cat = f"3.{idx_cat}"
            secoes_sumario.append(f"  {num_cat}. {cat_nome}")
            if not (len(subitens_selecionados) == 1 and subitens_selecionados[0] == "Geral"):
                for idx_sub, sub_nome in enumerate(subitens_selecionados, start=1):
                    secoes_sumario.append(f"    {num_cat}.{idx_sub} {sub_nome}")
            idx_cat += 1

    secoes_sumario.append("4. OBSERVAÇÕES COMPLEMENTARES")
    secoes_sumario.append("ANEXO A – PLANILHA DE CONTROLE")

    doc.add_heading('SUMÁRIO', level=1)
    for linha_sum in secoes_sumario:
        doc.add_paragraph(linha_sum)
    doc.add_page_break()

    # --------------------------------------------------------------------------
    # RESUMO
    # --------------------------------------------------------------------------
    doc.add_heading('RESUMO', level=1)
    doc.add_paragraph(
        "Visando o perfeito estado de conservação e funcionamento de todo o sistema, equipamentos e acessórios, conforme as normas vigentes: "
        "NBR 13.971 (Sistema de Refrigeração, Condicionamento de Ar e Ventilação – Manutenção Programada) da ABNT e Portaria 3.523 do Ministério da Saúde "
        "e Resolução – RE N° 9 de 16 de Janeiro de 2003, (“Orientação Técnica de padrões e referências de Qualidade do Ar Interior, em ambientes climatizados "
        "artificialmente, de uso público e coletivo”) – da Diretoria Colegiada da Agência Nacional de Vigilância Sanitária – ANVISA e Lei Federal 13.589.\n\n"
        "Torna-se necessário utilizar Planos de Manutenção Operação e Controle, para gerenciar a rotina de verificação e inspeção dentro dos prazos estipulados, "
        "ou para a identificação e solução de problemas relacionados ao sistema de AVAC – Aquecimento, Ventilação e Ar-Condicionado. Além disso em acordo com as normas: "
        "Utilizar na limpeza dos componentes do sistema de climatização, produtos biodegradáveis e devidamente registrados no Ministério da Saúde para esse fim. "
        "Preservar a captação de ar externo livre de possíveis fontes de poluentes externas que apresentam riscos à saúde humana."
    )

    # --------------------------------------------------------------------------
    # 1. DADOS GERAIS
    # --------------------------------------------------------------------------
    doc.add_heading('1. DADOS GERAIS', level=1)
    doc.add_heading('1.1 Identificação do Ambiente ou Conjunto de Ambientes:', level=2)
    t1_1 = doc.add_table(rows=3, cols=2); t1_1.style = 'Table Grid'
    t1_1.rows[0].cells[0].text = f"Nome: {dados.get('nome_ambiente', '-')}"
    t1_1.rows[1].cells[0].text = f"Endereço completo: {dados.get('end_ambiente', '-')}"
    t1_1.rows[1].cells[1].text = f"N.º: {dados.get('num_ambiente', '-')}"
    t1_1.rows[2].cells[0].text = f"Complemento: {dados.get('comp_ambiente', '-')} | Bairro: {dados.get('bairro_ambiente', '-')}"
    t1_1.rows[2].cells[1].text = f"Cidade: {dados.get('cidade_ambiente', '-')} | UF: {dados.get('uf_ambiente', '-')}"

    doc.add_heading('1.2 Identificação do Proprietário:', level=2)
    t1_2 = doc.add_table(rows=4, cols=2); t1_2.style = 'Table Grid'
    t1_2.rows[0].cells[0].text = f"Nome/Razão Social: {dados.get('nome_proprietario', '-')}"
    t1_2.rows[1].cells[0].text = f"CNPJ N°: {dados.get('cnpj_proprietario', '-')}"
    t1_2.rows[2].cells[0].text = f"Endereço completo: {dados.get('end_proprietario', '-')}"
    t1_2.rows[2].cells[1].text = f"N.º: {dados.get('num_proprietario', '-')}"
    t1_2.rows[3].cells[0].text = f"Telefone: {dados.get('tel_proprietario', '-')}"
    t1_2.rows[3].cells[1].text = f"Email: {dados.get('email_proprietario', '-')}"

    doc.add_heading('1.3 Identificação do Responsável Técnico:', level=2)
    t1_3 = doc.add_table(rows=4, cols=2); t1_3.style = 'Table Grid'
    t1_3.rows[0].cells[0].text = "Nome / Razão Social: SIARCON ENGENHARIA EIRELI - EPP"
    t1_3.rows[1].cells[0].text = "CNPJ N°: 02.541.727/0001-01"
    t1_3.rows[2].cells[0].text = "Endereço completo: Rua Professor Estevan Lange Adrien | N.º: 450"
    t1_3.rows[2].cells[1].text = "Bairro: Jd. Ns. Sra. Do Amparo | Cidade: Limeira | UF: SP"
    t1_3.rows[3].cells[0].text = "Telefone: (19) 3701-7300"
    t1_3.rows[3].cells[1].text = "Email: contato@siarcon.com.br"

    # --------------------------------------------------------------------------
    # 2. MAPEAMENTO DO SISTEMA HVAC
    # --------------------------------------------------------------------------
    doc.add_heading('2. MAPEAMENTO DO SISTEMA HVAC', level=1)
    doc.add_heading('2.1 Relação de ambientes climatizados:', level=2)
    t_amb = doc.add_table(rows=1, cols=6); t_amb.style = 'Table Grid'
    hdr_amb = t_amb.rows[0].cells
    hdr_amb[0].text = "Tipo de Atividade"; hdr_amb[1].text = "Fixos"; hdr_amb[2].text = "Flutuantes"
    hdr_amb[3].text = "Identificação do Ambiente"; hdr_amb[4].text = "Área m²"; hdr_amb[5].text = "Carga Térmica"
    for amb in dados.get('lista_ambientes', []):
        r_amb = t_amb.add_row().cells
        r_amb[0].text = str(amb.get('atividade', '')); r_amb[1].text = str(amb.get('fixos', ''))
        r_amb[2].text = str(amb.get('flutuantes', '')); r_amb[3].text = str(amb.get('identificacao', ''))
        r_amb[4].text = str(amb.get('area', '')); r_amb[5].text = str(amb.get('carga', ''))

    doc.add_heading('2.2 Relação de equipamentos presentes no sistema:', level=2)
    t_eq = doc.add_table(rows=1, cols=4); t_eq.style = 'Table Grid'
    hdr_eq = t_eq.rows[0].cells
    hdr_eq[0].text = "Equipamento"; hdr_eq[1].text = "Localização"; hdr_eq[2].text = "KW"; hdr_eq[3].text = "TAG"
    for eq in dados.get('lista_equipamentos', []):
        r_eq = t_eq.add_row().cells
        r_eq[0].text = str(eq.get('equipamento', '')); r_eq[1].text = str(eq.get('localizacao', ''))
        r_eq[2].text = str(eq.get('kw', '')); r_eq[3].text = str(eq.get('tag', ''))

    # --------------------------------------------------------------------------
    # 3. PLANO DE MANUTENÇÃO (COM ITENS NORMATIVOS E ITENS CUSTOMIZADOS)
    # --------------------------------------------------------------------------
    doc.add_heading('3. PLANO DE MANUTENÇÃO, OPERAÇÃO E CONTROLE', level=1)
    doc.add_paragraph(
        "Nesta seção encontram-se os itens que devem ser verificados periodicamente de cada equipamento, a nível de componente, "
        "conforme indicados em ABNT NBR 13.971 e determinações do Guia Orientativo de Manutenção.\n"
        "Legenda: M = Mensal | T = Trimestral | S = Semestral | A = Anual\n"
        "P = Atividades periódicas | NP = Atividades a serem executadas se necessário"
    )

    rotinas_customizadas = dados.get('rotinas_customizadas', {})
    idx_cat = 1
    for cat_nome, subitens_selecionados in selecao_itens.items():
        if subitens_selecionados:
            num_cat = f"3.{idx_cat}"
            doc.add_heading(f"{num_cat} {cat_nome}", level=2)

            if not (len(subitens_selecionados) == 1 and subitens_selecionados[0] == "Geral"):
                for idx_sub, sub_nome in enumerate(subitens_selecionados, start=1):
                    num_sub = f"{num_cat}.{idx_sub}"
                    doc.add_heading(f"{num_sub} {sub_nome}", level=3)
                    t_sys = doc.add_table(rows=1, cols=3); t_sys.style = 'Table Grid'
                    h_sys = t_sys.rows[0].cells
                    h_sys[0].text = "Descrição da atividade"; h_sys[1].text = "Periodicidade"; h_sys[2].text = "Prevista"
                    
                    rotinas_base = ESTRUTURA_PMOC_SIARCON.get(cat_nome, {}).get("subitens", {}).get(sub_nome, [])
                    rotinas_extras = rotinas_customizadas.get(f"{cat_nome} > {sub_nome}", [])
                    for rt in (rotinas_base + rotinas_extras):
                        r_sys = t_sys.add_row().cells
                        r_sys[0].text = rt['item']; r_sys[1].text = rt['frequencia']; r_sys[2].text = rt['tipo']
            else:
                t_sys = doc.add_table(rows=1, cols=3); t_sys.style = 'Table Grid'
                h_sys = t_sys.rows[0].cells
                h_sys[0].text = "Descrição da atividade"; h_sys[1].text = "Periodicidade"; h_sys[2].text = "Prevista"
                rotinas_base = ESTRUTURA_PMOC_SIARCON.get(cat_nome, {}).get("subitens", {}).get("Geral", [])
                rotinas_extras = rotinas_customizadas.get(f"{cat_nome} > Geral", [])
                for rt in (rotinas_base + rotinas_extras):
                    r_sys = t_sys.add_row().cells
                    r_sys[0].text = rt['item']; r_sys[1].text = rt['frequencia']; r_sys[2].text = rt['tipo']
            idx_cat += 1

    # --------------------------------------------------------------------------
    # 4. OBSERVAÇÕES COMPLEMENTARES
    # --------------------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading('4. OBSERVAÇÕES COMPLEMENTARES', level=1)
    doc.add_paragraph(
        "As práticas de manutenção acima devem ser aplicadas em conjunto com as recomendações de manutenção mecânica da NBR 13.971 - "
        "Sistemas de Refrigeração, Condicionamento de Ar e Ventilação - Manutenção Programada da ABNT, assim como aos edifícios da Administração "
        "Pública Federal o disposto no capítulo Práticas de Manutenção, Anexo 3, itens 2.6.3 e 2.6.4 da Portaria n.º 2296/97, de 23 de julho de 1997, "
        "Práticas de Projeto, Construção e Manutenção dos Edifícios Públicos Federais, do Ministério da Administração Federal e Reforma do Estado – MARE. "
        "O somatório das práticas de manutenção para garantia do ar e manutenção programada visando o bom funcionamento e desempenho térmico dos sistemas, "
        "permitirá o correto controle dos ajustes das variáveis de manutenção e controle dos poluentes dos ambientes.\n\n"
        "Todos os produtos utilizados na limpeza dos componentes dos sistemas de climatização devem ser biodegradáveis e estarem devidamente registrados "
        "no Ministério da Saúde para esse fim.\n\n"
        "Toda verificação deve ser seguida dos procedimentos necessários para o funcionamento correto do sistema de climatização."
    )

    # --------------------------------------------------------------------------
    # ANEXO A – PLANILHA DE ACOMPANHAMENTO – PMOC
    # --------------------------------------------------------------------------
    doc.add_page_break()
    h_anexo = doc.add_heading('ANEXO A - PLANILHA DE ACOMPANHAMENTO', 1)
    h_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lista_equip = dados.get('lista_equipamentos', [])
    equipamentos_anexo = lista_equip if lista_equip else [{"equipamento": "Climatização Geral HVAC", "localizacao": "-", "tag": "CH-01", "kw": 30}]

    for eq in equipamentos_anexo[:2]:
        t_anx = doc.add_table(rows=0, cols=14)
        t_anx.style = 'Table Grid'
        t_anx.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        r_eq1 = t_anx.add_row().cells
        r_eq1[0].text = f"Descrição do equipamento: {eq.get('equipamento', '')}"
        r_eq1[0].paragraphs[0].runs[0].bold = True
        
        r_eq2 = t_anx.add_row().cells
        r_eq2[0].text = f"Setor: {dados.get('nome_ambiente', '-')}   |   Local: {eq.get('localizacao', '-')}"
        r_eq2[0].paragraphs[0].runs[0].bold = True
        
        r_eq3 = t_anx.add_row().cells
        r_eq3[0].text = f"Capacidade em BTU/h: {eq.get('kw', 0)*3412:.0f}   |   Fabricante: SIARCON   |   Nº de série: -"
        r_eq3[0].paragraphs[0].runs[0].bold = True
        
        r_eq4 = t_anx.add_row().cells
        r_eq4[0].text = f"Modelo: -   |   Patrimônio: -   |   Tag: {eq.get('tag', '-')}"
        r_eq4[0].paragraphs[0].runs[0].bold = True

        for row_c in (r_eq1, r_eq2, r_eq3, r_eq4):
            for i in range(1, 14):
                row_c[0].merge(row_c[i])
            row_c[0].paragraphs[0].paragraph_format.space_after = Pt(2)
            aplicar_fundo_celula(row_c[0], "F9F9F9")

        r_hdr = t_anx.add_row().cells
        r_hdr[0].text = "Periodicidade"; r_hdr[1].text = "Descrição das atividades"
        meses = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
        for idx_m, nome_m in enumerate(meses):
            r_hdr[idx_m + 2].text = nome_m
            r_hdr[idx_m + 2].paragraphs[0].runs[0].font.size = Pt(8)
            r_hdr[idx_m + 2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_hdr[0].paragraphs[0].runs[0].bold = True; r_hdr[1].paragraphs[0].runs[0].bold = True
        r_hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in r_hdr:
            aplicar_fundo_celula(cell, "EAEAEA")

        for _ in range(15):
            r_vazia = t_anx.add_row().cells
            r_vazia[0].text = ""; r_vazia[1].text = ""
            for m in range(2, 14):
                r_vazia[m].text = ""

        r_leg = t_anx.add_row().cells
        r_leg[0].text = "M - Mensal     T - Trimestral     S - Semestral     A - Anual     N - Não se Aplica     OK - Conforme"
        for i in range(1, 14):
            r_leg[0].merge(r_leg[i])
        r_leg[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_leg[0].paragraphs[0].runs[0].font.size = Pt(8); r_leg[0].paragraphs[0].runs[0].bold = True
        aplicar_fundo_celula(r_leg[0], "F2F2F2")

        r_ano = t_anx.add_row().cells
        r_ano[0].text = "Ano:"
        r_ano[0].merge(r_ano[1]); r_ano[0].paragraphs[0].runs[0].bold = True
        meses_abrev = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun", "Jul", "Ago", "Set", "Out", "Nov", "Dez"]
        for idx_ab, m_ab in enumerate(meses_abrev):
            r_ano[idx_ab + 2].text = m_ab
            r_ano[idx_ab + 2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_ano[idx_ab + 2].paragraphs[0].runs[0].font.size = Pt(8)
            aplicar_fundo_celula(r_ano[idx_ab + 2], "F9F9F9")

        rows_assinaturas = [
            "Dia do mês em que a\npreventiva foi realizada",
            "Nome do Técnico",
            "Assinatura"
        ]
        for label_ass in rows_assinaturas:
            r_ass = t_anx.add_row().cells
            r_ass[0].text = label_ass
            r_ass[0].merge(r_ass[1])
            r_ass[0].paragraphs[0].runs[0].font.size = Pt(8.5)
            r_ass[0].paragraphs[0].runs[0].bold = True
            r_ass[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# ==============================================================================
# INTERFACE STREAMLIT - GUIA ORIENTATIVO DE MANUTENÇÃO
# ==============================================================================
st.set_page_config(page_title=f"SIARCON | {DISCIPLINA_ATUAL}", page_icon="📑", layout="wide")
st.title("📑 Guia Orientativo de Manutenção (SIARCON)")
st.caption("Gerador Normativo de Plano de Manutenção — SIARCON ENGENHARIA")

id_projeto = st.session_state.get('id_projeto_editar')
dados_edit = utils_db.buscar_projeto_por_id(id_projeto) if id_projeto else {}

tab1, tab2, tab3, tab4 = st.tabs([
    "1. Dados Gerais (1.1 a 1.3)",
    "2. Mapeamento HVAC (2.1 e 2.2)",
    "3. Seleção da Matriz e Rotinas",
    "4. Emissão Oficial DOCX"
])

with tab1:
    st.subheader("1.1 Identificação do Ambiente e Código")
    col_a1, col_a2 = st.columns(2)
    nome_ambiente = col_a1.text_input("Nome do Empreendimento / Ambiente:", value=dados_edit.get('nome_ambiente', ''))
    end_ambiente = col_a1.text_input("Endereço COMPLETO:", value=dados_edit.get('end_ambiente', ''))
    num_ambiente = col_a2.text_input("Número:", value=dados_edit.get('num_ambiente', ''))
    codigo_doc = col_a2.text_input("Código Institucional (GOM):", value=dados_edit.get('codigo_doc', f'GOM-{datetime.now().year}-00-00'))

    st.divider()
    st.subheader("1.2 Identificação do Proprietário")
    col_p1, col_p2 = st.columns(2)
    nome_prop = col_p1.text_input("Nome / Razão Social do Cliente:", value=dados_edit.get('nome_proprietario', ''))
    cnpj_prop = col_p1.text_input("CNPJ Nº:", value=dados_edit.get('cnpj_proprietario', ''))
    tel_prop = col_p2.text_input("Telefone do Contato:", value=dados_edit.get('tel_proprietario', ''))
    email_prop = col_p2.text_input("E-mail de Contato:", value=dados_edit.get('email_proprietario', ''))

with tab2:
    st.subheader("2.1 Relação de Ambientes Climatizados")
    if 'lista_ambientes' not in st.session_state:
        st.session_state['lista_ambientes'] = converter_para_estrutura(dados_edit.get('lista_ambientes', []), list) or [
            {"atividade": "Escritório / Administrativo", "fixos": 0, "flutuantes": 0, "identificacao": "", "area": 0, "carga": ""}
        ]
        
    for i, amb in enumerate(st.session_state['lista_ambientes']):
        c1, c2, c3, c4, c5 = st.columns([3, 1, 1, 3, 2])
        amb['atividade'] = c1.text_input("Atividade", value=amb.get('atividade', ''), key=f"at_{i}")
        amb['fixos'] = c2.number_input("Fixos", value=int(amb.get('fixos', 0)), key=f"fix_{i}")
        amb['flutuantes'] = c3.number_input("Flutuantes", value=int(amb.get('flutuantes', 0)), key=f"flut_{i}")
        amb['identificacao'] = c4.text_input("Identificação do Ambiente", value=amb.get('identificacao', ''), key=f"id_{i}")
        amb['carga'] = c5.text_input("Carga Térmica", value=amb.get('carga', ''), key=f"cg_{i}")

    if st.button("➕ Adicionar Novo Ambiente", key="add_amb_btn"):
        st.session_state['lista_ambientes'].append({"atividade": "", "fixos": 0, "flutuantes": 0, "identificacao": "", "area": 0, "carga": ""})
        st.rerun()

    st.divider()
    st.subheader("2.2 Relação de Equipamentos Presentes no Sistema")
    if 'lista_equipamentos' not in st.session_state:
        st.session_state['lista_equipamentos'] = converter_para_estrutura(dados_edit.get('lista_equipamentos', []), list) or [
            {"equipamento": "", "localizacao": "", "kw": 0, "tag": ""}
        ]
        
    for j, eq in enumerate(st.session_state['lista_equipamentos']):
        c_eq1, c_eq2, c_eq3, c_eq4 = st.columns([3, 3, 1, 2])
        eq['equipamento'] = c_eq1.text_input("Equipamento", value=eq.get('equipamento', ''), key=f"eq_{j}")
        eq['localizacao'] = c_eq2.text_input("Localização", value=eq.get('localizacao', ''), key=f"loc_{j}")
        eq['kw'] = c_eq3.number_input("KW", value=int(eq.get('kw', 0)), key=f"kw_{j}")
        eq['tag'] = c_eq4.text_input("TAG", value=eq.get('tag', ''), key=f"tag_{j}")

    if st.button("➕ Adicionar Novo Equipamento", key="add_eq_btn"):
        st.session_state['lista_equipamentos'].append({"equipamento": "", "localizacao": "", "kw": 0, "tag": ""})
        st.rerun()

with tab3:
    st.subheader("3. Seleção de Categorias e Subtópicos Mapeados")
    opcoes_categoria = list(ESTRUTURA_PMOC_SIARCON.keys())
    categorias_selecionadas = st.multiselect(
        "Selecione as categorias de equipamentos presentes na obra:",
        options=opcoes_categoria,
        default=opcoes_categoria[:3]
    )

    selecao_subitens = {}
    if categorias_selecionadas:
        st.divider()
        for cat in categorias_selecionadas:
            subitens_disp = list(ESTRUTURA_PMOC_SIARCON[cat]["subitens"].keys())
            if len(subitens_disp) == 1 and subitens_disp[0] == "Geral":
                selecao_subitens[cat] = ["Geral"]
            else:
                st.markdown(f"**🔹 Subitens de '{cat}':**")
                selecionados = st.multiselect(
                    f"Escolha os subtópicos para {cat}:",
                    options=subitens_disp,
                    default=subitens_disp,
                    key=f"sub_{cat}"
                )
                selecao_subitens[cat] = selecionados

        st.divider()
        st.subheader("3.1 Criação de Rotinas e Atividades Personalizadas")
        st.write("Adicione atividades ou comandos específicos que serão incorporados à tabela do relatório:")
        
        if 'rotinas_customizadas' not in st.session_state:
            st.session_state['rotinas_customizadas'] = {}

        col_c1, col_c2 = st.columns(2)
        cat_destino = col_c1.selectbox("Categoria de Destino:", categorias_selecionadas)
        sub_destino = col_c2.selectbox("Subtópico de Destino:", selecao_subitens.get(cat_destino, ["Geral"]))
        
        c_atv1, c_atv2, c_atv3 = st.columns([4, 1, 1])
        desc_atividade = c_atv1.text_input("Descrição da Atividade:")
        periodicidade_sel = c_atv2.selectbox("Periodicidade:", ["M", "T", "S", "A", "Quando necessário"])
        prevista_sel = c_atv3.selectbox("Prevista:", ["P", "NP"])

        if st.button("➕ Adicionar Atividade à Tabela"):
            chave_destino = f"{cat_destino} > {sub_destino}"
            if chave_destino not in st.session_state['rotinas_customizadas']:
                st.session_state['rotinas_customizadas'][chave_destino] = []
            st.session_state['rotinas_customizadas'][chave_destino].append({
                "item": desc_atividade,
                "frequencia": periodicidade_sel,
                "tipo": prevista_sel
            })
            st.success(f"Rotina adicionada a '{chave_destino}' com sucesso!")

        st.markdown("#### 📋 Visualização de Atividades Personalizadas Adicionadas:")
        for chave, itens_cust in st.session_state['rotinas_customizadas'].items():
            if itens_cust:
                st.write(f"**Destino:** `{chave}`")
                st.table([
                    {"Descrição da Atividade": i['item'], "Periodicidade": i['frequencia'], "Prevista": i['tipo']}
                    for i in itens_cust
                ])

with tab4:
    st.subheader("4. Emissão do Guia Orientativo de Manutenção SIARCON")
    st.write("O arquivo .docx gerado contará com a nomenclatura 'Guia Orientativo de Manutenção' no cabeçalho em todas as páginas, sumário automático e as rotinas personalizadas.")
    
    dados_pmoc = {
        '_id': dados_edit.get('_id'),
        'disciplina': DISCIPLINA_ATUAL,
        'codigo_doc': codigo_doc,
        'nome_ambiente': nome_ambiente,
        'end_ambiente': end_ambiente,
        'num_ambiente': num_ambiente,
        'cidade_ambiente': 'Limeira',
        'uf_ambiente': 'SP',
        'nome_proprietario': nome_prop,
        'cnpj_proprietario': cnpj_prop,
        'tel_proprietario': tel_prop,
        'email_proprietario': email_prop,
        'lista_ambientes': st.session_state.get('lista_ambientes', []),
        'lista_equipamentos': st.session_state.get('lista_equipamentos', []),
        'selecao_subitens': selecao_subitens,
        'rotinas_customizadas': st.session_state.get('rotinas_customizadas', {}),
        'data_inicio': dados_edit.get('data_inicio', date.today().strftime("%Y-%m-%d"))
    }
    
    col_b1, col_b2 = st.columns(2)
    with col_b1:
        if st.button("☁️ SALVAR REGISTRO NO DB SIARCON"):
            if utils_db.registrar_projeto(dados_pmoc):
                st.success("✅ Guia Orientativo salvo com sucesso no banco de dados!")
    with col_b2:
        if st.button("💾 SALVAR E GERAR DOCX OFICIAL", type="primary"):
            if utils_db.registrar_projeto(dados_pmoc):
                st.success("✅ Documento exportado em 100% de conformidade! Baixe abaixo:")
                st.session_state['btn_docx_pmoc_oficial'] = True
                
        if st.session_state.get('btn_docx_pmoc_oficial', False):
            b_docx = gerar_docx_pmoc(dados_pmoc)
            st.download_button(
                label="📥 BAIXAR GUIA ORIENTATIVO DE MANUTENÇÃO (.DOCX)",
                data=b_docx,
                file_name=f"Guia_Orientativo_SIARCON_{nome_ambiente.replace(' ', '_') or 'Obra'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
