import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time
from datetime import date, datetime
import utils_db

# --- 🔒 SEGURANÇA ---
if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado. Faça login.")
    st.stop()

# ============================================================================
# 1. CONFIGURAÇÕES ESPECÍFICAS DESTA DISCIPLINA
# ============================================================================
DISCIPLINA_ATUAL = "Dutos"
TEXTO_RESUMO_PADRAO = "Este escopo contempla o fornecimento de rede de dutos, conforme detalhamento a seguir."

# Listas Técnicas
ITENS_MATRIZ = [
    "Fabricação de Dutos (Chapa/MPU)", "Montagem de Dutos", "Isolamento Térmico",
    "Suportação e Fixação", "Instalação de Grelhas/Difusores", "Instalação de Dampers",
    "Dutos Flexíveis", "Conexão com Equipamentos", "Testes de Estanqueidade", 
    "Posicionamento dos equipamentos (Ventiladores / Exaustores)",
    "Posicionamento dos equipamentos (Fancoil / UTA)"
]

PADRAO_TECNICO = [
    "Fabricação e Montagem de Dutos em Chapa Galvanizada (TDC)", "Fabricação e Montagem de Dutos em MPU",
    "Fabricação e Montagem de Dutos em chapa preta", "Aplicação de Isolamento Térmico (Lã de Vidro/Lã de Rocha)",
    "Aplicação de Isolamento Térmico (Borracha Elastomérica)", "Instalação de Suportes e Tirantes",
    "Montagem de Rede de Dutos TDC", "Montagem de Rede de Dutos MPU", "Montagem de Rede de Dutos chapa preta",
    "Montagem de Rede de Dutos Circulares", "Montagem de Rede de Dutos Flexíveis", "Instalação de Dampers de Regulagem",
    "Instalação de Dampers Corta-Fogo", "Instalação de Grelhas, Difusores e Venezianas",
    "Vedação de Flanges e Juntas (Silicone/Fita)", "Conexão de Dutos aos Equipamentos (Fancoils/UTA)",
    "Instalação de Portas de Inspeção", "Posicionamento dos equipamentos", "Fabricação de dutos TDC",
    "Fabricação de dutos em chapa preta"
]

PADRAO_QUALIDADE = [
    "Preparação e teste de 100% da rede de dutos", "Preparação e teste por amostragem de rede de dutos",
    "Todos os dutos devem ser higienizados durante a instalação",
    "Todos os dutos devem ter suas bocas fechadas ao final do dia, com filme plástico",
    "Acompanhamento do trabalho de TAB", "Nivelamento e Alinhamento da Rede",
    "Inspeção de Vedação das Juntas", "Verificação de Fixação dos Suportes"
]

# Lista SMS Fixa para o Documento
SMS_PADRAO_DOC = [
    "Ficha de registro", "ASO (Atestado de Saúde Ocupacional)", "Ficha de EPI", "Ordem de Serviço",
    "Certificados de Treinamento", "NR-06 (Equipamento de Proteção Individual)",
    "NR-12 (Segurança em Máquinas e Equipamentos)",
    "Comprovações de recolhimento de INSS, FGTS e folha de pagamento"
]

# Lista SMS para Seleção (sem as obrigatórias)
LISTA_NRS_SELECAO = [
    "NR-01 (Disposições Gerais)", "NR-03 (Embargo e Interdição)", "NR-04 (SESMT)", "NR-05 (CIPA)", 
    "NR-07 (PCMSO)", "NR-08 (Edificações)", "NR-09 (Avaliação e Controle de Exposições)", 
    "NR-10 (Eletricidade)", "NR-11 (Transporte e Movimentação)", "NR-13 (Vasos de Pressão)", 
    "NR-15 (Insalubridade)", "NR-16 (Periculosidade)", "NR-17 (Ergonomia)", "NR-18 (Construção Civil)", 
    "NR-19 (Explosivos)", "NR-20 (Inflamáveis)", "NR-21 (Trabalho a Céu Aberto)", "NR-23 (Incêndios)", 
    "NR-24 (Condições Sanitárias)", "NR-25 (Resíduos)", "NR-26 (Sinalização)", "NR-28 (Fiscalização)", 
    "NR-33 (Espaços Confinados)", "NR-35 (Trabalho em Altura)", "NR-38 (Limpeza Urbana)"
]

# ============================================================================
# 2. FUNÇÕES AUXILIARES
# ============================================================================
st.set_page_config(page_title=f"Escopo {DISCIPLINA_ATUAL}", page_icon="📝", layout="wide")
if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

cat_tecnica_db = f"tecnico_{DISCIPLINA_ATUAL.lower()}"
id_projeto = st.session_state.get('id_projeto_editar')
dados_edit = {}
if id_projeto:
    t = utils_db.buscar_projeto_por_id(id_projeto)
    if t: dados_edit = t

def formatar_moeda(valor):
    try:
        if not valor: return ""
        v = float(str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip())
        return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except: return valor

def gerar_docx(dados):
    doc = Document()
    try: style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    except: pass
    
    # Título
    doc.add_heading(f'ESCOPO - {dados["disciplina"].upper()}', 0)
    
    # --- 1. DADOS (TABELA) ---
    doc.add_heading('1. DADOS DA OBRA', 1)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    # Dados para a tabela
    data_hj = datetime.now().strftime("%d/%m/%Y")
    info_rows = [
        ("CLIENTE", dados['cliente']),
        ("OBRA", dados['obra']),
        ("FORNECEDOR", dados['fornecedor']),
        ("ENGENHARIA", dados['responsavel']),
        ("SUPRIMENTOS", dados['resp_suprimentos']),
        ("PROJETOS REFERÊNCIA", dados.get('projetos_referencia', '-')),
        ("DATA / REVISÃO", f"{data_hj}  |  Rev: {dados.get('revisao','-')}")
    ]
    
    for idx, (label, value) in enumerate(info_rows):
        row = table.rows[idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)

    # --- 2. TÉCNICO ---
    doc.add_heading('2. ESCOPO TÉCNICO', 1)
    doc.add_paragraph("RESUMO:", style='Strong')
    doc.add_paragraph(dados.get('resumo_escopo', ''))
    
    if dados.get('tecnico_livre'):
        doc.add_paragraph("OBSERVAÇÕES GERAIS:", style='Strong')
        doc.add_paragraph(dados['tecnico_livre'])
    
    doc.add_paragraph("DETALHAMENTO:", style='Strong')
    # Itens selecionados com seus complementos
    itens_selecionados = dados.get('itens_tecnicos', [])
    comentarios = dados.get('comentarios_itens', {})
    
    for item in itens_selecionados:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        # Se tiver complemento, adiciona na frente ou embaixo
        comp = comentarios.get(item, "")
        if comp:
            p.add_run(f": {comp}")

    # --- 3. QUALIDADE ---
    doc.add_heading('3. PADRÃO DE QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []):
        doc.add_paragraph(item, style='List Bullet')

    # --- 4. MATRIZ ---
    doc.add_heading('4. MATRIZ DE RESPONSABILIDADES', 1)
    # Tabela com 4 colunas: Item, Siarcon, Fornecedor, Fora
    tm = doc.add_table(rows=1, cols=4)
    tm.style = 'Table Grid'
    hdr = tm.rows[0].cells
    hdr[0].text = "ITEM"; hdr[1].text = "SIARCON"; hdr[2].text = "FORNECEDOR"; hdr[3].text = "FORA DO ESCOPO"
    
    for item_nome, resp in dados.get('matriz', {}).items():
        row = tm.add_row().cells
        row[0].text = item_nome
        row[1].text = "X" if resp == "SIARCON" else ""
        row[2].text = "X" if resp == "FORNECEDOR" else ""
        row[3].text = "X" if resp == "FORA DO ESCOPO" else ""

    # --- 5. SMS ---
    doc.add_heading('5. SEGURANÇA (SMS)', 1)
    # Itens Padrão Obrigatórios
    for item in SMS_PADRAO_DOC:
        doc.add_paragraph(item, style='List Bullet')
    # Itens Extras Selecionados
    for nr in dados.get('nrs_selecionadas', []):
        doc.add_paragraph(nr, style='List Bullet')
    if dados.get('sms_livre'):
        doc.add_paragraph(dados['sms_livre'])

    # --- 6. COMERCIAL ---
    doc.add_heading('6. COMERCIAL', 1)
    val_fmt = formatar_moeda(dados.get('valor_total',''))
    doc.add_paragraph(f"Valor Global: {val_fmt} (valor fixo e irreajustável)")
    doc.add_paragraph(f"Condição de Pagamento: {dados.get('condicao_pgto','')}")
    if dados.get('obs_gerais'):
        doc.add_paragraph(f"Obs: {dados['obs_gerais']}")

    b = io.BytesIO(); doc.save(b); b.seek(0); return b

# ============================================================================
# 3. INTERFACE DO USUÁRIO
# ============================================================================
st.title(f"📝 {DISCIPLINA_ATUAL}")
if dados_edit: st.info(f"Editando: {dados_edit.get('obra')} | Cliente: {dados_edit.get('cliente')}")
opcoes = st.session_state.get('opcoes_db', {})

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Cadastro", "Técnico", "Matriz", "SMS", "Comercial"])

# --- TAB 1: CADASTRO ---
with tab1:
    c1, c2 = st.columns(2)
    cliente = c1.text_input("Cliente", value=dados_edit.get('cliente', ''))
    obra = c1.text_input("Obra", value=dados_edit.get('obra', ''))
    
    # Seleção de Fornecedor
    db_forn = utils_db.listar_fornecedores()
    lista_nomes = [""] + [f['Fornecedor'] for f in db_forn]
    val_forn_db = dados_edit.get('fornecedor', '')
    idx_f = lista_nomes.index(val_forn_db) if val_forn_db in lista_nomes else 0
    sel_forn = c1.selectbox("Fornecedor (Banco de Dados):", lista_nomes, index=idx_f)
    
    # Cadastro Rápido de Fornecedor
    with c1.expander("Cadastrar Novo Fornecedor"):
        novo_forn_nome = st.text_input("Razão Social")
        novo_forn_cnpj = st.text_input("CNPJ")
        if st.button("Salvar Fornecedor"):
            if novo_forn_nome:
                # Função fictícia de salvar (precisa ter no utils_db ou salvar na lista local)
                # Para simplificar, vou assumir que você vai usar o campo de texto livre abaixo se não salvar
                st.toast("Funcionalidade de salvar fornecedor no DB") 
    
    # Campo final de fornecedor (prevalece o texto)
    forn = c1.text_input("Razão Social (Final):", value=sel_forn if sel_forn else val_forn_db)
    cnpj = c1.text_input("CNPJ:", value=dados_edit.get('cnpj_fornecedor', ''))
    
    resp_eng = c2.text_input("Engenharia SIARCON", value=dados_edit.get('responsavel', ''))
    resp_sup = c2.text_input("Suprimentos SIARCON", value=dados_edit.get('resp_suprimentos', ''))
    revisao = c2.text_input("Revisão", value=dados_edit.get('revisao', 'R-00'))
    
    # Projetos de Referência (Upload ou Texto)
    st.divider()
    c2.write("📂 **Projetos de Referência**")
    uploads = c2.file_uploader("Arraste os arquivos aqui para pegar os nomes", accept_multiple_files=True)
    nomes_arquivos = ""
    if uploads:
        nomes_arquivos = "\n".join([f.name for f in uploads])
    
    val_proj_salvo = dados_edit.get('projetos_referencia', '')
    if nomes_arquivos and not val_proj_salvo:
        val_proj_final = nomes_arquivos
    elif nomes_arquivos and val_proj_salvo:
        val_proj_final = val_proj_salvo + "\n" + nomes_arquivos
    else:
        val_proj_final = val_proj_salvo
        
    projetos_ref = c2.text_area("Lista de Projetos:", value=val_proj_final, height=100, help="Nomes dos arquivos de projeto")

# --- TAB 2: TÉCNICO ---
with tab2:
    # Resumo
    val_resumo = dados_edit.get('resumo_escopo', '')
    if not val_resumo: val_resumo = TEXTO_RESUMO_PADRAO
    resumo = st.text_area("Resumo do Escopo:", value=val_resumo, height=80)
    
    st.divider()
    
    # Adicionar Novo Item ao Banco
    c_add1, c_add2 = st.columns([4, 1])
    novo_item_db = c_add1.text_input("Criar novo item no Banco de Dados:", key="new_item_db")
    if c_add2.button("💾 Criar", key="btn_new_item"):
        if utils_db.aprender_novo_item(cat_tecnica_db, novo_item_db):
            st.session_state['opcoes_db'] = utils_db.carregar_opcoes()
            st.success("Criado!"); time.sleep(0.5); st.rerun()

    # Seleção de Itens
    lista_tec_final = sorted(list(set(opcoes.get(cat_tecnica_db, []) + PADRAO_TECNICO)))
    itens_salvos = dados_edit.get('itens_tecnicos', [])
    if isinstance(itens_salvos, str): 
        try: itens_salvos = eval(itens_salvos) 
        except: itens_salvos = []
    elif not isinstance(itens_salvos, list): itens_salvos = []
    
    opcoes_finais = sorted(list(set(lista_tec_final + itens_salvos)))
    itens_tec = st.multiselect("Selecione os Itens do Escopo:", opcoes_finais, default=itens_salvos)
    
    # CAMPO INDIVIDUAL PARA CADA ITEM (COMPLEMENTO)
    comentarios_salvos = dados_edit.get('comentarios_itens', {})
    if isinstance(comentarios_salvos, str):
        try: comentarios_salvos = eval(comentarios_salvos)
        except: comentarios_salvos = {}
        
    comentarios_novos = {}
    if itens_tec:
        st.caption("📝 Detalhe os itens selecionados (Marca, Modelo, Local, etc.):")
        for item in itens_tec:
            val_coment = comentarios_salvos.get(item, "")
            comentarios_novos[item] = st.text_input(f"Detalhe para '{item}':", value=val_coment)
    
    st.divider()
    tec_livre = st.text_area("Observações Gerais / Itens Livres:", value=dados_edit.get('tecnico_livre', ''))
    
    st.divider()
    st.markdown("#### Padrão de Qualidade")
    # Qualidade
    lista_qual_final = sorted(list(set(opcoes.get(f"qualidade_{DISCIPLINA_ATUAL.lower()}", []) + PADRAO_QUALIDADE)))
    itens_salvos_q = dados_edit.get('itens_qualidade', [])
    if isinstance(itens_salvos_q, str): 
        try: itens_salvos_q = eval(itens_salvos_q) 
        except: itens_salvos_q = []
    elif not isinstance(itens_salvos_q, list): itens_salvos_q = []
    
    opcoes_finais_q = sorted(list(set(lista_qual_final + itens_salvos_q)))
    itens_qual = st.multiselect("Itens de Qualidade:", opcoes_finais_q, default=itens_salvos_q)

# --- TAB 3: MATRIZ ---
with tab3:
    escolhas = {}
    matriz_salva = dados_edit.get('matriz', {})
    if isinstance(matriz_salva, str): 
        try: matriz_salva = eval(matriz_salva) 
        except: matriz_salva = {}
    elif not isinstance(matriz_salva, dict): matriz_salva = {}
    
    st.write("Defina a responsabilidade de cada item:")
    for item in ITENS_MATRIZ:
        c_m1, c_m2 = st.columns([2, 3])
        c_m1.write(f"**{item}**")
        
        # Recupera valor salvo ou padrão
        val_saved = matriz_salva.get(item, "SIARCON")
        opts = ["SIARCON", "FORNECEDOR", "FORA DO ESCOPO"]
        try: idx_res = opts.index(val_saved)
        except: idx_res = 0
        
        escolhas[item] = c_m2.radio(f"resp_{item}", opts, index=idx_res, horizontal=True, label_visibility="collapsed")
        st.divider()

# --- TAB 4: SMS ---
with tab4:
    st.info("ℹ️ Os itens padrão (Ficha, ASO, EPI, NR-06, NR-12, INSS...) já serão incluídos automaticamente no documento.")
    
    nrs_salvas = dados_edit.get('nrs_selecionadas', [])
    if isinstance(nrs_salvas, str): 
        try: nrs_salvas = eval(nrs_salvas) 
        except: nrs_salvas = []
    elif not isinstance(nrs_salvas, list): nrs_salvas = []
    
    opcoes_sms = sorted(list(set(LISTA_NRS_SELECAO + nrs_salvas)))
    nrs = st.multiselect("Selecione NRs Adicionais:", opcoes_sms, default=nrs_salvas)
    sms_livre = st.text_area("Outras exigências de SMS:", value=dados_edit.get('sms_livre', ''))

# --- TAB 5: COMERCIAL ---
with tab5:
    val = st.text_input("Valor Global (R$):", value=dados_edit.get('valor_total', ''))
    st.caption("Nota: O texto '(valor fixo e irreajustável)' será adicionado automaticamente no documento.")
    
    pgto = st.text_area("Condição de Pagamento:", value=dados_edit.get('condicao_pgto', ''))
    obs = st.text_area("Observações Comerciais:", value=dados_edit.get('obs_gerais', ''))
    
    # Status Kanban (Corrigido para bater com o Dashboard)
    lista_st = ["Não Iniciado", "Engenharia", "Obras", "Suprimentos", "Finalizado"]
    st_at = dados_edit.get('status', 'Não Iniciado')
    
    # Mapa de compatibilidade para status antigos
    mapa_fix = {"Em Elaboração": "Engenharia", "Em Cotação": "Suprimentos", "Em Análise Obras": "Obras", "Concluído": "Finalizado"}
    st_at = mapa_fix.get(st_at, st_at)
    
    idx_st = lista_st.index(st_at) if st_at in lista_st else 0
    status = st.selectbox("Status do Projeto:", lista_st, index=idx_st)

# ============================================================================
# 4. SALVAMENTO
# ============================================================================
st.markdown("---")
dados = {
    '_id': dados_edit.get('_id'), 
    'disciplina': DISCIPLINA_ATUAL, 
    'cliente': cliente, 'obra': obra, 'fornecedor': forn, 'cnpj_fornecedor': cnpj,
    'responsavel': resp_eng, 'resp_suprimentos': resp_sup, 'revisao': revisao, 
    'projetos_referencia': projetos_ref, # Novo campo
    'resumo_escopo': resumo, 
    'itens_tecnicos': itens_tec, 
    'comentarios_itens': comentarios_novos, # Novo campo (detalhes dos itens)
    'tecnico_livre': tec_livre, 
    'itens_qualidade': itens_qual, 
    'matriz': escolhas, 
    'nrs_selecionadas': nrs, 'sms_livre': sms_livre, 
    'valor_total': val, 'condicao_pgto': pgto, 'obs_gerais': obs, 
    'status': status, 
    'data_inicio': dados_edit.get('data_inicio', date.today().strftime("%Y-%m-%d"))
}

col_b1, col_b2 = st.columns(2)

if col_b1.button("☁️ SALVAR"):
    if utils_db.registrar_projeto(dados):
        st.success("Salvo com sucesso!")
        time.sleep(1)
    else:
        st.error("Erro ao salvar! Verifique a conexão.")

if col_b2.button("💾 SALVAR E DOCX", type="primary"):
    if utils_db.registrar_projeto(dados):
        b = gerar_docx(dados)
        st.download_button(f"📥 Baixar DOCX", b, f"Escopo_{DISCIPLINA_ATUAL}.docx")
    else:
        st.error("Erro ao salvar! Não foi possível gerar o DOCX.")
