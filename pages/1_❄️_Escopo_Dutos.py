import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import zipfile
from datetime import date, timedelta
import urllib.parse
import utils_db

# --- CONFIGURAÇÃO ---
st.set_page_config(page_title="Escopo Dutos | SIARCON", page_icon="❄️", layout="wide")

# --- FUNÇÕES AUXILIARES ---
def adicionar_item_callback(categoria, key_input):
    novo_item = st.session_state.get(key_input, "")
    if novo_item:
        retorno = utils_db.aprender_novo_item(categoria, novo_item)
        if retorno is True:
            st.session_state[key_input] = ""
            if 'opcoes_db' in st.session_state: del st.session_state['opcoes_db']
            st.toast(f"✅ Item '{novo_item}' salvo em {categoria}!", icon="✅")
        elif retorno == "Duplicado":
            st.toast(f"⚠️ O item '{novo_item}' já existe na lista!", icon="⚠️")
        else:
            st.toast("❌ Erro ao conectar com a planilha.", icon="❌")
    else:
        st.toast("⚠️ Digite algo antes de salvar.", icon="✍️")

def callback_atualizar_nomes_anexos():
    arquivos = st.session_state.get("uploader_anexos", [])
    if arquivos:
        nomes = [f.name for f in arquivos]
        texto_final = "; ".join(nomes)
        st.session_state["input_proj_ref"] = texto_final

# --- LÓGICA DE EDIÇÃO ---
dados_edicao = {}
id_linha_edicao = None

if 'modo_edicao' in st.session_state and st.session_state['modo_edicao']:
    st.info("✏️ MODO EDIÇÃO ATIVO: Editando registro.")
    dados_edicao = st.session_state.get('dados_projeto', {})
    id_linha_edicao = dados_edicao.get('_id_linha')
    
    if st.button("❌ Cancelar Edição (Limpar)"):
        st.session_state['modo_edicao'] = False
        st.session_state['dados_projeto'] = {}
        st.rerun()

# --- CARREGAR DADOS ---
if 'opcoes_db' not in st.session_state:
    with st.spinner("Carregando banco de dados..."):
        st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

# --- FUNÇÃO DOCX ---
def gerar_docx(dados):
    document = Document()
    try:
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    except: pass

    # Cabeçalho
    section = document.sections[0]
    header = section.header
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
        
    p_head = header.add_paragraph()
    p_head.text = "Departamento de Operações SIARCON"
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.style.font.bold = True
    p_head.style.font.size = Pt(14)
    p_head.style.font.name = 'Calibri'

    document.add_paragraph("\n")
    title = document.add_heading('Escopo de fornecimento - Rede de dutos', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_rev = document.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {dados['revisao']}")
    p_rev.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_heading('1. OBJETIVO E RESUMO', level=1)
    table = document.add_table(rows=6, cols=2)
    try: table.style = 'Table Grid'
    except: pass
    
    ref_proj = dados['projetos_ref'] 
    
    infos = [
        ("Cliente:", dados['cliente']), 
        ("Local/Obra:", dados['obra']), 
        ("Projetos Ref:", ref_proj), 
        ("Fornecedor:", dados['fornecedor']), 
        ("Resp. Engenharia:", dados['responsavel']),
        ("Resp. Obras:", dados['resp_obras'])
    ]
    
    for i, (k, v) in enumerate(infos):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = v
    
    document.add_paragraph(f"\nResumo: {dados['resumo_escopo']}")

    document.add_heading('2. TÉCNICO', level=1)
    for item in dados['itens_tecnicos']: document.add_paragraph(item, style='List Bullet')
    if dados['tecnico_livre']: document.add_paragraph(dados['tecnico_livre'], style='List Bullet')

    document.add_heading('3. QUALIDADE', level=1)
    for item in dados['itens_qualidade']: document.add_paragraph(item, style='List Bullet')
    if dados['qualidade_livre']: document.add_paragraph(dados['qualidade_livre'], style='List Bullet')

    document.add_heading('4. MATRIZ RESPONSABILIDADES', level=1)
    table_m = document.add_table(rows=1, cols=3)
    try: table_m.style = 'Table Grid'
    except: pass
    h = table_m.rows[0].cells
    h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = dados['fornecedor'].upper()
    
    for item, resp in dados['matriz'].items():
        row = table_m.add_row().cells
        row[0].text = item
        if resp == "SIARCON": 
            row[1].text = "X"; row[1].paragraphs[0].alignment = 1
        else: 
            row[2].text = "X"; row[2].paragraphs[0].alignment = 1

    document.add_heading('5. SMS', level=1)
    docs_padrao = ["Ficha de registro", "ASO (Atestado de Saúde Ocupacional)", "Ficha de EPI", "Ordem de Serviço", "Certificados de Treinamento", "NR-06 (Equipamento de Proteção Individual)"]
    for doc in docs_padrao: document.add_paragraph(doc, style='List Bullet')
    for doc in dados['nrs_selecionadas']: document.add_paragraph(doc, style='List Bullet')

    document.add_heading('6. CRONOGRAMA', level=1)
    document.add_paragraph(f"Início: {dados['data_inicio'].strftime('%d/%m/%Y')}")
    document.add_paragraph(f"Prazo limite para envio de documentação: {dados['dias_integracao']} dias antes da integração.")
    
    if dados.get('data_fim'):
        document.add_paragraph(f"Previsão de Término: {dados['data_fim'].strftime('%d/%m/%Y')}")

    num_secao = 7
    if dados['obs_gerais']: 
        document.add_heading(f'{num_secao}. OBSERVAÇÕES GERAIS', level=1)
        document.add_paragraph(dados['obs_gerais'])
        num_secao += 1 

    if dados['status'] == "Contratação Finalizada":
        document.add_heading(f'{num_secao}. COMERCIAL', level=1)
        document.add_paragraph(f"Total: {dados['valor_total']} | Pagamento: {dados['condicao_pgto']}")
        if dados['info_comercial']: document.add_paragraph(dados['info_comercial'])
    
    # Rodapé
    footer = section.footer
    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
        
    p_foot = footer.add_paragraph()
    p_foot.text = "SIARCON - Controlando condições ambientais com excelência"
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.style.font.size = Pt(9)
    p_foot.style.font.italic = True

    document.add_paragraph("\n\n")
    document.add_paragraph("_"*60)
    document.add_paragraph(f"DE ACORDO: {dados['fornecedor']}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- FRONT END ---
st.title("❄️ Escopo de Dutos")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["1. Cadastro", "2. Técnico", "3. Matriz", "4. SMS", "5. Comercial e Status"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        val_cli = dados_edicao.get('Cliente', '')
        cliente = st.text_input("Cliente", value=val_cli)
        
        val_obra = dados_edicao.get('Obra', '')
        obra = st.text_input("Obra", value=val_obra)
        
        val_forn = dados_edicao.get('Fornecedor', '')
        if val_forn == "PROPONENTE DE DUTOS": val_forn = "" 
        fornecedor_input = st.text_input("Fornecedor", value=val_forn, placeholder="Deixe em branco p/ genérico")
        if not fornecedor_input: fornecedor_final = "PROPONENTE DE DUTOS"
        else: fornecedor_final = fornecedor_input
            
    with c2:
        col_resp1, col_resp2 = st.columns(2)
        with col_resp1:
            val_resp = dados_edicao.get('Responsável', 'Engenharia')
            responsavel = st.text_input("Resp. Engenharia", value=val_resp)
        with col_resp2:
            val_resp_obras = dados_edicao.get('Responsável Obras', '') 
            resp_obras = st.text_input("Resp. Obras", value=val_resp_obras)
        
        revisao = st.text_input("Revisão", "R-00")
        
        if "input_proj_ref" not in st.session_state:
            st.session_state["input_proj_ref"] = dados_edicao.get('projetos_ref', '')

        projetos_ref = st.text_input("Projetos Ref.", key="input_proj_ref", help="Preenchido automaticamente pelos anexos.")
        
        email_suprimentos = st.text_input("📧 E-mail:", value="suprimentos@siarcon.com.br")

    resumo_escopo = st.text_area("Resumo")
    
    arquivos_anexos = st.file_uploader(
        "Anexos (O nome irá para 'Projetos Ref.' | Baixados separadamente)", 
        accept_multiple_files=True,
        key="uploader_anexos",
        on_change=callback_atualizar_nomes_anexos
    )

with tab2:
    st.subheader("Técnico")
    opcoes_tec = st.session_state['opcoes_db'].get('tecnico', [])
    itens_tecnicos = st.multiselect("Selecione:", options=opcoes_tec)
    
    col_add, col_free = st.columns(2)
    with col_add:
        st.text_input("➕ Novo item (Técnico)", key="input_novo_tec")
        st.button("Salvar Item Técnico", on_click=adicionar_item_callback, args=("tecnico", "input_novo_tec"))
                
    with col_free: tecnico_livre = st.text_area("Texto Livre (Técnico)")
    st.divider()
    
    st.subheader("Qualidade")
    opcoes_qual = st.session_state['opcoes_db'].get('qualidade', [])
    itens_qualidade = st.multiselect("Selecione Qualidade:", options=opcoes_qual)
    
    c_q1, c_q2 = st.columns(2)
    with c_q1:
        st.text_input("➕ Novo item (Qualidade)", key="input_novo_qual")
        st.button("Salvar Qualidade", on_click=adicionar_item_callback, args=("qualidade", "input_novo_qual"))

    with c_q2: qualidade_livre = st.text_input("Texto Livre (Qualidade)")

with tab3:
    escolhas_matriz = {}
    itens_matriz = ["Materiais de dutos (Chapa, canto, isolamento)", "Materiais de difusão de ar", "Consumíveis (Discos, brocas, etc)", "Plataformas elevatórias e/ou andaimes", "Ferramentas manuais", "Escadas tipo \"A\"", "Alimentação, viagem, hospedagem", "Epis", "Uniformes"]
    
    nome_na_matriz = fornecedor_final.upper() if fornecedor_final else "PROPONENTE"
    st.info(f"Matriz de responsabilidades para: **{nome_na_matriz}**")
    
    for item in itens_matriz:
        c1, c2 = st.columns([3,2])
        c1.write(f"**{item}**")
        escolhas_matriz[item] = c2.radio(f"r_{item}", ["SIARCON", nome_na_matriz], horizontal=True, label_visibility="collapsed")
        st.divider()

with tab4:
    opcoes_sms = st.session_state['opcoes_db'].get('sms', [])
    nrs = st.multiselect("SMS Adicional:", options=opcoes_sms)
    st.text_input("➕ Novo Doc SMS", key="input_novo_sms")
    st.button("Salvar SMS", on_click=adicionar_item_callback, args=("sms", "input_novo_sms"))
    st.divider()
    
    c_data1, c_data2 = st.columns(2)
    with c_data1: d_ini = st.date_input("Data de Início")
    with c_data2: d_int = st.number_input("Dias Integração", 5)

    st.markdown("---")
    usar_data_fim = st.checkbox("Informar previsão de término no documento?", value=True)
    if usar_data_fim: d_fim = st.date_input("Previsão de Término", date.today()+timedelta(days=30))
    else: d_fim = None

with tab5:
    st.subheader("Informações Comerciais")
    st.caption("ℹ️ Apenas no Contrato Final.")
    val_total = dados_edicao.get('Valor', 'R$ 0,00')
    valor = st.text_input("Total", value=val_total)
    pgto = st.text_area("Pagamento")
    info = st.text_input("Info Extra")
    obs = st.text_area("Obs")
    st.divider()
    
    st.subheader("🚦 Workflow")
    status_atual = dados_edicao.get('Status', 'Em Elaboração (Engenharia)')
    lista_status = ["Em Elaboração (Engenharia)", "Aguardando Obras", "Recebido (Suprimentos)", "Enviado para Cotação", "Em Negociação", "Contratação Finalizada"]
    try: idx_status = lista_status.index(status_atual)
    except: idx_status = 0
    novo_status = st.selectbox("Fase Atual:", lista_status, index=idx_status)
    if novo_status == "Contratação Finalizada": st.warning("⚠️ Informe o nome da empresa vencedora na aba 1.")

st.markdown("---")

# --- BOTÕES FINAIS ---
if status_atual == "Contratação Finalizada" and 'modo_edicao' in st.session_state:
    st.error("🔒 CONTRATAÇÃO FINALIZADA. Edição bloqueada.")
    st.download_button("📥 Baixar Contrato Final", gerar_docx(dados_edicao).getvalue(), f"Escopo_{val_forn}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    label_botao = "💾 ATUALIZAR PROJETO" if id_linha_edicao else "🚀 REGISTRAR"

    if st.button(label_botao, type="primary", use_container_width=True):
        erro = False
        if novo_status == "Contratação Finalizada":
            if fornecedor_final == "PROPONENTE DE DUTOS" or not fornecedor_final.strip():
                st.error("⛔ Preencha o nome do fornecedor na Aba 1!"); erro = True
        
        if not erro:
            val_proj_ref = st.session_state.get("input_proj_ref", "")
            dados = {
                'cliente': cliente, 'obra': obra, 
                'fornecedor': fornecedor_final,
                'responsavel': responsavel, 
                'resp_obras': resp_obras,
                'revisao': revisao, 
                'projetos_ref': val_proj_ref,
                'resumo_escopo': resumo_escopo,
                'itens_tecnicos': itens_tecnicos, 'tecnico_livre': tecnico_livre,
                'itens_qualidade': itens_qualidade, 'qualidade_livre': qualidade_livre,
                'matriz': escolhas_matriz, 'nrs_selecionadas': nrs,
                'data_inicio': d_ini, 'dias_integracao': d_int, 'data_fim': d_fim, 
                'obs_gerais': obs, 'valor_total': valor, 'condicao_pgto': pgto, 'info_comercial': info,
                'status': novo_status,
                'disciplina': 'Dutos', # <--- FORÇA DISCIPLINA
                'nomes_anexos': [f.name for f in arquivos_anexos] if arquivos_anexos else []
            }
            
            docx_buffer = gerar_docx(dados)
            nome_arq = f"Escopo_{fornecedor_final.replace(' ', '_')}.docx"
            
            zip_buffer = None
            if arquivos_anexos:
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    for arquivo in arquivos_anexos:
                        zf.writestr(arquivo.name, arquivo.getvalue())
                zip_buffer.seek(0)
            
            with st.spinner("Salvando..."):
                utils_db.registrar_projeto(dados, id_linha=id_linha_edicao)
                st.success(f"✅ Projeto Atualizado! Fase: {novo_status}")

            st.divider()
            
            st.markdown("### 📥 Downloads e Notificação")
            col_d1, col_d2, col_d3 = st.columns([1.5, 1.5, 2])
            
            with col_d1:
                st.download_button(
                    "📄 1. Baixar Escopo (DOCX)", 
                    data=docx_buffer.getvalue(), 
                    file_name=nome_arq, 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                    use_container_width=True
                )
            
            with col_d2:
                if zip_buffer:
                    st.download_button(
                        "📎 2. Baixar Anexos (ZIP)", 
                        data=zip_buffer, 
                        file_name=f"Anexos_{obra}.zip", 
                        mime="application/zip", 
                        use_container_width=True
                    )
                else:
                    st.info("Sem anexos.")

            with col_d3:
                assunto_cot = f"Atualização: {obra} - {novo_status}"
                corpo_cot = f"Olá,\n\nSegue documento atualizado.\nObra: {obra}\nStatus: {novo_status}"
                link_cot = f"mailto:{email_suprimentos}?subject={urllib.parse.quote(assunto_cot)}&body={urllib.parse.quote(corpo_cot)}"
                
                html_botao = f"""
                <a href="{link_cot}" target="_blank" style="text-decoration:none;">
                    <button style="width:100%; background-color:#FF4B4B; color:white; border:none; padding:10px; border-radius:5px; font-weight:bold; cursor:pointer;">
                        📧 Notificar por E-mail
                    </button>
                </a>
                """
                st.markdown(html_botao, unsafe_allow_html=True)
