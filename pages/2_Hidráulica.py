import streamlit as st
from docx import Document
from docx.shared import Pt
import io
from datetime import date
import utils_db

# --- 🔒 BLOCO DE SEGURANÇA ---
if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado. Faça login no Dashboard.")
    st.stop()

# ============================================================================
# CONFIGURAÇÃO: HIDRÁULICA
# ============================================================================
DISCIPLINA_ATUAL = "Hidraulica" # Sem acento para compatibilidade com banco antigo

ITENS_MATRIZ = [
    "Tubulações e Conexões", "Válvulas e Acessórios", "Bombas e Equipamentos",
    "Suportes e Fixações", "Isolamento Térmico", "Pintura e Identificação",
    "Teste de Estanqueidade", "Limpeza da Rede", "ART de Execução"
]

# --- LISTAS PADRÃO (Para não ficar vazio) ---
PADRAO_TECNICO = [
    "Rede de Água Gelada (CAG)",
    "Rede de Água de Condensação",
    "Rede de Drenagem de Condensado",
    "Instalação de Válvulas de Controle (PICV/Duas Vias)",
    "Instalação de Válvulas de Bloqueio (Borboleta/Esfera)",
    "Instalação de Filtros Y",
    "Instalação de Termômetros e Manômetros",
    "Montagem de Skids de Bombas",
    "Isolamento em Borracha Elastomérica",
    "Isolamento em Poliestireno Expandido (Isopor)",
    "Revestimento em Alumínio Corrugado",
    "Suportação (Leitos/Abraçadeiras)",
    "Tratamento Químico Inicial"
]

PADRAO_QUALIDADE = [
    "Teste de Pressão Hidrostática (24h)",
    "Ensaio de Liquido Penetrante (Soldas)",
    "Verificação de Pintura e Identificação",
    "Check-list de Montagem de Bombas",
    "Laudo de Estanqueidade",
    "Certificado de Materiais (Tubos/Conexões)",
    "Databook de Obra"
]
# ============================================================================

st.set_page_config(page_title="Escopo Hidráulica", page_icon="💧", layout="wide")

# --- CARGA DE DADOS ---
if 'opcoes_db' not in st.session_state or st.sidebar.button("🔄 Forçar Recarga"):
    with st.spinner("Lendo banco..."):
        st.cache_data.clear()
        st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

cat_tecnica_db = f"tecnico_{DISCIPLINA_ATUAL.lower()}"
cat_qualidade_db = f"qualidade_{DISCIPLINA_ATUAL.lower()}"

# --- VERIFICAÇÃO DE EDIÇÃO ---
id_projeto = st.session_state.get('id_projeto_editar')
dados_edit = {}
if id_projeto:
    temp_dados = utils_db.buscar_projeto_por_id(id_projeto)
    if temp_dados and temp_dados.get('disciplina') == DISCIPLINA_ATUAL:
        dados_edit = temp_dados

def formatar_moeda(valor):
    try:
        v = float(str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip())
        return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except: return valor

def gerar_docx(dados):
    doc = Document()
    try: style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    except: pass
    
    doc.add_heading(f'Escopo - {dados["disciplina"]}', 0)
    doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {dados.get('revisao','-')}")
    
    doc.add_heading('1. DADOS GERAIS', 1)
    t = doc.add_table(rows=1, cols=2)
    try: t.style = 'Table Grid'
    except: pass
    infos = [("Cliente", dados['cliente']), ("Obra", dados['obra']), ("Fornecedor", dados['fornecedor']),
             ("Engenharia", dados['responsavel']), ("Suprimentos", dados['resp_suprimentos'])]
    for k, v in infos:
        row = t.add_row().cells; row[0].text = k; row[0].paragraphs[0].runs[0].bold = True; row[1].text = str(v)

    doc.add_heading('2. ESCOPO TÉCNICO', 1)
    doc.add_paragraph(f"Resumo: {dados.get('resumo_escopo','')}")
    if dados.get('tecnico_livre'): 
        doc.add_paragraph("Obs Técnicas:", style='List Bullet')
        doc.add_paragraph(dados['tecnico_livre'])
    for item in dados.get('itens_tecnicos', []): doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []): doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4. MATRIZ DE RESPONSABILIDADE', 1)
    tm = doc.add_table(rows=1, cols=3)
    try: tm.style = 'Table Grid'
    except: pass
    h = tm.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = "FORNECEDOR"
    for i, r in dados.get('matriz', {}).items():
        row = tm.add_row().cells; row[0].text = i
        if r == "SIARCON": row[1].text = "X"
        else: row[2].text = "X"

    doc.add_heading('5. SMS', 1)
    if dados.get('sms_livre'): doc.add_paragraph(dados['sms_livre'])
    for nr in dados.get('nrs_selecionadas', []): doc.add_paragraph(nr, style='List Bullet')

    doc.add_heading('6. COMERCIAL', 1)
    doc.add_paragraph(f"Valor: {formatar_moeda(dados.get('valor_total',''))}")
    doc.add_paragraph(f"Pagamento: {dados.get('condicao_pgto','')}")
    if dados.get('obs_gerais'): doc.add_paragraph(f"Obs: {dados['obs_gerais']}")

    b = io.BytesIO(); doc.save(b); b.seek(0); return b

# --- INTERFACE ---
st.title(f"💧 {DISCIPLINA_ATUAL}")
if dados_edit: st.info(f"Modo Edição: {dados_edit.get('obra')}")
opcoes = st.session_state.get('opcoes_db', {})

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Cadastro", "Técnico", "Matriz", "SMS", "Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    cliente = c1.text_input("Cliente", value=dados_edit.get('cliente', ''))
    obra = c1.text_input("Obra", value=dados_edit.get('obra', ''))
    
    db_forn = utils_db.listar_fornecedores()
    lista_nomes = [""] + [f['Fornecedor'] for f in db_forn]
    val_forn_db = dados_edit.get('fornecedor', '')
    index_forn = lista_nomes.index(val_forn_db) if val_forn_db in lista_nomes else 0
    sel_forn = c1.selectbox("Fornecedor (Banco):", lista_nomes, index=index_forn)
    
    val_final_forn = sel_forn if sel_forn else dados_edit.get('fornecedor', '')
    forn = c1.text_input("Razão Social:", value=val_final_forn)
    cnpj = c1.text_input("CNPJ:", value=dados_edit.get('cnpj_fornecedor', ''))
    
    resp_eng = c2.text_input("Engenharia", value=dados_edit.get('responsavel', ''))
    resp_sup = c2.text_input("Suprimentos", value=dados_edit.get('resp_suprimentos', ''))
    revisao = c2.text_input("Revisão", value=dados_edit.get('revisao', 'R-00'))
    resumo = c2.text_area("Resumo Escopo", value=dados_edit.get('resumo_escopo', ''))

with tab2:
    st.subheader("Itens Técnicos")
    # MISTURA BANCO + PADRÃO
    lista_tec_db = opcoes.get(cat_tecnica_db, [])
    lista_tec_final = sorted(list(set(lista_tec_db + PADRAO_TECNICO)))
    
    itens_salvos = dados_edit.get('itens_tecnicos', [])
    if isinstance(itens_salvos, str): itens_salvos = eval(itens_salvos)
    opcoes_finais = sorted(list(set(lista_tec_final + itens_salvos)))
    
    itens_tec = st.multiselect("Selecione Itens:", opcoes_finais, default=itens_salvos)
    
    novo_tec = st.text_input("Novo Item Técnico (Salvar no DB):")
    if st.button("💾 Adicionar Item"):
        if utils_db.aprender_novo_item(cat_tecnica_db, novo_tec): st.toast("Salvo!"); st.rerun()
            
    tec_livre = st.text_area("Texto Livre (Técnico):", value=dados_edit.get('tecnico_livre', ''))
    
    st.divider()
    st.subheader("Qualidade")
    # MISTURA BANCO + PADRÃO
    lista_qual_db = opcoes.get(cat_qualidade_db, [])
    lista_qual_final = sorted(list(set(lista_qual_db + PADRAO_QUALIDADE)))
    
    itens_salvos_q = dados_edit.get('itens_qualidade', [])
    if isinstance(itens_salvos_q, str): itens_salvos_q = eval(itens_salvos_q)
    opcoes_finais_q = sorted(list(set(lista_qual_final + itens_salvos_q)))

    itens_qual = st.multiselect("Selecione Itens:", opcoes_finais_q, default=itens_salvos_q)

with tab3:
    escolhas = {}
    matriz_salva = dados_edit.get('matriz', {})
    if isinstance(matriz_salva, str): matriz_salva = eval(matriz_salva)
    
    for item in ITENS_MATRIZ:
        c_m1, c_m2 = st.columns([2,1])
        c_m1.write(f"**{item}**")
        val_padrao = 1 if (item in matriz_salva and matriz_salva[item] != "SIARCON") else 0
        escolhas[item] = c_m2.radio(item, ["SIARCON", "FORNECEDOR"], index=val_padrao, horizontal=True, label_visibility="collapsed", key=f"m_{item}")
        st.divider()

with tab4:
    nrs_salvas = dados_edit.get('nrs_selecionadas', [])
    if isinstance(nrs_salvas, str): nrs_salvas = eval(nrs_salvas)
    opcoes_sms = sorted(list(set(opcoes.get('sms', []) + nrs_salvas)))
    nrs = st.multiselect("NRs:", opcoes_sms, default=nrs_salvas)
    sms_livre = st.text_area("Texto Livre (Segurança):", value=dados_edit.get('sms_livre', ''))

with tab5:
    val = st.text_input("Valor Total (R$)", value=dados_edit.get('valor_total', ''))
    pgto = st.text_area("Pagamento", value=dados_edit.get('condicao_pgto', ''))
    obs = st.text_area("Obs Gerais", value=dados_edit.get('obs_gerais', ''))
    
    lista_status = ["Em Elaboração", "Em Análise Obras", "Em Cotação", "Finalizado", "Concluído"]
    st_atual = dados_edit.get('status', 'Em Elaboração')
    idx_st = lista_status.index(st_atual) if st_atual in lista_status else 0
    status = st.selectbox("Status", lista_status, index=idx_st)

st.markdown("---")
dados = {
    '_id': dados_edit.get('_id'), 'disciplina': DISCIPLINA_ATUAL, 
    'cliente': cliente, 'obra': obra, 'fornecedor': forn, 'cnpj_fornecedor': cnpj,
    'responsavel': resp_eng, 'resp_suprimentos': resp_sup, 'revisao': revisao, 'resumo_escopo': resumo,
    'itens_tecnicos': itens_tec, 'tecnico_livre': tec_livre, 'itens_qualidade': itens_qual, 'matriz': escolhas, 
    'nrs_selecionadas': nrs, 'sms_livre': sms_livre, 'valor_total': val, 'condicao_pgto': pgto, 'obs_gerais': obs,
    'status': status, 'data_inicio': dados_edit.get('data_inicio', date.today().strftime("%Y-%m-%d"))
}

c_b1, c_b2 = st.columns(2)
if c_b1.button("☁️ SALVAR"):
    if utils_db.registrar_projeto(dados): st.success("Salvo!"); st.toast("Salvo")
    else: st.error("Erro")

if c_b2.button("💾 SALVAR E DOCX", type="primary"):
    utils_db.registrar_projeto(dados)
    b = gerar_docx(dados)
    st.download_button(f"📥 Baixar DOCX", b, f"Escopo_{DISCIPLINA_ATUAL}.docx")
