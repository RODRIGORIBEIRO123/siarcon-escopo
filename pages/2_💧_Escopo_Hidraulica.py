import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
from datetime import date, timedelta
import utils_db

st.set_page_config(page_title="Escopo Hidr√°ulica | SIARCON", page_icon="üíß", layout="wide")

def adicionar_item_callback(categoria, key_input):
    novo = st.session_state.get(key_input, "")
    if novo:
        if utils_db.aprender_novo_item(categoria, novo) is True:
            st.session_state[key_input] = ""
            if 'opcoes_db' in st.session_state: del st.session_state['opcoes_db']
            st.toast(f"‚úÖ Item salvo!", icon="‚úÖ")

def atualizar_anexos():
    arquivos = st.session_state.get("uploader_anexos", [])
    if arquivos: st.session_state["input_proj_ref"] = "; ".join([f.name for f in arquivos])

def formatar_moeda_brl(valor):
    if not valor: return ""
    try:
        limpo = str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip()
        float_val = float(limpo)
        formatado = f"{float_val:,.2f}"
        main, decimal = formatado.split('.')
        main = main.replace(',', '.')
        return f"R$ {main},{decimal}"
    except: return valor

dados_edicao = {}
id_linha_edicao = None
if 'modo_edicao' in st.session_state and st.session_state['modo_edicao']:
    dados_edicao = st.session_state.get('dados_projeto', {})
    id_linha_edicao = dados_edicao.get('_id_linha')
    if st.button("‚ùå Cancelar Edi√ß√£o"):
        st.session_state['modo_edicao'] = False; st.rerun()

if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

def gerar_docx(dados):
    document = Document()
    try: style = document.styles['Normal']; font = style.font; font.name = 'Calibri'; font.size = Pt(11)
    except: pass
    
    section = document.sections[0]; header = section.header
    for p in header.paragraphs: p._element.getparent().remove(p._element)
    p = header.add_paragraph(); p.text = "Departamento de Opera√ß√µes SIARCON"; p.alignment = 1; p.style.font.bold = True; p.style.font.size = Pt(14)

    document.add_paragraph("\n")
    document.add_heading('Escopo de fornecimento - Rede Hidr√°ulica', 0).alignment = 1
    
    revisao_txt = dados.get('revisao', 'R-00')
    document.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {revisao_txt}").alignment = 2

    document.add_heading('1. OBJETIVO', 1)
    table = document.add_table(rows=7, cols=2)
    try: table.style = 'Table Grid'
    except: pass
    
    infos = [
        ("Cliente:", dados.get('cliente', '')), ("Obra:", dados.get('obra', '')), ("Ref:", dados.get('projetos_ref', '')), 
        ("Fornecedor:", dados.get('fornecedor', '')), ("Resp. Eng:", dados.get('responsavel', '')), 
        ("Resp. Obras:", dados.get('resp_obras', '')), ("Resp. Suprimentos:", dados.get('resp_suprimentos', ''))
    ]
    for i, (k, v) in enumerate(infos): 
        if i < len(table.rows):
            row = table.rows[i]; row.cells[0].text = k; row.cells[0].paragraphs[0].runs[0].bold = True; row.cells[1].text = v
    document.add_paragraph(f"\nResumo: {dados.get('resumo_escopo', '')}")

    document.add_heading('2. T√âCNICO', 1)
    for item in dados.get('itens_tecnicos', []): document.add_paragraph(item, style='List Bullet')
    if dados.get('tecnico_livre'): document.add_paragraph(dados['tecnico_livre'], style='List Bullet')

    document.add_heading('3. QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []): document.add_paragraph(item, style='List Bullet')
    if dados.get('qualidade_livre'): document.add_paragraph(dados['qualidade_livre'], style='List Bullet')

    document.add_heading('4. MATRIZ RESPONSABILIDADES', 1)
    t_m = document.add_table(rows=1, cols=3); 
    try: t_m.style = 'Table Grid'
    except: pass
    h = t_m.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = dados.get('fornecedor', 'FORNECEDOR').upper()
    for item, resp in dados.get('matriz', {}).items():
        row = t_m.add_row().cells; row[0].text = item
        if resp == "SIARCON": row[1].text = "X"; row[1].paragraphs[0].alignment = 1
        else: row[2].text = "X"; row[2].paragraphs[0].alignment = 1

    document.add_heading('5. SMS', 1)
    docs = ["Ficha de registro", "ASO", "Ficha EPI", "Ordem de Servi√ßo", "Certificados Treinamento", "NR-06"]
    for d in docs: document.add_paragraph(d, style='List Bullet')
    for d in dados.get('nrs_selecionadas', []): document.add_paragraph(d, style='List Bullet')

    document.add_heading('6. CRONOGRAMA', 1)
    if dados.get('data_inicio'): document.add_paragraph(f"In√≠cio: {dados['data_inicio'].strftime('%d/%m/%Y')}")
    if dados.get('data_fim'): document.add_paragraph(f"T√©rmino: {dados['data_fim'].strftime('%d/%m/%Y')}")

    num_secao = 7
    if dados.get('obs_gerais'): 
        document.add_heading(f'{num_secao}. OBSERVA√á√ïES', 1); document.add_paragraph(dados['obs_gerais']); num_secao += 1 
    
    if dados.get('status') == "Contrata√ß√£o Finalizada":
        document.add_heading(f'{num_secao}. COMERCIAL', 1)
        p_val = document.add_paragraph(); p_val.add_run("Valor Global: ").bold = True
        p_val.add_run(f"{formatar_moeda_brl(dados.get('valor_total', ''))} (Fixo e irreajust√°vel)")
        p_pgto = document.add_paragraph(); p_pgto.add_run("Condi√ß√£o de Pagamento: ").bold = True
        p_pgto.add_run(dados.get('condicao_pgto', ''))
        if dados.get('info_comercial'): document.add_paragraph(dados['info_comercial'])
    
    section = document.sections[0]; footer = section.footer
    for p in footer.paragraphs: p._element.getparent().remove(p._element)
    pf = footer.add_paragraph(); pf.text = "SIARCON Engenharia, controlando condi√ß√µes ambientais com excel√™ncia."
    pf.alignment = 1; pf.style.font.size = Pt(9); pf.style.font.italic = True
    document.add_paragraph("\n\n"); document.add_paragraph("_"*60); document.add_paragraph(f"DE ACORDO: {dados.get('fornecedor', '')}")
    b = io.BytesIO(); document.save(b); b.seek(0); return b

# --- INTERFACE ---
st.title("üíß Escopo de Hidr√°ulica")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["1. Cadastro", "2. T√©cnico", "3. Matriz", "4. SMS", "5. Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        cliente = st.text_input("Cliente", value=dados_edicao.get('Cliente', ''))
        obra = st.text_input("Obra", value=dados_edicao.get('Obra', ''))
        forn = st.text_input("Fornecedor", value=dados_edicao.get('Fornecedor', ''))
        if not forn: forn = "PROPONENTE HIDR√ÅULICA"
        email_sup = st.text_input("E-mail Suprimentos", value="suprimentos@siarcon.com.br")
    with c2:
        c_r1, c_r2, c_r3 = st.columns(3)
        resp_eng = c_r1.text_input("Resp. Eng.", value=dados_edicao.get('Respons√°vel', ''))
        resp_obr = c_r2.text_input("Resp. Obras", value=dados_edicao.get('Respons√°vel Obras', ''))
        resp_sup = c_r3.text_input("Resp. Suprim.", value=dados_edicao.get('resp_suprimentos', ''))
        revisao = st.text_input("Revis√£o", "R-00")
        if "input_proj_ref" not in st.session_state: st.session_state["input_proj_ref"] = dados_edicao.get('projetos_ref', '')
        proj_ref = st.text_input("Projetos Ref.", key="input_proj_ref")
    resumo = st.text_area("Resumo")
    anexos = st.file_uploader("Anexos", accept_multiple_files=True, key="uploader_anexos", on_change=atualizar_anexos)

with tab2:
    st.subheader("T√©cnico")
    key_tec = "multi_tec_hidraulica"
    if key_tec not in st.session_state: st.session_state[key_tec] = dados_edicao.get('itens_tecnicos', [])
    itens_tec = st.multiselect("Itens T√©cnicos:", sorted(st.session_state['opcoes_db'].get('tecnico_hidraulica', [])), key=key_tec)
    
    c_add, c_free = st.columns(2)
    c_add.text_input("Novo T.", key="n_tec"); c_add.button("Salvar T.", on_click=adicionar_item_callback, args=("tecnico_hidraulica", "n_tec"))
    tec_livre = c_free.text_area("Livre T.")
    
    st.divider()
    st.subheader("Qualidade")
    key_qual = "multi_qual_hidraulica"
    if key_qual not in st.session_state: st.session_state[key_qual] = dados_edicao.get('itens_qualidade', [])
    itens_qual = st.multiselect("Itens Qualidade:", sorted(st.session_state['opcoes_db'].get('qualidade_hidraulica', [])), key=key_qual)
    
    c_q1, c_q2 = st.columns(2)
    c_q1.text_input("Novo Q.", key="n_qual"); c_q1.button("Salvar Q.", on_click=adicionar_item_callback, args=("qualidade_hidraulica", "n_qual"))
    qual_livre = c_q2.text_input("Livre Q.")

with tab3:
    escolhas = {}
    itens_m = ["Tubula√ß√µes/Conex√µes", "Isolamento T√©rmico", "Insumos Solda", "Consum√≠veis", "Cavaletes/Drenos", "V√°lvulas", "Plataformas/Andaimes", "Ferramentas", "Escadas", "Alimenta√ß√£o/Viagem", "EPIs", "Uniformes"]
    nome_m = forn.upper(); st.info(f"Matriz: {nome_m}")
    for i in itens_m:
        c1, c2 = st.columns([3,2])
        c1.write(f"**{i}**")
        escolhas[i] = c2.radio(f"m_{i}", ["SIARCON", nome_m], horizontal=True, label_visibility="collapsed")
        st.divider()

with tab4:
    nrs = st.multiselect("SMS Adicional:", sorted(st.session_state['opcoes_db'].get('sms', [])))
    c_d1, c_d2 = st.columns(2)
    d_ini = c_d1.date_input("In√≠cio"); d_int = c_d2.number_input("Dias Int.", 5)
    uf = st.checkbox("Data Fim?", True); d_fim = st.date_input("Fim", date.today()+timedelta(days=30)) if uf else None

with tab5:
    val = st.text_input("Valor Total (Ex: 35000.00)", dados_edicao.get('Valor', '')); pgto = st.text_area("Pagamento"); info = st.text_input("Info"); obs = st.text_area("Obs")
    st.divider()
    status = st.selectbox("Status", ["Em Elabora√ß√£o (Engenharia)", "Aguardando Obras", "Recebido (Suprimentos)", "Enviado para Cota√ß√£o", "Em Negocia√ß√£o", "Contrata√ß√£o Finalizada"], index=0)

st.markdown("---")
if status == "Contrata√ß√£o Finalizada" and 'modo_edicao' in st.session_state:
    st.error("üîí Finalizado."); st.download_button("üì• Baixar", gerar_docx(dados_edicao).getvalue(), f"Escopo_{forn}.docx")
else:
    if st.button("üíæ SALVAR / ATUALIZAR", type="primary"):
        dados = {
            'cliente': cliente, 'obra': obra, 'fornecedor': forn, 'responsavel': resp_eng, 'resp_obras': resp_obr,
            'resp_suprimentos': resp_sup, 'revisao': revisao, 'projetos_ref': proj_ref, 'resumo_escopo': resumo,
            'itens_tecnicos': itens_tec, 'tecnico_livre': tec_livre,
            'itens_qualidade': itens_qual, 'qualidade_livre': qual_livre,
            'matriz': escolhas, 'nrs_selecionadas': nrs,
            'data_inicio': d_ini, 'dias_integracao': d_int, 'data_fim': d_fim,
            'obs_gerais': obs, 'valor_total': val, 'condicao_pgto': pgto, 'info_comercial': info,
            'status': status, 'disciplina': 'Hidr√°ulica',
            'nomes_anexos': [f.name for f in anexos] if anexos else []
        }
        docx = gerar_docx(dados); utils_db.registrar_projeto(dados, id_linha_edicao)
        st.success("‚úÖ Salvo!"); st.download_button("üì• Baixar DOCX", docx.getvalue(), f"Escopo_{forn.replace(' ', '_')}.docx")
