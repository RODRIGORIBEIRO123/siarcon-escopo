import streamlit as st
from docx import Document
from docx.shared import Pt
import io
from datetime import date
import utils_db

# ============================================================================
# 🚨 ATENÇÃO: MUDE ESTAS DUAS VARIÁVEIS PARA CADA ARQUIVO!
# ============================================================================
DISCIPLINA_ATUAL = "Eletrica"
ITENS_MATRIZ = ["Cabos e Fios", "Eletrocalhas", "Quadros/Painéis", "Disjuntores", "Infraestrutura", "Montagem", "Testes"]
# ============================================================================

st.set_page_config(page_title=f"Escopo {DISCIPLINA_ATUAL}", page_icon="📝", layout="wide")

# --- CARGA DE DADOS ---
if 'opcoes_db' not in st.session_state or st.sidebar.button("🔄 Recarregar Dados"):
    with st.spinner("Sincronizando com Banco de Dados..."):
        st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

def formatar_moeda(valor):
    try:
        v = float(str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip())
        return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except: return valor

# --- GERADOR DE DOCX ---
def gerar_docx(dados):
    doc = Document()
    try: style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    except: pass
    
    doc.add_heading(f'Escopo - {dados["disciplina"]}', 0)
    doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {dados.get('revisao','-')}")
    
    doc.add_heading('1. DADOS GERAIS', 1)
    t = doc.add_table(rows=1, cols=2)
    try: t.style = 'Table Grid'
    except: pass
    
    infos = [("Cliente", dados['cliente']), ("Obra", dados['obra']), ("Fornecedor", dados['fornecedor']),
             ("Engenharia", dados['responsavel']), ("Suprimentos", dados['resp_suprimentos'])]
    for k, v in infos:
        row = t.add_row().cells; row[0].text = k; row[0].paragraphs[0].runs[0].bold = True; row[1].text = str(v)

    doc.add_heading('2. ESCOPO TÉCNICO', 1)
    doc.add_paragraph(f"Resumo: {dados.get('resumo_escopo','')}")
    
    # Campo Livre
    if dados.get('tecnico_livre'): 
        doc.add_paragraph("Observações Técnicas Gerais:", style='List Bullet')
        doc.add_paragraph(dados['tecnico_livre'])
        
    doc.add_paragraph("Itens Específicos:", style='List Bullet')
    for item in dados.get('itens_tecnicos', []): doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []): doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4. MATRIZ DE RESPONSABILIDADE', 1)
    tm = doc.add_table(rows=1, cols=3)
    try: tm.style = 'Table Grid'
    except: pass
    h = tm.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = "FORNECEDOR"
    for i, r in dados.get('matriz', {}).items():
        row = tm.add_row().cells; row[0].text = i
        if r == "SIARCON": row[1].text = "X"
        else: row[2].text = "X"

    doc.add_heading('5. SMS E SEGURANÇA', 1)
    for nr in dados.get('nrs_selecionadas', []): doc.add_paragraph(nr, style='List Bullet')

    doc.add_heading('6. COMERCIAL', 1)
    doc.add_paragraph(f"Valor: {formatar_moeda(dados.get('valor_total',''))}")
    doc.add_paragraph(f"Pagamento: {dados.get('condicao_pgto','')}")
    if dados.get('obs_gerais'): doc.add_paragraph(f"Obs: {dados['obs_gerais']}")

    b = io.BytesIO(); doc.save(b); b.seek(0); return b

# --- INTERFACE ---
st.title(f"🛠️ {DISCIPLINA_ATUAL}")
opcoes = st.session_state.get('opcoes_db', {'tecnico':[], 'qualidade':[], 'sms':[]})

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Cadastro", "Técnico", "Matriz", "SMS", "Comercial"])

with tab1:
    st.warning("⚠️ Atenção: Os campos do fornecedor (abaixo) devem ser preenchidos exclusivamente pelo time de Suprimentos.")
    c1, c2 = st.columns(2)
    cliente = c1.text_input("Cliente")
    obra = c1.text_input("Obra")
    
    db_forn = utils_db.listar_fornecedores()
    sel_forn = c1.selectbox("Fornecedor (Banco):", [""] + [f['Fornecedor'] for f in db_forn])
    cnpj_auto = next((str(f['CNPJ']) for f in db_forn if f['Fornecedor'] == sel_forn), "") if sel_forn else ""
    
    forn = c1.text_input("Razão Social:", value=sel_forn)
    cnpj = c1.text_input("CNPJ:", value=cnpj_auto)
    
    resp_eng = c2.text_input("Engenharia")
    resp_sup = c2.text_input("Suprimentos")
    revisao = c2.text_input("Revisão", "R-00")
    resumo = c2.text_area("Resumo Escopo")

with tab2:
    st.subheader("Itens Técnicos")
    k_tec = f"tec_{DISCIPLINA_ATUAL.lower()}"
    itens_tec = st.multiselect("Selecione Itens:", opcoes.get('tecnico', []), key=k_tec)
    
    c_add, c_txt = st.columns(2)
    novo_tec = c_add.text_input("Novo Item Técnico (DB):", key=f"new_{k_tec}")
    if c_add.button("💾 Adicionar Técnico", key=f"btn_{k_tec}"):
        if utils_db.aprender_novo_item("tecnico", novo_tec):
            st.toast("Item técnico salvo!"); st.rerun()
            
    tec_livre = st.text_area("📝 Informações Gerais e Complementares (Texto Livre):", height=150, placeholder="Descreva detalhes específicos...")
    
    st.divider()
    st.subheader("Controle de Qualidade")
    k_qual = f"qual_{DISCIPLINA_ATUAL.lower()}"
    itens_qual = st.multiselect("Selecione Itens de Qualidade:", opcoes.get('qualidade', []), key=k_qual)
    
    c_add_q, c_vazio = st.columns(2)
    novo_qual = c_add_q.text_input("Novo Item Qualidade (DB):", key=f"new_q_{k_qual}")
    if c_add_q.button("💾 Adicionar Qualidade", key=f"btn_q_{k_qual}"):
        if utils_db.aprender_novo_item("qualidade", novo_qual):
             st.toast("Item qualidade salvo!"); st.rerun()

with tab3:
    escolhas = {}
    nome_f = forn.split(' ')[0].upper() if forn else "FORN"
    for item in ITENS_MATRIZ:
        c_m1, c_m2 = st.columns([2,1])
        c_m1.write(f"**{item}**")
        escolhas[item] = c_m2.radio(item, ["SIARCON", nome_f], horizontal=True, label_visibility="collapsed", key=f"mtz_{DISCIPLINA_ATUAL}_{item}")
        st.divider()

with tab4:
    st.subheader("Normas Regulamentadoras (NRs)")
    nrs = st.multiselect("Selecione as NRs Aplicáveis:", opcoes.get('sms', []), key=f"sms_{DISCIPLINA_ATUAL}")

with tab5:
    c_v1, c_v2 = st.columns(2)
    val = c_v1.text_input("Valor Total (R$)")
    pgto = c_v2.text_area("Condição de Pagamento")
    obs = st.text_area("Observações Gerais")
    status = st.selectbox("Status", ["Em Elaboração", "Finalizado"])

st.markdown("---")

dados_projeto = {
    'disciplina': DISCIPLINA_ATUAL, 'cliente': cliente, 'obra': obra,
    'fornecedor': forn, 'cnpj_fornecedor': cnpj,
    'responsavel': resp_eng, 'resp_suprimentos': resp_sup,
    'revisao': revisao, 'resumo_escopo': resumo,
    'itens_tecnicos': itens_tec, 'tecnico_livre': tec_livre,
    'itens_qualidade': itens_qual, 'matriz': escolhas, 'nrs_selecionadas': nrs,
    'valor_total': val, 'condicao_pgto': pgto, 'obs_gerais': obs,
    'status': status, 'data_inicio': date.today().strftime("%Y-%m-%d")
}

col_b1, col_b2 = st.columns(2)

if col_b1.button("☁️ APENAS SALVAR (Banco de Dados)"):
    if not cliente or not obra: st.error("Preencha Cliente e Obra.")
    else:
        if utils_db.registrar_projeto(dados_projeto):
            st.success("✅ Salvo na nuvem!"); st.toast("Salvo!")
        else: st.error("Erro ao salvar.")

if col_b2.button("💾 SALVAR E GERAR ARQUIVO (Download)", type="primary"):
    if not cliente or not obra: st.error("Preencha Cliente e Obra.")
    else:
        utils_db.registrar_projeto(dados_projeto)
        docx_buffer = gerar_docx(dados_projeto)
        nome_arquivo = f"Escopo_{DISCIPLINA_ATUAL}_{cliente}_{obra}.docx".replace(" ", "_")
        st.success("✅ Dados salvos!")
        st.download_button(f"📥 Baixar: {nome_arquivo}", docx_buffer, nome_arquivo)
