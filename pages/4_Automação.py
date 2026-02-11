import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time
from datetime import date, datetime
import utils_db

if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado."); st.stop()

DISCIPLINA_ATUAL = "Automação"
TEXTO_RESUMO_PADRAO = "Este escopo contempla o fornecimento de sistema de automação, conforme detalhamento a seguir."

# --- LISTAS AUTOMAÇÃO (MANTIDAS) ---
ITENS_MATRIZ = [
    "Controladores (DDC/PLC)", "Sensores e Atuadores", "Infraestrutura de Rede",
    "Cabeamento de Controle", "Painéis de Automação", "Software Supervisório",
    "Comissionamento/Start-up", "Treinamento Operacional", "Licenças de Software"
]

PADRAO_TECNICO = [
    "Fornecimento e Instalação de Controladores (DDC)", "Instalação de Sensores de Temperatura/Umidade",
    "Instalação de Sensores de Pressão Diferencial", "Instalação de Atuadores de Válvulas e Dampers",
    "Lançamento de Cabos de Rede (CAT6 / RS-485)", "Montagem de Painéis de Automação",
    "Interligação Elétrica dos Periféricos", "Programação de Lógica de Controle",
    "Desenvolvimento de Telas Gráficas (Supervisório)", "Integração com Equipamentos (Chiller/Fancoil - BACnet/Modbus)",
    "Start-up e Testes Funcionais"
]

PADRAO_QUALIDADE = [
    "Teste de Ponto a Ponto à frio", "Teste de ponto a Ponto à quente",
    "Emissão de relatório de comissionamento dos pontos", "Emissão de memorial de lógica de controle",
    "Teste de Lógica de Controle", "Teste de Falha de Comunicação", "Verificação de Calibração de Sensores",
    "Backup da Programação Entregue", "Treinamento da Equipe de Operação", "Manual de Operação do Sistema",
    "Lista de Pontos (I/O List) As-Built", "Lista de spare parts"
]

SMS_PADRAO_DOC = [
    "Ficha de registro", "ASO (Atestado de Saúde Ocupacional)", "Ficha de EPI", "Ordem de Serviço",
    "Certificados de Treinamento", "NR-06 (Equipamento de Proteção Individual)",
    "NR-12 (Segurança em Máquinas e Equipamentos)",
    "Comprovações de recolhimento de INSS, FGTS e folha de pagamento"
]

LISTA_NRS_SELECAO = [
    "NR-01 (Disposições Gerais)", "NR-03 (Embargo e Interdição)", "NR-04 (SESMT)", "NR-05 (CIPA)", 
    "NR-07 (PCMSO)", "NR-08 (Edificações)", "NR-09 (Avaliação e Controle de Exposições)", 
    "NR-10 (Eletricidade)", "NR-11 (Transporte e Movimentação)", "NR-13 (Vasos de Pressão)", 
    "NR-15 (Insalubridade)", "NR-16 (Periculosidade)", "NR-17 (Ergonomia)", "NR-18 (Construção Civil)", 
    "NR-19 (Explosivos)", "NR-20 (Inflamáveis)", "NR-21 (Trabalho a Céu Aberto)", "NR-23 (Incêndios)", 
    "NR-24 (Condições Sanitárias)", "NR-25 (Resíduos)", "NR-26 (Sinalização)", "NR-28 (Fiscalização)", 
    "NR-33 (Espaços Confinados)", "NR-35 (Trabalho em Altura)", "NR-38 (Limpeza Urbana)"
]

st.set_page_config(page_title=f"Escopo {DISCIPLINA_ATUAL}", page_icon="🤖", layout="wide")
if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

cat_tecnica_db = f"tecnico_{DISCIPLINA_ATUAL.lower()}"
id_projeto = st.session_state.get('id_projeto_editar')
dados_edit = {}
if id_projeto:
    t = utils_db.buscar_projeto_por_id(id_projeto)
    if t: dados_edit = t

def formatar_moeda(valor):
    try:
        if not valor: return ""
        v = float(str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip())
        return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except: return valor

def gerar_docx(dados):
    doc = Document()
    try: style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    except: pass
    
    # Título Centralizado
    head = doc.add_heading(f'ESCOPO - {dados["disciplina"].upper()}', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo (20pt)
    sub = doc.add_paragraph('SIARCON ENGENHARIA')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(20)
    
    # 1. DADOS
    doc.add_heading('1. DADOS DA OBRA', 1)
    table = doc.add_table(rows=8, cols=2); table.style = 'Table Grid'
    data_hj = datetime.now().strftime("%d/%m/%Y")
    info_rows = [
        ("CLIENTE", dados['cliente']), ("OBRA", dados['obra']), ("FORNECEDOR", dados['fornecedor']),
        ("ENGENHARIA", dados['responsavel']), ("OBRAS", dados.get('resp_obras', '')),
        ("SUPRIMENTOS", dados['resp_suprimentos']),
        ("PROJETOS REFERÊNCIA", dados.get('projetos_referencia', '-')),
        ("DATA / REVISÃO", f"{data_hj}  |  Rev: {dados.get('revisao','-')}")
    ]
    for idx, (label, value) in enumerate(info_rows):
        row = table.rows[idx]
        row.cells[0].text = label; row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)

    # 2. TÉCNICO
    doc.add_heading('2. ESCOPO TÉCNICO', 1)
    doc.add_paragraph(dados.get('resumo_escopo', ''))
    if dados.get('tecnico_livre'):
        p = doc.add_paragraph(); p.add_run("OBSERVAÇÕES GERAIS:").bold = True
        doc.add_paragraph(dados['tecnico_livre'])
    
    p = doc.add_paragraph(); p.add_run("DETALHAMENTO:").bold = True
    comentarios = dados.get('comentarios_itens', {})
    for item in dados.get('itens_tecnicos', []):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item); run.bold = True
        if comentarios.get(item): p.add_run(f": {comentarios[item]}")

    # 3. QUALIDADE
    doc.add_heading('3. PADRÃO DE QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []): doc.add_paragraph(item, style='List Bullet')

    # 4. MATRIZ
    doc.add_heading('4. MATRIZ DE RESPONSABILIDADES', 1)
    tm = doc.add_table(rows=1, cols=4); tm.style = 'Table Grid'
    h = tm.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = "FORNECEDOR"; h[3].text = "FORA DO ESCOPO"
    for k, v in dados.get('matriz', {}).items():
        row = tm.add_row().cells; row[0].text = k
        row[1].text = "X" if v == "SIARCON" else ""; row[2].text = "X" if v == "FORNECEDOR" else ""
        row[3].text = "X" if v == "FORA DO ESCOPO" else ""

    # 5. SMS
    doc.add_heading('5. SEGURANÇA (SMS)', 1)
    for i in SMS_PADRAO_DOC: doc.add_paragraph(i, style='List Bullet')
    for nr in dados.get('nrs_selecionadas', []): doc.add_paragraph(nr, style='List Bullet')
    if dados.get('sms_livre'): doc.add_paragraph(dados['sms_livre'])

    # 6. COMERCIAL
    doc.add_heading('6. COMERCIAL', 1)
    doc.add_paragraph(f"Valor Global: {formatar_moeda(dados.get('valor_total',''))} (valor fixo e irreajustável)")
    doc.add_paragraph(f"Condição de Pagamento: {dados.get('condicao_pgto','')}")
    if dados.get('obs_gerais'): doc.add_paragraph(f"Obs: {dados['obs_gerais']}")

    b = io.BytesIO(); doc.save(b); b.seek(0); return b

# --- INTERFACE ---
st.title(f"🤖 {DISCIPLINA_ATUAL}")
if dados_edit: st.info(f"Editando: {dados_edit.get('obra')} | Cliente: {dados_edit.get('cliente')}")
opcoes = st.session_state.get('opcoes_db', {})

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Cadastro", "Técnico", "Matriz", "SMS", "Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    cliente = c1.text_input("Cliente", value=dados_edit.get('cliente', ''))
    obra = c1.text_input("Obra", value=dados_edit.get('obra', ''))
    
    db_forn = utils_db.listar_fornecedores(); lista_nomes = [""] + [f['Fornecedor'] for f in db_forn]
    val_forn_db = dados_edit.get('fornecedor', ''); idx_f = lista_nomes.index(val_forn_db) if val_forn_db in lista_nomes else 0
    sel_forn = c1.selectbox("Fornecedor (DB):", lista_nomes, index=idx_f)
    with c1.expander("Cadastrar Novo Fornecedor"):
        n_forn = st.text_input("Razão Social"); n_cnpj = st.text_input("CNPJ")
        if st.button("Salvar Fornecedor"): st.toast("Função DB necessária")
    forn = c1.text_input("Razão Social (Final):", value=sel_forn if sel_forn else val_forn_db)
    cnpj = c1.text_input("CNPJ:", value=dados_edit.get('cnpj_fornecedor', ''))
    
    resp_eng = c2.text_input("Engenharia SIARCON", value=dados_edit.get('responsavel', ''))
    resp_obras = c2.text_input("Obras SIARCON", value=dados_edit.get('resp_obras', ''))
    resp_sup = c2.text_input("Suprimentos SIARCON", value=dados_edit.get('resp_suprimentos', ''))
    revisao = c2.text_input("Revisão", value=dados_edit.get('revisao', 'R-00'))
    
    st.divider(); c2.write("📂 **Projetos de Referência**")
    val_proj_salvo = dados_edit.get('projetos_referencia', '')
    if not val_proj_salvo: val_proj_salvo = ""
    
    set_arquivos = set([x.strip() for x in val_proj_salvo.split('\n') if x.strip()])
    
    uploads = c2.file_uploader("Arraste arquivos aqui", accept_multiple_files=True)
    if uploads:
        for f in uploads:
            set_arquivos.add(f.name)
            
    val_proj_final = "\n".join(sorted(list(set_arquivos)))
    projetos_ref = c2.text_area("Lista de Projetos:", value=val_proj_final, height=100)

with tab2:
    val_resumo = dados_edit.get('resumo_escopo', '') or TEXTO_RESUMO_PADRAO
    resumo = st.text_area("Resumo:", value=val_resumo, height=80)
    st.divider()
    
    c_a1, c_a2 = st.columns([4,1])
    novo_item = c_a1.text_input("Novo Item DB:", key="add_db")
    if c_a2.button("💾 Criar"): 
        if utils_db.aprender_novo_item(cat_tecnica_db, novo_item): st.session_state['opcoes_db'] = utils_db.carregar_opcoes(); st.rerun()

    lista_tec_final = sorted(list(set(opcoes.get(cat_tecnica_db, []) + PADRAO_TECNICO)))
    
    itens_salvos = dados_edit.get('itens_tecnicos', [])
    if isinstance(itens_salvos, str) and itens_salvos.strip():
        try: itens_salvos = eval(itens_salvos)
        except: itens_salvos = []
    elif not isinstance(itens_salvos, list): itens_salvos = []
        
    opcoes_finais = sorted(list(set(lista_tec_final + itens_salvos)))
    itens_tec = st.multiselect("Itens do Escopo:", opcoes_finais, default=itens_salvos)
    
    comentarios_salvos = dados_edit.get('comentarios_itens', {})
    if isinstance(comentarios_salvos, str) and comentarios_salvos.strip():
        try: comentarios_salvos = eval(comentarios_salvos)
        except: comentarios_salvos = {}
    elif not isinstance(comentarios_salvos, dict): comentarios_salvos = {}
        
    comentarios_novos = {}
    if itens_tec:
        st.caption("📝 Detalhe os itens (Marca, Local, etc.):")
        for i in itens_tec:
            comentarios_novos[i] = st.text_input(f"Detalhe '{i}':", value=comentarios_salvos.get(i, ""))
            
    st.divider()
    tec_livre = st.text_area("Observações Gerais:", value=dados_edit.get('tecnico_livre', ''))
    st.divider(); st.markdown("#### Qualidade")
    
    lista_qual = sorted(list(set(opcoes.get(f"qualidade_{DISCIPLINA_ATUAL.lower()}", []) + PADRAO_QUALIDADE)))
    
    itens_salvos_q = dados_edit.get('itens_qualidade', [])
    if isinstance(itens_salvos_q, str) and itens_salvos_q.strip():
        try: itens_salvos_q = eval(itens_salvos_q)
        except: itens_salvos_q = []
    elif not isinstance(itens_salvos_q, list): itens_salvos_q = []
        
    opcoes_qual = sorted(list(set(lista_qual + itens_salvos_q)))
    itens_qual = st.multiselect("Itens Qualidade:", opcoes_qual, default=itens_salvos_q)

with tab3:
    escolhas = {}
    matriz_salva = dados_edit.get('matriz', {})
    if isinstance(matriz_salva, str) and matriz_salva.strip():
        try: matriz_salva = eval(matriz_salva)
        except: matriz_salva = {}
    elif not isinstance(matriz_salva, dict): matriz_salva = {}
    
    st.write("Responsabilidades:")
    for item in ITENS_MATRIZ:
        c_m1, c_m2 = st.columns([2, 3])
        c_m1.write(f"**{item}**")
        opts = ["SIARCON", "FORNECEDOR", "FORA DO ESCOPO"]
        idx = opts.index(matriz_salva.get(item, "SIARCON")) if matriz_salva.get(item) in opts else 0
        escolhas[item] = c_m2.radio(f"r_{item}", opts, index=idx, horizontal=True, label_visibility="collapsed")
        st.divider()

with tab4:
    st.info("Itens padrão (Ficha, ASO, EPI, NRs 06/12...) inclusos automaticamente.")
    
    nrs_salvas = dados_edit.get('nrs_selecionadas', [])
    if isinstance(nrs_salvas, str) and nrs_salvas.strip():
        try: nrs_salvas = eval(nrs_salvas)
        except: nrs_salvas = []
    elif not isinstance(nrs_salvas, list): nrs_salvas = []
        
    opcoes_sms = sorted(list(set(LISTA_NRS_SELECAO + nrs_salvas)))
    nrs = st.multiselect("NRs Adicionais:", opcoes_sms, default=nrs_salvas)
    sms_livre = st.text_area("Outras exigências:", value=dados_edit.get('sms_livre', ''))

with tab5:
    val = st.text_input("Valor Global (R$):", value=dados_edit.get('valor_total', ''))
    pgto = st.text_area("Condição Pagamento:", value=dados_edit.get('condicao_pgto', ''))
    obs = st.text_area("Obs Comerciais:", value=dados_edit.get('obs_gerais', ''))
    
    lista_st = ["Não Iniciado", "Engenharia", "Obras", "Suprimentos", "Finalizado"]
    st_at = dados_edit.get('status', 'Não Iniciado')
    mapa_fix = {"Em Elaboração": "Engenharia", "Em Cotação": "Suprimentos", "Em Análise Obras": "Obras", "Concluído": "Finalizado"}
    st_at = mapa_fix.get(st_at, st_at)
    idx_st = lista_st.index(st_at) if st_at in lista_st else 0
    status = st.selectbox("Status:", lista_st, index=idx_st)

st.markdown("---")
dados = {
    '_id': dados_edit.get('_id'), 'disciplina': DISCIPLINA_ATUAL, 'cliente': cliente, 'obra': obra, 
    'fornecedor': forn, 'cnpj_fornecedor': cnpj, 'responsavel': resp_eng, 
    'resp_obras': resp_obras,
    'resp_suprimentos': resp_sup, 
    'revisao': revisao, 'projetos_referencia': projetos_ref, 'resumo_escopo': resumo, 
    'itens_tecnicos': itens_tec, 'comentarios_itens': comentarios_novos, 'tecnico_livre': tec_livre, 
    'itens_qualidade': itens_qual, 'matriz': escolhas, 'nrs_selecionadas': nrs, 'sms_livre': sms_livre, 
    'valor_total': val, 'condicao_pgto': pgto, 'obs_gerais': obs, 'status': status, 
    'data_inicio': dados_edit.get('data_inicio', date.today().strftime("%Y-%m-%d"))
}

col_b1, col_b2 = st.columns(2)
if col_b1.button("☁️ SALVAR"):
    if utils_db.registrar_projeto(dados): st.success("Salvo!"); time.sleep(1)
    else: st.error("Erro ao salvar.")

if col_b2.button("💾 SALVAR E DOCX", type="primary"):
    if utils_db.registrar_projeto(dados):
        f_forn = forn.strip() or "Fornecedor"; f_obra = obra.strip() or "Obra"
        b = gerar_docx(dados); st.download_button(f"Baixar DOCX", b, f"Escopo_{DISCIPLINA_ATUAL} - {f_forn} - {f_obra}.docx")
    else: st.error("Erro ao salvar.")
