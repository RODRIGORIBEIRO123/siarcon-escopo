import streamlit as st
from docx import Document
from docx.shared import Pt
import io
import zipfile
from datetime import date, timedelta
import utils_db

st.set_page_config(page_title="Escopo Automação | SIARCON", page_icon="🤖", layout="wide")

def adicionar_item_callback(categoria, key_input):
    novo = st.session_state.get(key_input, "")
    if novo:
        if utils_db.aprender_novo_item(categoria, novo) is True:
            st.session_state[key_input] = ""
            if 'opcoes_db' in st.session_state: del st.session_state['opcoes_db']
            st.toast("✅ Salvo!", icon="✅")

def atualizar_anexos():
    arquivos = st.session_state.get("uploader_anexos", [])
    if arquivos: st.session_state["input_proj_ref"] = "; ".join([f.name for f in arquivos])

def formatar_moeda_brl(valor):
    if not valor: return ""
    try:
        limpo = str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip()
        float_val = float(limpo)
        formatado = f"{float_val:,.2f}"
        main, decimal = formatado.split('.')
        main = main.replace(',', '.')
        return f"R$ {main},{decimal}"
    except: return valor

dados_edicao = {}
id_linha_edicao = None
if 'modo_edicao' in st.session_state and st.session_state['modo_edicao']:
    dados_edicao = st.session_state.get('dados_projeto', {})
    id_linha_edicao = dados_edicao.get('_id_linha')
    if st.button("❌ Cancelar Edição"):
        st.session_state['modo_edicao'] = False; st.rerun()

if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

def gerar_docx(dados):
    document = Document()
    try: style = document.styles['Normal']; font = style.font; font.name = 'Calibri'; font.size = Pt(11)
    except: pass
    
    p = document.add_paragraph(); p.text = "Departamento de Operações SIARCON"; p.alignment = 1; p.style.font.bold = True; p.style.font.size = Pt(14)
    document.add_paragraph("\n"); document.add_heading('Escopo - Automação e Controle', 0).alignment = 1
    
    revisao_txt = dados.get('revisao', 'R-00')
    document.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {revisao_txt}").alignment = 2

    document.add_heading('1. OBJETIVO', 1)
    table = document.add_table(rows=7, cols=2)
    try: table.style = 'Table Grid'
    except: pass
    infos = [
        ("Cliente:", dados.get('cliente', '')), ("Obra:", dados.get('obra', '')), ("Ref:", dados.get('projetos_ref', '')), 
        ("Fornecedor:", dados.get('fornecedor', '')), ("Resp. Eng:", dados.get('responsavel', '')), 
        ("Resp. Obras:", dados.get('resp_obras', '')), ("Resp. Suprimentos:", dados.get('resp_suprimentos', ''))
    ]
    for i, (k, v) in enumerate(infos): 
        if i < len(table.rows):
            row = table.rows[i]; row.cells[0].text = k; row.cells[0].paragraphs[0].runs[0].bold = True; row.cells[1].text = v
    document.add_paragraph(f"\nResumo: {dados.get('resumo_escopo', '')}")

    document.add_heading('2. TÉCNICO', 1)
    for item in dados.get('itens_tecnicos', []): document.add_paragraph(item, style='List Bullet')
    if dados.get('tecnico_livre'): document.add_paragraph(dados['tecnico_livre'], style='List Bullet')

    document.add_heading('3. QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []): document.add_paragraph(item, style='List Bullet')
    if dados.get('qualidade_livre'): document.add_paragraph(dados['qualidade_livre'], style='List Bullet')

    document.add_heading('4. MATRIZ', 1)
    t_m = document.add_table(rows=1, cols=3); 
    try: t_m.style = 'Table Grid'
    except: pass
    h = t_m.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = dados.get('fornecedor', 'FORNECEDOR').upper()
    for item, resp in dados.get('matriz', {}).items():
        row = t_m.add_row().cells; row[0].text = item
        if resp == "SIARCON": row[1].text = "X"
        else: row[2].text = "X"

    document.add_heading('5. SMS', 1)
    for d in ["NR-10", "Ficha EPI"] + dados.get('nrs_selecionadas', []): document.add_paragraph(d, style='List Bullet')
    
    document.add_heading('6. CRONOGRAMA', 1)
    if dados.get('data_inicio'): document.add_paragraph(f"Início: {dados['data_inicio'].strftime('%d/%m/%Y')}")
    if dados.get('data_fim'): document.add_paragraph(f"Término: {dados['data_fim'].strftime('%d/%m/%Y')}")

    num_secao = 7
    if dados.get('obs_gerais'): 
        document.add_heading(f'{num_secao}. OBSERVAÇÕES', 1); document.add_paragraph(dados['obs_gerais']); num_secao += 1 
    
    if dados.get('status') == "Contratação Finalizada":
        document.add_heading(f'{num_secao}. COMERCIAL', 1)
        p_val = document.add_paragraph(); p_val.add_run("Valor Global: ").bold = True
        p_val.add_run(f"{formatar_moeda_brl(dados.get('valor_total', ''))} (Fixo e irreajustável)")
        p_pgto = document.add_paragraph(); p_pgto.add_run("Condição de Pagamento: ").bold = True
        p_pgto.add_run(dados.get('condicao_pgto', ''))

    section = document.sections[0]; footer = section.footer
    for p in footer.paragraphs: p._element.getparent().remove(p._element)
    pf = footer.add_paragraph(); pf.text = "SIARCON Engenharia, controlando condições ambientais com excelência."
    pf.alignment = 1; pf.style.font.size = Pt(9); pf.style.font.italic = True
    document.add_paragraph("\n\n"); document.add_paragraph("_"*60); document.add_paragraph(f"DE ACORDO: {dados.get('fornecedor', '')}")
    b = io.BytesIO(); document.save(b); b.seek(0); return b

# --- FRONT ---
st.title("🤖 Escopo de Automação")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["1. Cadastro", "2. Técnico", "3. Matriz", "4. SMS", "5. Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        cli = st.text_input("Cliente", value=dados_edicao.get('Cliente', ''))
        obra = st.text_input("Obra", value=dados_edicao.get('Obra', ''))
        forn = st.text_input("Fornecedor", value=dados_edicao.get('Fornecedor', '')); 
        if not forn: forn = "PROPONENTE AUTOMAÇÃO"
        email_sup = st.text_input("E-mail Suprimentos", value="suprimentos@siarcon.com.br")
    with c2:
        c_r1, c_r2, c_r3 = st.columns(3)
        resp = c_r1.text_input("Resp. Eng.", value=dados_edicao.get('Responsável', ''))
        obr = c_r2.text_input("Resp. Obras", value=dados_edicao.get('Responsável Obras', ''))
        resp_sup = c_r3.text_
