import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time
import ast
from datetime import date, datetime
import utils_db

if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado."); st.stop()

DISCIPLINA_ATUAL = "TAB"
TEXTO_RESUMO_PADRAO = "Este escopo contempla o fornecimento de serviços de TAB / Comissionamento de sistemas, conforme detalhamento a seguir."

ITENS_MATRIZ = [
    "Instrumentação Calibrada (Balômetro/Anemômetro)", "Mão de Obra Especializada",
    "Relatórios Técnicos", "Balanceamento de Ar", "Balanceamento Hidrônico",
    "Testes de Estanqueidade de Dutos", "Medição de Ruído/Vibração",
    "Ajuste de Polias e Correias", "Start-up Assistido"
]

PADRAO_TECNICO = [
    "Medição e Ajuste de Vazão de Ar em Difusores/Grelhas", "Medição e Ajuste de Vazão de Ar em Caixas VAV",
    "Medição de Pressão Estática em Ventiladores", "Medição de Corrente e Tensão de Motores",
    "Ajuste de Rotação (Troca de Polias/Inversor)", "Balanceamento Hidrônico de Chillers e Fancoils",
    "Regulagem de Válvulas de Balanceamento (PICV/Estática)", "Medição de Diferencial de Pressão (Água/Ar)",
    "Teste de Fumaça em Dutos (Estanqueidade)", "Verificação de Setpoints de Temperatura/Umidade",
    "Medição de Nível de Ruído (dB)"
]

PADRAO_QUALIDADE = [
    "Certificados de Calibração dos Instrumentos (RBC)", "Relatório Fotográfico das Medições",
    "Etiquetagem dos Pontos Balanceados", "Comparativo Projeto x Executado",
    "Verificação de Fechamento de Forro", "Limpeza dos Filtros antes do TAB"
]

SMS_PADRAO_DOC = [
    "Ficha de registro", "ASO (Atestado de Saúde Ocupacional)", "Ficha de EPI", "Ordem de Serviço",
    "Certificados de Treinamento", "NR-06 (Equipamento de Proteção Individual)",
    "NR-12 (Segurança em Máquinas e Equipamentos)",
    "Comprovações de recolhimento de INSS, FGTS e folha de pagamento"
]

LISTA_NRS_SELECAO = [
    "NR-01 (Disposições Gerais)", "NR-03 (Embargo e Interdição)", "NR-04 (SESMT)", "NR-05 (CIPA)", 
    "NR-07 (PCMSO)", "NR-08 (Edificações)", "NR-09 (Avaliação e Controle de Exposições)", 
    "NR-10 (Eletricidade)", "NR-11 (Transporte e Movimentação)", "NR-13 (Vasos de Pressão)", 
    "NR-15 (Insalubridade)", "NR-16 (Periculosidade)", "NR-17 (Ergonomia)", "NR-18 (Construção Civil)", 
    "NR-19 (Explosivos)", "NR-20 (Inflamáveis)", "NR-21 (Trabalho a Céu Aberto)", "NR-23 (Incêndios)", 
    "NR-24 (Condições Sanitárias)", "NR-25 (Resíduos)", "NR-26 (Sinalização)", "NR-28 (Fiscalização)", 
    "NR-33 (Espaços Confinados)", "NR-35 (Trabalho em Altura)", "NR-38 (Limpeza Urbana)"
]

st.set_page_config(page_title=f"Escopo {DISCIPLINA_ATUAL}", page_icon="⚖️", layout="wide")

st.markdown("""
    <style>
    div[data-testid="stDownloadButton"] button { background-color: #28a745 !important; color: white !important; border-color: #28a745 !important; width: 100%; }
    div[data-testid="stDownloadButton"] button:hover { background-color: #218838 !important; border-color: #1e7e34 !important; }
    </style>
""", unsafe_allow_html=True)

if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()
cat_tecnica_db = f"tecnico_{DISCIPLINA_ATUAL.lower()}"
id_projeto = st.session_state.get('id_projeto_editar')
dados_edit = utils_db.buscar_projeto_por_id(id_projeto) if id_projeto else {}

def converter_para_estrutura(valor, tipo_esperado=list):
    """Função blindada para converter strings do DB com segurança e evitar TypeErrors"""
    if isinstance(valor, tipo_esperado):
        return valor
    if isinstance(valor, str) and valor.strip():
        try:
            res = ast.literal_eval(valor)
            if isinstance(res, tipo_esperado):
                return res
        except:
            pass
    return [] if tipo_esperado == list else {}

def formatar_moeda(valor):
    if not valor: return ""
    v_str = str(valor).replace('R$', '').strip()
    if not v_str: return ""
    try:
        if ',' in v_str and '.' in v_str: v_clean = v_str.replace('.', '').replace(',', '.')
        elif ',' in v_str: v_clean = v_str.replace(',', '.')
        elif '.' in v_str:
            parts = v_str.split('.')
            if len(parts[-1]) == 2: v_clean = v_str
            elif len(parts[-1]) == 3: v_clean = v_str.replace('.', '')
            else: v_clean = v_str
        else: v_clean = v_str
        v_float = float(v_clean)
        return f"R$ {v_float:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except: return f"R$ {valor}" if not str(valor).startswith("R$") else str(valor)

def gerar_docx(dados):
    doc = Document()
    try: style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    except: pass
    
    head = doc.add_heading(f'ESCOPO - {dados["disciplina"].upper()}', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('SIARCON ENGENHARIA')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True; sub.runs[0].font.size = Pt(20)
    
    doc.add_heading('1. DADOS DA OBRA', 1)
    table = doc.add_table(rows=9, cols=2); table.style = 'Table Grid'
    info_rows = [
        ("CLIENTE", dados['cliente']), ("OBRA", dados['obra']), ("FORNECEDOR", dados['fornecedor']),
        ("CNPJ FORNECEDOR", dados.get('cnpj_fornecedor', '-')), ("ENGENHARIA", dados['responsavel']), 
        ("OBRAS", dados.get('resp_obras', '')), ("SUPRIMENTOS", dados['resp_suprimentos']), 
        ("PROJETOS REFERÊNCIA", dados.get('projetos_referencia', '-')), 
        ("DATA / REVISÃO", f"{datetime.now().strftime('%d/%m/%Y')}  |  Rev: {dados.get('revisao','-')}")
    ]
    for idx, (label, value) in enumerate(info_rows):
        row = table.rows[idx]
        row.cells[0].text = label; row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)

    doc.add_heading('2. ESCOPO TÉCNICO', 1)
    doc.add_paragraph(dados.get('resumo_escopo', ''))
    if dados.get('tecnico_livre'):
        p = doc.add_paragraph(); p.add_run("OBSERVAÇÕES GERAIS:").bold = True
        doc.add_paragraph(dados['tecnico_livre'])
    p = doc.add_paragraph(); p.add_run("DETALHAMENTO:").bold = True
    comentarios = dados.get('comentarios_itens', {})
    for item in dados.get('itens_tecnicos', []):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item); run.bold = True
        if comentarios.get(item): p.add_run(f": {comentarios[item]}")

    num_secao = 3
    if dados.get('itens_qualidade'):
        doc.add_heading(f'{num_secao}. PADRÃO DE QUALIDADE', 1)
        for item in dados['itens_qualidade']: doc.add_paragraph(item, style='List Bullet')
        num_secao += 1

    doc.add_heading(f'{num_secao}. MATRIZ DE RESPONSABILIDADES', 1)
    tm = doc.add_table(rows=1, cols=4); tm.style = 'Table Grid'
    h = tm.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = "FORNECEDOR"; h[3].text = "FORA DO ESCOPO"
    for k, v in dados.get('matriz', {}).items():
        row = tm.add_row().cells; row[0].text = k
        row[1].text = "X" if v == "SIARCON" else ""; row[2].text = "X" if v == "FORNECEDOR" else ""
        row[3].text = "X" if v == "FORA DO ESCOPO" else ""
    num_secao += 1

    doc.add_heading(f'{num_secao}. SEGURANÇA (SMS)', 1)
    for i in SMS_PADRAO_DOC: doc.add_paragraph(i, style='List Bullet')
    for nr in dados.get('nrs_selecionadas', []): doc.add_paragraph(nr, style='List Bullet')
    if dados.get('sms_livre'): doc.add_paragraph(dados['sms_livre'])
    num_secao += 1

    doc.add_heading(f'{num_secao}. COMERCIAL', 1)
    if dados.get('valor_total'): doc.add_paragraph(f"Valor Global: {dados['valor_total']} (valor fixo e irreajustável)")
    if dados.get('valor_siarcon'): doc.add_paragraph(f"Faturamento SIARCON: {dados['valor_siarcon']}")
    if dados.get('flag_fat_direto') and dados.get('valor_direto'): doc.add_paragraph(f"Faturamento Direto: {dados['valor_direto']}")
    doc.add_paragraph(f"Condição de Pagamento: {dados.get('condicao_pgto','')}")
    if dados.get('obs_gerais'): doc.add_paragraph(f"Obs: {dados['obs_gerais']}")

    b = io.BytesIO(); doc.save(b); b.seek(0); return b

st.title(f"⚖️ {DISCIPLINA_ATUAL}")
if dados_edit: st.info(f"Editando: {dados_edit.get('obra')} | Cliente: {dados_edit.get('cliente')}")
opcoes = st.session_state.get('opcoes_db', {})

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Cadastro", "Técnico", "Matriz", "SMS", "Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    cliente = c1.text_input("Cliente", value=dados_edit.get('cliente', ''))
    obra = c1.text_input("Obra", value=dados_edit.get('obra', ''))
    db_forn = utils_db.listar_fornecedores(); lista_nomes = [""] + [f['Fornecedor'] for f in db_forn]
    val_forn_db = dados_edit.get('fornecedor', ''); idx_f = lista_nomes.index(val_forn_db) if val_forn_db in lista_nomes else 0
    sel_forn = c1.selectbox("Fornecedor (DB):", lista_nomes, index=idx_f)
    forn = c1.text_input("Razão Social (Final):", value=sel_forn if sel_forn else val_forn_db)
    cnpj = c1.text_input("CNPJ:", value=dados_edit.get('cnpj_fornecedor', ''))
    resp_eng = c2.text_input("Engenharia SIARCON", value=dados_edit.get('responsavel', ''))
    resp_obras = c2.text_input("Obras SIARCON", value=dados_edit.get('resp_obras', ''))
    resp_sup = c2.text_input("Suprimentos SIARCON", value=dados_edit.get('resp_suprimentos', ''))
    revisao = c2.text_input("Revisão", value=dados_edit.get('revisao', 'R-00'))
    st.divider(); c2.write("📂 **Projetos de Referência**")
    val_proj_salvo = dados_edit.get('projetos_referencia', '')
    set_arquivos = set([x.strip() for x in val_proj_salvo.split('\n') if x.strip()]) if val_proj_salvo else set()
    uploads = c2.file_uploader("Arraste arquivos aqui", accept_multiple_files=True)
    if uploads:
        for f in uploads: set_arquivos.add(f.name)
    projetos_ref = c2.text_area("Lista de Projetos:", value="\n".join(sorted(list(set_arquivos))), height=100)

with tab2:
    resumo = st.text_area("Resumo:", value=dados_edit.get('resumo_escopo', '') or TEXTO_RESUMO_PADRAO, height=80)
    st.divider()
    c_a1, c_a2 = st.columns([4,1])
    novo_item = c_a1.text_input("Novo Item DB (Técnico):", key="add_db_tec")
    lista_tec_final = [str(x) for x in sorted(list(set(opcoes.get(cat_tecnica_db, []) + PADRAO_TECNICO))) if x]
    
    if c_a2.button("💾 Criar Téc."): 
        if novo_item and utils_db.aprender_novo_item(cat_tecnica_db, novo_item): 
            st.session_state['opcoes_db'] = utils_db.carregar_opcoes(); st.toast("✅ Salvo no banco!")
            
    itens_salvos = converter_para_estrutura(dados_edit.get('itens_tecnicos', []), list)
    itens_salvos_limpos = [str(x) for x in itens_salvos if x]
    
    opcoes_multiselect = sorted(list(set(lista_tec_final + itens_salvos_limpos)))
    itens_tec = st.multiselect("Itens do Escopo:", opcoes=opcoes_multiselect, default=itens_salvos_limpos)
    
    comentarios_salvos = converter_para_estrutura(dados_edit.get('comentarios_itens', {}), dict)
    comentarios_novos = {i: st.text_input(f"Detalhe '{i}':", value=comentarios_salvos.get(i, "")) for i in itens_tec}
    
    st.divider()
    tec_livre = st.text_area("Observações Gerais:", value=dados_edit.get('tecnico_livre', ''))
    st.divider(); st.markdown("#### Qualidade")
    c_q1, c_q2 = st.columns([4,1])
    novo_item_q = c_q1.text_input("Novo Item DB (Qualidade):", key="add_db_qual")
    lista_qual = [str(x) for x in sorted(list(set(opcoes.get(f"qualidade_{DISCIPLINA_ATUAL.lower()}", []) + PADRAO_QUALIDADE))) if x]
    
    if c_q2.button("💾 Criar Qual."): 
        if novo_item_q and utils_db.aprender_novo_item(f"qualidade_{DISCIPLINA_ATUAL.lower()}", novo_item_q): 
            st.session_state['opcoes_db'] = utils_db.carregar_opcoes(); st.toast("✅ Qualidade salva no banco!")
            
    itens_salvos_q = converter_para_estrutura(dados_edit.get('itens_qualidade', []), list)
    itens_salvos_q_limpos = [str(x) for x in itens_salvos_q if x]
    
    opcoes_qual_multiselect = sorted(list(set(lista_qual + itens_salvos_q_limpos)))
    itens_qual = st.multiselect("Itens Qualidade:", sorted(list(set(opcoes_qual_multiselect))), default=itens_salvos_q_limpos)

with tab3:
    escolhas = {}
    matriz_salva = converter_para_estrutura(dados_edit.get('matriz', {}), dict)
    st.write("Responsabilidades:")
    for item in ITENS_MATRIZ:
        c_m1, c_m2 = st.columns([2, 3])
        c_m1.write(f"**{item}**")
        opts = ["SIARCON", "FORNECEDOR", "FORA DO ESCOPO"]
        idx = opts.index(matriz_salva.get(item, "SIARCON")) if matriz_salva.get(item) in opts else 0
        escolhas[item] = c_m2.radio(f"r_{item}", opts, index=idx, horizontal=True, label_visibility="collapsed")
        st.divider()

with tab4:
    nrs_salvas = converter_para_estrutura(dados_edit.get('nrs_selecionadas', []), list)
    nrs_salvas_limpas = [str(x) for x in nrs_salvas if x]
    
    opcoes_nrs_select = sorted(list(set(LISTA_NRS_SELECAO + nrs_salvas_limpas)))
    nrs = st.multiselect("NRs Adicionais:", opcoes=opcoes_nrs_select, default=nrs_salvas_limpas)
    sms_livre = st.text_area("Outras exigências:", value=dados_edit.get('sms_livre', ''))

with tab5:
    v_total_db = str(dados_edit.get('valor_total', '')).replace('R$', '').strip()
    val = st.text_input("Valor Global (R$):", value=v_total_db)
    v_siarcon_db = str(dados_edit.get('valor_siarcon', '')).replace('R$', '').strip()
    val_siarcon = st.text_input("Valor Faturamento SIARCON (R$):", value=v_siarcon_db)
    flag_fat_direto = st.checkbox("Opção de faturamento direto", value=dados_edit.get('flag_fat_direto', False))
    val_direto = ""
    if flag_fat_direto:
        v_direto_db = str(dados_edit.get('valor_direto', '')).replace('R$', '').strip()
        val_direto = st.text_input("Valor Previsto Faturamento Direto (R$):", value=v_direto_db)
    pgto = st.text_area("Condição Pagamento:", value=dados_edit.get('condicao_pgto', ''))
    obs = st.text_area("Obs Comerciais:", value=dados_edit.get('obs_gerais', ''))
    lista_st = ["Não Iniciado", "Engenharia", "Obras", "Suprimentos", "Finalizado"]
    status_atual = dados_edit.get('status')
    idx_status = lista_st.index(status_atual) if status_atual in lista_st else 0
    status = st.selectbox("Status:", lista_st, index=idx_status)

st.markdown("---")
dados = {
    '_id': dados_edit.get('_id'), 'disciplina': DISCIPLINA_ATUAL, 'cliente': cliente, 'obra': obra, 
    'fornecedor': forn, 'cnpj_fornecedor': cnpj, 'responsavel': resp_eng, 
    'resp_obras': resp_obras, 'resp_suprimentos': resp_sup, 'revisao': revisao, 'projetos_referencia': projetos_ref, 
    'resumo_escopo': resumo, 'itens_tecnicos': itens_tec, 'comentarios_itens': comentarios_novos, 
    'tecnico_livre': tec_livre, 'itens_qualidade': itens_qual, 'matriz': escolhas, 'nrs_selecionadas': nrs, 
    'sms_livre': sms_livre, 'valor_total': formatar_moeda(val), 'valor_siarcon': formatar_moeda(val_siarcon),
    'flag_fat_direto': flag_fat_direto, 'valor_direto': formatar_moeda(val_direto) if flag_fat_direto else "",
    'condicao_pgto': pgto, 'obs_gerais': obs, 'status': status, 'data_inicio': dados_edit.get('data_inicio', date.today().strftime("%Y-%m-%d"))
}
col_b1, col_b2 = st.columns(2)
with col_b1:
    if st.button("☁️ APENAS SALVAR NO DB"):
        if utils_db.registrar_projeto(dados): st.success("Salvo no banco de dados!")
with col_b2:
    if st.button("💾 SALVAR E PREPARAR DOCX", type="primary"):
        if utils_db.registrar_projeto(dados):
            st.success("Salvo! Baixe abaixo."); st.session_state[f'btn_docx_{DISCIPLINA_ATUAL}'] = True
    if st.session_state.get(f'btn_docx_{DISCIPLINA_ATUAL}', False):
        b = gerar_docx(dados)
        st.download_button("📥 BAIXAR DOCX GERADO", b, file_name=f"Escopo_{DISCIPLINA_ATUAL}.docx")
