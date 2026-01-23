import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
from datetime import date, timedelta
import utils_db

# --- CONFIGURAÇÃO ---
st.set_page_config(page_title="Escopo TAB | SIARCON", page_icon="💨", layout="wide")

# --- FUNÇÕES ---
def adicionar_item_callback(categoria, key_input):
    novo = st.session_state.get(key_input, "")
    if novo:
        if utils_db.aprender_novo_item(categoria, novo) is True:
            st.session_state[key_input] = ""
            if 'opcoes_db' in st.session_state: del st.session_state['opcoes_db']
            st.toast(f"✅ Item salvo!", icon="✅")

def atualizar_anexos():
    arquivos = st.session_state.get("uploader_anexos", [])
    if arquivos: st.session_state["input_proj_ref"] = "; ".join([f.name for f in arquivos])

# --- EDIÇÃO ---
dados_edicao = {}
id_linha_edicao = None
if 'modo_edicao' in st.session_state and st.session_state['modo_edicao']:
    dados_edicao = st.session_state.get('dados_projeto', {})
    id_linha_edicao = dados_edicao.get('_id_linha')
    if st.button("❌ Cancelar Edição"):
        st.session_state['modo_edicao'] = False; st.rerun()

if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

# --- GERAR DOCX ---
def gerar_docx(dados):
    document = Document()
    try:
        style = document.styles['Normal']; font = style.font; font.name = 'Calibri'; font.size = Pt(11)
    except: pass

    # Cabeçalho
    section = document.sections[0]; header = section.header
    for p in header.paragraphs: p._element.getparent().remove(p._element)
    p = header.add_paragraph()
    p.text = "Departamento de Operações SIARCON"; p.alignment = 1; p.style.font.bold = True; p.style.font.size = Pt(14)

    document.add_paragraph("\n")
    document.add_heading('Escopo de fornecimento - TAB (Testes e Balanceamento)', 0).alignment = 1
    document.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')} | Rev: {dados['revisao']}").alignment = 2

    document.add_heading('1. OBJETIVO E RESUMO', 1)
    table = document.add_table(rows=6, cols=2); 
    try: table.style = 'Table Grid'
    except: pass
    infos = [("Cliente:", dados['cliente']), ("Local/Obra:", dados['obra']), ("Projetos Ref:", dados['projetos_ref']), ("Fornecedor:", dados['fornecedor']), ("Resp. Engenharia:", dados['responsavel']), ("Resp. Obras:", dados['resp_obras'])]
    for i, (k, v) in enumerate(infos): row = table.rows[i]; row.cells[0].text = k; row.cells[0].paragraphs[0].runs[0].bold = True; row.cells[1].text = v
    
    document.add_paragraph(f"\nResumo: {dados['resumo_escopo']}")

    document.add_heading('2. TÉCNICO', 1)
    for item in dados['itens_tecnicos']: document.add_paragraph(item, style='List Bullet')
    if dados['tecnico_livre']: document.add_paragraph(dados['tecnico_livre'], style='List Bullet')

    document.add_heading('3. QUALIDADE', 1)
    for item in dados['itens_qualidade']: document.add_paragraph(item, style='List Bullet')
    if dados['qualidade_livre']: document.add_paragraph(dados['qualidade_livre'], style='List Bullet')

    document.add_heading('4. MATRIZ RESPONSABILIDADES', 1)
    t_m = document.add_table(rows=1, cols=3); 
    try: t_m.style = 'Table Grid'
    except: pass
    h = t_m.rows[0].cells; h[0].text = "ITEM"; h[1].text = "SIARCON"; h[2].text = dados['fornecedor'].upper()
    for item, resp in dados['matriz'].items():
        row = t_m.add_row().cells; row[0].text = item
        if resp == "SIARCON": row[1].text = "X"; row[1].paragraphs[0].alignment = 1
        else: row[2].text = "X"; row[2].paragraphs[0].alignment = 1

    document.add_heading('5. SMS', 1)
    docs = ["Ficha de registro", "ASO", "Ficha EPI", "Ordem de Serviço", "Certificados de Treinamento", "NR-35 (Trabalho em Altura)"]
    for d in docs: document.add_paragraph(d, style='List Bullet')
    for d in dados['nrs_selecionadas']: document.add_paragraph(d, style='List Bullet')

    document.add_heading('6. CRONOGRAMA', 1)
    document.add_paragraph(f"Início: {dados['data_inicio'].strftime('%d/%m/%Y')}")
    if dados.get('data_fim'): document.add_paragraph(f"Término: {dados['data_fim'].strftime('%d/%m/%Y')}")

    if dados['obs_gerais']: document.add_heading('7. OBSERVAÇÕES', 1); document.add_paragraph(dados['obs_gerais'])
    if dados['status'] == "Contratação Finalizada":
        document.add_heading('8. COMERCIAL', 1); document.add_paragraph(f"Total: {dados['valor_total']} | Pagamento: {dados['condicao_pgto']}")
        if dados['info_comercial']: document.add_paragraph(dados['info_comercial'])
    
    section = document.sections[0]; footer = section.footer
    for p in footer.paragraphs: p._element.getparent().remove(p._element)
    pf = footer.add_paragraph(); pf.text = "SIARCON - Excelência em TAB"; pf.alignment = 1; pf.style.font.size = Pt(9); pf.style.font.italic = True
    
    document.add_paragraph("\n\n"); document.add_paragraph("_"*60); document.add_paragraph(f"DE ACORDO: {dados['fornecedor']}")
    
    b = io.BytesIO(); document.save(b); b.seek(0); return b

# --- INTERFACE ---
st.title("💨 Escopo de TAB")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["1. Cadastro", "2. Técnico", "3. Matriz", "4. SMS", "5. Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        cliente = st.text_input("Cliente", value=dados_edicao.get('Cliente', ''))
        obra = st.text_input("Obra", value=dados_edicao.get('Obra', ''))
        forn = st.text_input("Fornecedor", value=dados_edicao.get('Fornecedor', ''), placeholder="Padrão: PROPONENTE TAB")
        if not forn: forn = "PROPONENTE TAB"
    with c2:
        c_r1, c_r2 = st.columns(2)
        resp_eng = c_r1.text_input("Resp. Eng.", value=dados_edicao.get('Responsável', ''))
        resp_obr = c_r2.text_input("Resp. Obras", value=dados_edicao.get('Responsável Obras', ''))
        revisao = st.text_input("Revisão", "R-00")
        if "input_proj_ref" not in st.session_state: st.session_state["input_proj_ref"] = dados_edicao.get('projetos_ref', '')
        proj_ref = st.text_input("Projetos Ref.", key="input_proj_ref")
        email = st.text_input("E-mail:", "suprimentos@siarcon.com.br")
    
    resumo = st.text_area("Resumo")
    anexos = st.file_uploader("Anexos", accept_multiple_files=True, key="uploader_anexos", on_change=atualizar_anexos)

with tab2:
    st.subheader("Técnico (TAB)")
    padroes_tec = [
        "Medição de vazão e pressão em todos os difusores e grelhas",
        "Ajuste de rotação de ventiladores e troca de polias se necessário",
        "Balanceamento hidrônico de água gelada (Válvulas de balanceamento)",
        "Medição de grandezas elétricas (Tensão e Corrente) dos equipamentos",
        "Emissão de relatório técnico com comparativo Projeto x Realizado",
        "Utilização de instrumentos calibrados (Certificado RBC)",
        "Marcação física dos pontos de ajuste (Lacre)"
    ]
    # Usa chave generica ou cria nova no utils_db se quiser
    opcoes = list(set(padroes_tec + st.session_state['opcoes_db'].get('tecnico', [])))
    itens_tec = st.multiselect("Selecione:", opcoes)
    
    c_add, c_free = st.columns(2)
    c_add.text_input("Novo Item", key="n_tec"); c_add.button("Salvar", on_click=adicionar_item_callback, args=("tecnico", "n_tec"))
    tec_livre = c_free.text_area("Livre")

with tab3:
    escolhas = {}
    itens_m = ["Instrumentos de medição (Balometer, Manômetro, etc)", "Ferramentas manuais", "Polias e Correias (se necessário troca)", "Escadas e Andaimes", "Emissão de Relatórios", "Alimentação/Viagem", "EPIs/Uniformes"]
    nome_m = forn.upper()
    st.info(f"Matriz: {nome_m}")
    for i in itens_m:
        c1, c2 = st.columns([3,2])
        c1.write(f"**{i}**")
        escolhas[i] = c2.radio(f"m_{i}", ["SIARCON", nome_m], horizontal=True, label_visibility="collapsed")
        st.divider()

with tab4:
    # Apenas NR-35 e o padrao
    nrs = st.multiselect("SMS Adicional:", st.session_state['opcoes_db'].get('sms', []))
    c_d1, c_d2 = st.columns(2)
    d_ini = c_d1.date_input("Início"); d_int = c_d2.number_input("Dias Integração", 5)
    usar_fim = st.checkbox("Data Fim?", True)
    d_fim = st.date_input("Fim", date.today()+timedelta(days=10)) if usar_fim else None

    # Qualidade simplificada aqui
    st.subheader("Qualidade")
    padrao_qual = ["Instrumentos com certificado de calibração válido", "Relatórios seguindo norma ABNT/SMACNA"]
    itens_qual = st.multiselect("Itens:", padrao_qual)
    qual_livre = st.text_input("Qualidade Livre")

with tab5:
    val = st.text_input("Valor Total", dados_edicao.get('Valor', '')); pgto = st.text_area("Pagamento"); info = st.text_input("Info"); obs = st.text_area("Obs")
    st.divider()
    status = st.selectbox("Status", ["Em Elaboração (Engenharia)", "Aguardando Obras", "Recebido (Suprimentos)", "Enviado para Cotação", "Em Negociação", "Contratação Finalizada"], index=0)

st.markdown("---")
if status == "Contratação Finalizada" and 'modo_edicao' in st.session_state:
    st.error("🔒 Finalizado."); st.download_button("📥 Baixar", gerar_docx(dados_edicao).getvalue(), f"Escopo_{forn}.docx")
else:
    if st.button("💾 SALVAR / ATUALIZAR", type="primary"):
        dados = {
            'cliente': cliente, 'obra': obra, 'fornecedor': forn, 'responsavel': resp_eng, 'resp_obras': resp_obr,
            'revisao': revisao, 'projetos_ref': proj_ref, 'resumo_escopo': resumo,
            'itens_tecnicos': itens_tec, 'tecnico_livre': tec_livre,
            'itens_qualidade': itens_qual, 'qualidade_livre': qual_livre,
            'matriz': escolhas, 'nrs_selecionadas': nrs,
            'data_inicio': d_ini, 'dias_integracao': d_int, 'data_fim': d_fim,
            'obs_gerais': obs, 'valor_total': val, 'condicao_pgto': pgto, 'info_comercial': info,
            'status': status, 'disciplina': 'TAB', # <--- IMPORTANTE
            'nomes_anexos': [f.name for f in anexos] if anexos else []
        }
        
        docx = gerar_docx(dados); nome_a = f"Escopo_{forn.replace(' ', '_')}.docx"
        utils_db.registrar_projeto(dados, id_linha_edicao)
        st.success("✅ Salvo!"); st.download_button("📥 Baixar DOCX", docx.getvalue(), nome_a)
