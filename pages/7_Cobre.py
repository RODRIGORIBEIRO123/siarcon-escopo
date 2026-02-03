import streamlit as st
from docx import Document
from docx.shared import Pt
import io
from datetime import date
import utils_db

if 'logado' not in st.session_state or not st.session_state['logado']:
    st.warning("🔒 Acesso negado."); st.stop()

DISCIPLINA_ATUAL = "Cobre"
TEXTO_RESUMO_PADRAO = "Este escopo contempla o fornecimento de instalações frigorígenas, conforme detalhamento a seguir."

LISTA_NRS_COMPLETA = [
    "NR-01 (Disposições Gerais)", "NR-03 (Embargo e Interdição)", "NR-04 (SESMT)", "NR-05 (CIPA)", 
    "NR-06 (EPI)", "NR-07 (PCMSO)", "NR-08 (Edificações)", "NR-09 (Avaliação e Controle de Exposições)", 
    "NR-10 (Eletricidade)", "NR-11 (Transporte e Movimentação)", "NR-12 (Máquinas e Equipamentos)", 
    "NR-13 (Vasos de Pressão)", "NR-15 (Insalubridade)", "NR-16 (Periculosidade)", "NR-17 (Ergonomia)", 
    "NR-18 (Construção Civil)", "NR-19 (Explosivos)", "NR-20 (Inflamáveis)", "NR-21 (Trabalho a Céu Aberto)", 
    "NR-23 (Incêndios)", "NR-24 (Condições Sanitárias)", "NR-25 (Resíduos)", "NR-26 (Sinalização)", 
    "NR-28 (Fiscalização)", "NR-33 (Espaços Confinados)", "NR-35 (Trabalho em Altura)", "NR-38 (Limpeza Urbana)"
]

ITENS_MATRIZ = [
    "Tubulação de Cobre", "Isolamento Térmico (Elastomérico)",
    "Solda e Consumíveis (PPU/Prata)", "Carga de Gás Refrigerante",
    "Nitrogênio para Testes", "Suportação da Rede", "Refnets e Derivações"
]

PADRAO_TECNICO = [
    "Instalação de tubulação de cobre rígido/flexível", "Brasagem com fluxo de nitrogênio passante",
    "Instalação de Refnets (VRF)", "Aplicação de isolamento térmico blindado UV (Externo)",
    "Aplicação de isolamento térmico elastomérico (Interno)", "Fixação com braçadeiras e isoladores",
    "Interligação das unidades evaporadoras/condensadoras", "Sifões e liras de dilatação (conforme fabricante)",
    "Desidratação do sistema (Vácuo < 500 microns)", "Carga adicional de fluido refrigerante"
]

PADRAO_QUALIDADE = [
    "Teste de Estanqueidade (Pressurização N2 - 24h)", "Relatório de Vácuo (Vacuômetro digital)",
    "Cálculo de Carga Adicional (Software Fabricante)", "Inspeção visual das soldas",
    "Verificação de espessura do isolamento"
]

st.set_page_config(page_title="Escopo Cobre", page_icon="❄️", layout="wide")
if 'opcoes_db' not in st.session_state: st.session_state['opcoes_db'] = utils_db.carregar_opcoes()

id_projeto = st.session_state.get('id_projeto_editar')
dados_edit = {}
if id_projeto:
    t = utils_db.buscar_projeto_por_id(id_projeto)
    if t: dados_edit = t

def formatar_moeda(valor):
    try:
        v = float(str(valor).replace('R$', '').replace('.', '').replace(',', '.').strip())
        return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except: return valor

def gerar_docx(dados):
    doc = Document()
    try: style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    except: pass
    
    doc.add_heading(f'Escopo - {dados["disciplina"]}', 0)
    doc.add_paragraph(f"Rev: {dados.get('revisao','-')}")
    doc.add_heading('1. DADOS', 1)
    t = doc.add_table(rows=1, cols=2)
    infos = [("Cliente", dados['cliente']), ("Obra", dados['obra']), ("Fornecedor", dados['fornecedor']),
             ("Engenharia", dados['responsavel']), ("Suprimentos", dados['resp_suprimentos'])]
    for k, v in infos:
        row = t.add_row().cells; row[0].text = k; row[1].text = str(v)

    doc.add_heading('2. TÉCNICO', 1)
    doc.add_paragraph(f"Resumo: {dados.get('resumo_escopo','')}")
    if dados.get('tecnico_livre'): doc.add_paragraph(dados['tecnico_livre'])
    for item in dados.get('itens_tecnicos', []): doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. QUALIDADE', 1)
    for item in dados.get('itens_qualidade', []): doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4. MATRIZ DE RESPONSABILIDADES', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "ITEM"; hdr[1].text = "SIARCON"; hdr[2].text = "FORNECEDOR"
    for k, v in dados.get('matriz', {}).items():
        row = table.add_row().cells
        row[0].text = k
        row[1].text = "X" if v == "SIARCON" else ""
        row[2].text = "X" if v != "SIARCON" else ""

    doc.add_heading('5. SMS', 1)
    if dados.get('sms_livre'): doc.add_paragraph(dados['sms_livre'])
    for nr in dados.get('nrs_selecionadas', []): doc.add_paragraph(nr, style='List Bullet')

    doc.add_heading('6. COMERCIAL', 1)
    doc.add_paragraph(f"Valor: {formatar_moeda(dados.get('valor_total',''))}")
    doc.add_paragraph(f"Pagamento: {dados.get('condicao_pgto','')}")
    if dados.get('obs_gerais'): doc.add_paragraph(f"Obs: {dados['obs_gerais']}")

    b = io.BytesIO(); doc.save(b); b.seek(0); return b

st.title(f"❄️ {DISCIPLINA_ATUAL}")
if dados_edit: st.info(f"Editando: {dados_edit.get('obra')} | Cliente: {dados_edit.get('cliente')}")
opcoes = st.session_state.get('opcoes_db', {})

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Cadastro", "Técnico", "Matriz", "SMS", "Comercial"])

with tab1:
    c1, c2 = st.columns(2)
    cliente = c1.text_input("Cliente", value=dados_edit.get('cliente', ''))
    obra = c1.text_input("Obra", value=dados_edit.get('obra', ''))
    db_forn = utils_db.listar_fornecedores(); lista_nomes = [""] + [f['Fornecedor'] for f in db_forn]
    val_forn_db = dados_edit.get('fornecedor', ''); idx_f = lista_nomes.index(val_forn_db) if val_forn_db in lista_nomes else 0
    sel_forn = c1.selectbox("Fornecedor (DB):", lista_nomes, index=idx_f)
    forn = c1.text_input("Razão Social:", value=sel_forn if sel_forn else val_forn_db)
    cnpj = c1.text_input("CNPJ:", value=dados_edit.get('cnpj_fornecedor', ''))
    resp_eng = c2.text_input("Engenharia", value=dados_edit.get('responsavel', ''))
    resp_sup = c2.text_input("Suprimentos", value=dados_edit.get('resp_suprimentos', ''))
    revisao = c2.text_input("Revisão", value=dados_edit.get('revisao', 'R-00'))
    val_resumo = dados_edit.get('resumo_escopo', TEXTO_RESUMO_PADRAO)
    resumo = c2.text_area("Resumo", value=val_resumo, height=100)

with tab2:
    lista_tec_final = sorted(list(set(opcoes.get(f"tecnico_{DISCIPLINA_ATUAL.lower()}", []) + PADRAO_TECNICO)))
    itens_salvos = dados_edit.get('itens_tecnicos', [])
    if isinstance(itens_salvos, str): 
        try: itens_salvos = eval(itens_salvos)
        except: itens_salvos = []
    elif not isinstance(itens_salvos, list): itens_salvos = []
    opcoes_finais = sorted(list(set(lista_tec_final + itens_salvos)))
    itens_tec = st.multiselect("Itens Técnicos:", opcoes_finais, default=itens_salvos)
    tec_livre = st.text_area("Livre Técnico:", value=dados_edit.get('tecnico_livre', ''))
    st.divider()
    lista_qual_final = sorted(list(set(opcoes.get(f"qualidade_{DISCIPLINA_ATUAL.lower()}", []) + PADRAO_QUALIDADE)))
    itens_salvos_q = dados_edit.get('itens_qualidade', [])
    if isinstance(itens_salvos_q, str): 
        try: itens_salvos_q = eval(itens_salvos_q)
        except: itens_salvos_q = []
    elif not isinstance(itens_salvos_q, list): itens_salvos_q = []
    opcoes_finais_q = sorted(list(set(lista_qual_final + itens_salvos_q)))
    itens_qual = st.multiselect("Itens Qualidade:", opcoes_finais_q, default=itens_salvos_q)

with tab3:
    escolhas = {}
    matriz_salva = dados_edit.get('matriz', {})
    if isinstance(matriz_salva, str): 
        try: matriz_salva = eval(matriz_salva)
        except: matriz_salva = {}
    elif not isinstance(matriz_salva, dict): matriz_salva = {}
    for item in ITENS_MATRIZ:
        col_a, col_b = st.columns([2,1])
        col_a.write(f"**{item}**")
        val = 1 if (item in matriz_salva and matriz_salva[item] != "SIARCON") else 0
        escolhas[item] = col_b.radio(item, ["SIARCON", "FORNECEDOR"], index=val, horizontal=True, label_visibility="collapsed", key=f"m_{item}")
        st.divider()

with tab4:
    nrs_salvas = dados_edit.get('nrs_selecionadas', [])
    if isinstance(nrs_salvas, str): 
        try: nrs_salvas = eval(nrs_salvas)
        except: nrs_salvas = []
    elif not isinstance(nrs_salvas, list): nrs_salvas = []
    opcoes_sms = sorted(list(set(LISTA_NRS_COMPLETA + nrs_salvas)))
    nrs = st.multiselect("NRs:", opcoes_sms, default=nrs_salvas)
    sms_livre = st.text_area("Livre SMS:", value=dados_edit.get('sms_livre', ''))

with tab5:
    val = st.text_input("Valor", value=dados_edit.get('valor_total', ''))
    pgto = st.text_area("Pgto", value=dados_edit.get('condicao_pgto', ''))
    obs = st.text_area("Obs", value=dados_edit.get('obs_gerais', ''))
    lista_st = ["Em Elaboração", "Em Análise Obras", "Em Cotação", "Finalizado", "Concluído"]
    st_at = dados_edit.get('status', 'Em Elaboração')
    idx_st = lista_st.index(st_at) if st_at in lista_st else 0
    status = st.selectbox("Status", lista_st, index=idx_st)

st.markdown("---")
dados = {
    '_id': dados_edit.get('_id'), 'disciplina': DISCIPLINA_ATUAL, 'cliente': cliente, 'obra': obra, 
    'fornecedor': forn, 'cnpj_fornecedor': cnpj, 'responsavel': resp_eng, 'resp_suprimentos': resp_sup, 
    'revisao': revisao, 'resumo_escopo': resumo, 'itens_tecnicos': itens_tec, 'tecnico_livre': tec_livre, 
    'itens_qualidade': itens_qual, 'matriz': escolhas, 'nrs_selecionadas': nrs, 'sms_livre': sms_livre, 
    'valor_total': val, 'condicao_pgto': pgto, 'obs_gerais': obs, 'status': status, 
    'data_inicio': dados_edit.get('data_inicio', date.today().strftime("%Y-%m-%d"))
}

c1, c2 = st.columns(2)
if c1.button("☁️ SALVAR"): utils_db.registrar_projeto(dados); st.success("Salvo!")
if c2.button("💾 DOCX", type="primary"): 
    utils_db.registrar_projeto(dados); b = gerar_docx(dados)
    st.download_button("Baixar DOCX", b, f"Escopo_{DISCIPLINA_ATUAL}.docx")
